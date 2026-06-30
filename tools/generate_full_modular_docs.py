from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\RIESGO_LAVADO")
OUT = ROOT / "docs" / "5. Documentacion Modular"
CAPTURES = ROOT / "docs" / "5. Documentacion Modular" / "capturas"
DATE = "30/06/2026"
VERSION = "1.0-full"

spec = importlib.util.spec_from_file_location("base_docs", ROOT / "tools" / "generate_modular_docs.py")
base = importlib.util.module_from_spec(spec)
spec.loader.exec_module(base)  # type: ignore[union-attr]


MODULES = [
    {
        "id": 2,
        "slug": "Usuarios",
        "name": "Seguridad y Usuarios",
        "route": "/usuarios",
        "capture": "01_usuarios.png",
        "frontend": [
            "features/admin/usuarios/usuarios.component.ts/html",
            "core/services/auth.service.ts",
            "core/guards/auth.guard.ts, modulo.guard.ts, role.guard.ts",
        ],
        "endpoints": [
            ("GET", "/api/auth/usuarios", "Lista usuarios"),
            ("POST", "/api/auth/usuarios", "Crear usuario"),
            ("PUT", "/api/auth/usuarios/{uid}", "Editar usuario"),
            ("PUT", "/api/auth/usuarios/{uid}/estado", "Cambiar estado"),
            ("GET", "/api/auth/validar-dominio", "Validar Active Directory"),
            ("GET", "/api/catalogos/roles", "Catalogo roles"),
            ("GET", "/api/catalogos/dominios", "Catalogo dominios"),
            ("GET", "/api/catalogos/modulos", "Catalogo modulos"),
        ],
        "tables": ["RL_USUARIOS", "RL_ROLES", "RL_DOMINIO", "RL_MODULOS", "RL_USUARIO_MODULOS", "RL_AUDITORIA"],
        "flow": [
            "Ingresar con usuario administrador.",
            "Consultar listado de usuarios.",
            "Crear o editar datos personales, rol, dominio y modulos asignados.",
            "Validar usuario de dominio cuando aplique.",
            "Guardar cambios y revisar auditoria generada.",
        ],
        "rules": [
            "Debe seleccionarse al menos un modulo para guardar usuario.",
            "Usuarios de dominio requieren dominio y usuario AD valido.",
            "El boton guardar se bloquea cuando el formulario no procede.",
            "Usuarios inactivos no deben renovar sesion por refresh token.",
        ],
        "audits": ["Creacion de usuario", "Edicion de usuario", "Cambio de estado", "Cambio/recuperacion de contrasena"],
    },
    {
        "id": 3,
        "slug": "Configuracion",
        "name": "Configuracion del Sistema",
        "route": "/configuracion",
        "capture": "02_configuracion.png",
        "frontend": ["features/admin/configuracion/configuracion.component.ts", "core/services/configuracion.service.ts"],
        "endpoints": [
            ("GET", "/api/configuracion/sistema", "Obtener configuracion"),
            ("PUT", "/api/configuracion/sistema", "Guardar configuracion"),
            ("GET", "/api/configuracion/slides", "Listar slides"),
            ("POST", "/api/configuracion/slides", "Crear slide"),
            ("PUT", "/api/configuracion/slides/{id}", "Editar slide"),
            ("DELETE", "/api/configuracion/slides/{id}", "Eliminar slide"),
            ("POST", "/api/configuracion/slides/upload", "Subir imagen"),
        ],
        "tables": ["RL_CONFIG_SISTEMA", "RL_LOGIN_SLIDES", "RL_AUDITORIA"],
        "flow": [
            "Abrir configuracion del sistema.",
            "Actualizar identidad visual, timeout y parametros generales.",
            "Administrar slides de login e imagenes.",
            "Guardar y validar mensaje de exito o error.",
        ],
        "rules": [
            "Colores deben respetar formato hexadecimal.",
            "Timeout y maximo de intentos deben estar en rangos permitidos.",
            "Toda modificacion critica exige permisos de modulo 3.",
        ],
        "audits": ["Cambio de configuracion", "Creacion/edicion/eliminacion de slide", "Carga de imagen"],
    },
    {
        "id": 4,
        "slug": "Monitoreo_Listas",
        "name": "Monitoreo de Listas",
        "route": "/monitoreo-listas",
        "capture": "03_monitoreo_listas.png",
        "frontend": ["features/admin/monitoreo-listas/monitoreo-listas.component.ts", "core/services/listas.service.ts"],
        "endpoints": [
            ("GET", "/api/listas/juridicas", "Coincidencias juridicas"),
            ("GET", "/api/listas/naturales", "Coincidencias naturales"),
            ("GET", "/api/listas/empleados", "Coincidencias empleados"),
            ("POST", "/api/listas/positivos", "Registrar motivo positivo"),
            ("GET", "/api/listas/positivos/{noDocumento}/seguimientos", "Seguimientos"),
            ("POST", "/api/listas/positivos/{noDocumento}/seguimientos", "Crear seguimiento/evidencia"),
            ("PUT", "/api/listas/seguimientos/{detalleId}", "Editar seguimiento"),
            ("DELETE", "/api/listas/seguimientos/{detalleId}", "Eliminar seguimiento"),
            ("GET", "/api/listas/evidencias/{id}", "Descargar evidencia"),
            ("DELETE", "/api/listas/evidencias/{id}", "Eliminar evidencia"),
            ("POST", "/api/listas/positivos/{noDocumento}/reporte-impreso", "Auditar reporte"),
        ],
        "tables": ["DNP_IHSS.V_REPORTE_COINCIDENCIA", "RL_LISTA_POSITIVOS", "RL_DETALLE_LISTA", "RL_DETALLE_EVIDENCIA", "RL_AUDITORIA"],
        "flow": [
            "Seleccionar categoria juridicas, naturales o empleados.",
            "Filtrar por texto, estado o fechas.",
            "Revisar detalle de coincidencia.",
            "Registrar motivo de inclusion.",
            "Agregar seguimientos y evidencias.",
            "Exportar o generar reporte auditado.",
        ],
        "rules": [
            "Seguimiento requiere motivo previo.",
            "Comentario de seguimiento obligatorio.",
            "Eliminaciones requieren motivo.",
            "Reportes/exportaciones se cancelan si falla auditoria.",
        ],
        "audits": ["Registro positivo", "Seguimientos", "Evidencias", "Descargas", "Eliminaciones logicas", "Reportes y exportaciones"],
    },
    {
        "id": 5,
        "slug": "Auditoria",
        "name": "Auditoria y Bitacora",
        "route": "/bitacora",
        "capture": "04_bitacora.png",
        "frontend": ["features/admin/bitacora/bitacora.component.ts", "core/services/auditoria.service.ts"],
        "endpoints": [("GET", "/api/auditoria", "Consulta bitacora"), ("POST", "/api/auditoria/exportacion", "Registra exportacion")],
        "tables": ["RL_AUDITORIA", "RL_USUARIOS"],
        "flow": [
            "Abrir bitacora.",
            "Filtrar por busqueda, accion, modulo y fechas.",
            "Revisar detalle de valores anteriores y nuevos.",
            "Aplicar filtro rapido de documentos eliminados.",
        ],
        "rules": [
            "Fecha desde no puede ser mayor que fecha hasta.",
            "Endpoint critico sin auditoria no se considera cerrado.",
            "Debe conservar usuario, IP, tabla, registro y modulo.",
        ],
        "audits": ["Consulta de eventos registrados", "Exportaciones y reportes desde modulos criticos"],
    },
    {
        "id": 6,
        "slug": "Tipo_Listas",
        "name": "Tipos de Listas de Cautela",
        "route": "/tipo-listas",
        "capture": "05_tipo_listas.png",
        "frontend": ["features/admin/tipo-listas/tipo-listas.component.ts/html", "core/services/listas.service.ts"],
        "endpoints": [
            ("GET", "/api/listas/tipos-listas-cautela", "Listar tipos"),
            ("POST", "/api/listas/tipos-listas-cautela", "Crear tipo"),
            ("PUT", "/api/listas/tipos-listas-cautela/{id}", "Editar tipo"),
            ("DELETE", "/api/listas/tipos-listas-cautela/{id}", "Eliminar tipo"),
        ],
        "tables": ["DNP_IHSS.TIPO_LISTA_CAUTELA", "RL_AUDITORIA"],
        "flow": ["Consultar tipos existentes.", "Crear o editar descripcion/cantidad de columnas.", "Eliminar solo con confirmacion."],
        "rules": ["Descripcion obligatoria.", "Cantidad de columnas mayor o igual a 1.", "No eliminar si esta referenciado."],
        "audits": ["Creacion", "Edicion", "Eliminacion de tipo de lista"],
    },
    {
        "id": 7,
        "slug": "Cargar_Listas",
        "name": "Carga de Listas de Cautela",
        "route": "/cargar-listas",
        "capture": "06_cargar_listas.png",
        "frontend": ["features/admin/cargar-listas/cargar-listas.component.ts/html", "core/services/listas.service.ts"],
        "endpoints": [
            ("GET", "/api/listas/resumen", "Resumen de listas cargadas"),
            ("POST", "/api/listas/cautela/upload", "Carga de archivo"),
            ("GET", "/api/listas/{id}/exportar", "Exportar lista"),
            ("GET", "/api/listas/tipos-listas-cautela", "Tipos disponibles"),
        ],
        "tables": ["DNP_IHSS.LISTA_CAUTELA", "DNP_IHSS.TIPO_LISTA_CAUTELA", "RL_AUDITORIA"],
        "flow": ["Seleccionar tipo de lista.", "Seleccionar archivo CSV/XLSX/XML.", "Validar y cargar.", "Revisar resumen.", "Exportar detalle si aplica."],
        "rules": ["Extension permitida CSV, XLSX o XML.", "Validacion de columnas por tipo.", "Botones bloqueados durante carga/exportacion."],
        "audits": ["Carga de lista", "Exportacion de lista"],
    },
    {
        "id": 8,
        "slug": "Coincidencias_Patrono",
        "name": "Coincidencias Patrono",
        "route": "/coincidencias-patrono",
        "capture": "07_coincidencias_patrono.png",
        "frontend": ["features/admin/coincidencias-patrono/coincidencias-patrono.component.ts/html", "core/services/listas.service.ts"],
        "endpoints": [
            ("GET", "/api/listas/coincidencias-patrono/resumen", "Resumen por fecha"),
            ("GET", "/api/listas/coincidencias-patrono/detalle", "Detalle por fecha"),
            ("GET", "/api/listas/coincidencias-patrono/resumen-match", "Resumen dinamico"),
            ("PUT", "/api/listas/coincidencias-patrono/{id}/calificar", "Calificar coincidencia"),
            ("POST", "/api/auditoria/exportacion", "Auditar exportacion"),
        ],
        "tables": ["DNP_IHSS.V_REPORTE_COINCIDENCIA", "RL_CALIF_COINCIDENCIAS", "RL_AUDITORIA"],
        "flow": ["Consultar resumen por fecha.", "Abrir detalle.", "Filtrar y paginar.", "Calificar como positivo/falso positivo.", "Exportar con auditoria."],
        "rules": ["Solo tipos de patrono se califican en este modulo.", "Calificacion valida: Positivo o Falso Positivo.", "Exportacion cancelada si falla auditoria."],
        "audits": ["Calificacion de patrono", "Exportacion Excel"],
    },
    {
        "id": 9,
        "slug": "Coincidencias_Empleado",
        "name": "Coincidencias Empleado",
        "route": "/coincidencias-empleado",
        "capture": "08_coincidencias_empleado.png",
        "frontend": ["features/admin/coincidencias-empleado/coincidencias-empleado.component.ts/html", "core/services/listas.service.ts"],
        "endpoints": [
            ("GET", "/api/listas/coincidencias-empleado/resumen", "Resumen por fecha"),
            ("GET", "/api/listas/coincidencias-empleado/detalle", "Detalle por fecha"),
            ("GET", "/api/listas/coincidencias-empleado/resumen-match", "Resumen dinamico"),
            ("PUT", "/api/listas/coincidencias-empleado/{id}/calificar", "Calificar coincidencia"),
            ("POST", "/api/auditoria/exportacion", "Auditar exportacion"),
        ],
        "tables": ["DNP_IHSS.V_REPORTE_COINCIDENCIA", "RL_CALIF_COINCIDENCIAS", "RL_AUDITORIA"],
        "flow": ["Consultar resumen por fecha.", "Abrir detalle.", "Filtrar y paginar.", "Calificar como positivo/falso positivo.", "Exportar con auditoria."],
        "rules": ["Solo tipos de empleado se califican en este modulo.", "Calificacion valida: Positivo o Falso Positivo.", "Exportacion cancelada si falla auditoria."],
        "audits": ["Calificacion de empleado", "Exportacion Excel"],
    },
]


def add_control_full(doc: Document, module_name: str, doc_type: str) -> None:
    base.add_table(doc, ["Campo", "Valor"], [
        ("Modulo", module_name),
        ("Tipo de documento", doc_type),
        ("Version", VERSION),
        ("Fecha", DATE),
        ("Estado", "Version full generada en Word; lista para aprobacion del usuario."),
        ("Aprobacion", "Pendiente de revision/firma del usuario responsable."),
        ("Fuente", "Repositorio local C:/RIESGO_LAVADO"),
    ], widths=[1.8, 4.7])


def add_capture(doc: Document, capture_file: str, caption: str) -> None:
    base.add_heading(doc, "Captura actualizada", 1)
    image = CAPTURES / capture_file
    if image.exists():
        p = doc.add_paragraph()
        p.alignment = 1
        run = p.add_run()
        run.add_picture(str(image), width=Inches(6.2))
        cap = doc.add_paragraph(caption)
        cap.alignment = 1
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(90, 90, 90)
    else:
        base.add_para(doc, f"Pendiente: no se encontro la captura {capture_file}.")


def save_doc(doc: Document, filename: str) -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(OUT / filename)
    Document(OUT / filename)


def client_doc(m: dict) -> None:
    title = f"{m['name']} - Version Cliente"
    doc = Document()
    base.style_doc(doc, title)
    base.add_title(doc, title, "Guia funcional completa para uso, flujo, reglas, auditorias y capturas del modulo.")
    add_control_full(doc, m["name"], "Cliente funcional")
    base.add_heading(doc, "Objetivo funcional", 1)
    base.add_para(doc, f"Documentar el uso operativo del modulo {m['name']} desde la perspectiva del usuario final autorizado.")
    base.add_heading(doc, "Flujo funcional", 1)
    base.add_numbers(doc, m["flow"])
    base.add_heading(doc, "Ruta del frontend", 1)
    base.add_table(doc, ["Ruta", "Modulo", "Proteccion"], [(m["route"], f"Modulo {m['id']}", f"authGuard + moduloGuard({m['id']})")], widths=[2.0, 1.5, 3.0])
    base.add_heading(doc, "Endpoints usados por la pantalla", 1)
    base.add_table(doc, ["Metodo", "Endpoint", "Uso"], m["endpoints"], widths=[0.8, 2.8, 2.9])
    base.add_heading(doc, "Tablas involucradas", 1)
    base.add_bullets(doc, m["tables"])
    base.add_heading(doc, "Reglas de negocio", 1)
    base.add_bullets(doc, m["rules"])
    base.add_heading(doc, "Auditorias que genera", 1)
    base.add_bullets(doc, m["audits"])
    add_capture(doc, m["capture"], f"Captura actualizada del modulo {m['name']} tomada desde build local del frontend.")
    base.add_heading(doc, "Pendientes conocidos", 1)
    base.add_bullets(doc, [
        "Validar capturas contra ambiente con datos reales si se requiere evidencia productiva.",
        "Registrar aprobacion formal del usuario responsable en el control documental.",
    ])
    save_doc(doc, f"{m['id']:02d}_{m['slug']}_Version_Cliente.docx")


def developer_doc(m: dict) -> None:
    title = f"{m['name']} - Version Desarrollador"
    doc = Document()
    base.style_doc(doc, title)
    base.add_title(doc, title, "Documento tecnico completo para mantenimiento, control de cambios y auditoria del modulo.")
    add_control_full(doc, m["name"], "Tecnico desarrollador")
    base.add_heading(doc, "Componentes frontend", 1)
    base.add_bullets(doc, m["frontend"])
    base.add_heading(doc, "Ruta y seguridad frontend", 1)
    base.add_table(doc, ["Ruta", "Guard", "Modulo"], [(m["route"], f"moduloGuard({m['id']})", f"MOD_ID {m['id']}")], widths=[2.0, 2.0, 2.5])
    base.add_heading(doc, "Endpoints backend", 1)
    endpoint_rows = [(a, b, c, f"ModuloAuthorize({m['id']}) cuando aplica") for a, b, c in m["endpoints"]]
    base.add_table(doc, ["Metodo", "Endpoint", "Uso", "Seguridad"], endpoint_rows, widths=[0.7, 2.5, 2.2, 1.1])
    base.add_heading(doc, "Tablas involucradas", 1)
    base.add_bullets(doc, m["tables"])
    base.add_heading(doc, "Reglas tecnicas y validaciones", 1)
    base.add_bullets(doc, m["rules"])
    base.add_heading(doc, "Auditorias que genera", 1)
    base.add_bullets(doc, m["audits"])
    base.add_heading(doc, "Criterios de cierre tecnico", 1)
    base.add_bullets(doc, [
        "Endpoint critico protegido por autorizacion de modulo.",
        "Validacion visible en frontend y validacion critica en backend.",
        "Errores de backend presentados al usuario cuando aplique.",
        "Auditoria obligatoria en acciones criticas.",
        "Compilacion frontend sin errores.",
    ])
    add_capture(doc, m["capture"], f"Captura actualizada usada como referencia visual del modulo {m['name']}.")
    base.add_heading(doc, "Pendientes conocidos", 1)
    base.add_bullets(doc, [
        "Aprobacion final del documento por el usuario responsable.",
        "Actualizar este documento si cambian rutas, endpoints, tablas o reglas.",
    ])
    save_doc(doc, f"{m['id']:02d}_{m['slug']}_Version_Desarrollador.docx")


def evidence_doc() -> None:
    title = "Evidencias y Seguimientos - Version Full"
    doc = Document()
    base.style_doc(doc, title)
    base.add_title(doc, title, "Documento completo de flujo funcional, seguridad documental, auditoria y almacenamiento.")
    add_control_full(doc, "Evidencias y Seguimientos", "Transversal funcional/tecnico")
    base.add_heading(doc, "Flujo funcional", 1)
    base.add_numbers(doc, [
        "Registrar motivo positivo.",
        "Crear seguimiento con comentario obligatorio.",
        "Adjuntar evidencias opcionales.",
        "Validar extension, MIME, tamano y firma real.",
        "Guardar nombre original y nombre fisico GUID.",
        "Descargar mediante endpoint protegido.",
        "Eliminar logicamente con motivo obligatorio y conservar archivo fisico.",
    ])
    base.add_heading(doc, "Endpoints", 1)
    base.add_table(doc, ["Metodo", "Endpoint", "Uso"], [
        ("GET", "/api/listas/evidencias/politica", "Politica de evidencias"),
        ("POST", "/api/listas/positivos/{doc}/seguimientos", "Crear seguimiento/evidencias"),
        ("PUT", "/api/listas/seguimientos/{id}", "Editar seguimiento/evidencias"),
        ("GET", "/api/listas/evidencias/{id}", "Descargar o visualizar"),
        ("DELETE", "/api/listas/evidencias/{id}", "Eliminar logicamente evidencia"),
        ("DELETE", "/api/listas/seguimientos/{id}", "Eliminar logicamente seguimiento"),
    ], widths=[0.8, 3.0, 2.7])
    base.add_heading(doc, "Tablas involucradas", 1)
    base.add_bullets(doc, ["RL_LISTA_POSITIVOS", "RL_DETALLE_LISTA", "RL_DETALLE_EVIDENCIA", "RL_AUDITORIA"])
    base.add_heading(doc, "Reglas de seguridad documental", 1)
    base.add_bullets(doc, [
        "Extensiones permitidas configurables.",
        "Tamano maximo configurable.",
        "MIME y firma real verificables.",
        "Ruta de almacenamiento final configurable.",
        "Sin rutas publicas directas.",
        "Eliminacion logica conserva archivo fisico.",
    ])
    base.add_heading(doc, "Auditorias que genera", 1)
    base.add_bullets(doc, [
        "INSERT en RL_DETALLE_LISTA al crear seguimiento.",
        "INSERT en RL_DETALLE_EVIDENCIA al adjuntar evidencia.",
        "VER en RL_DETALLE_EVIDENCIA al visualizar o descargar evidencia.",
        "DELETE logico en RL_DETALLE_EVIDENCIA al eliminar evidencia con motivo.",
        "DELETE logico en RL_DETALLE_LISTA al eliminar seguimiento con motivo.",
    ])
    add_capture(doc, "03_monitoreo_listas.png", "Pantalla base desde donde se gestionan seguimientos y evidencias.")
    base.add_heading(doc, "Pendientes conocidos", 1)
    base.add_bullets(doc, [
        "Validar ubicacion final de almacenamiento en ambiente productivo.",
        "Revisar politica de firma real si se agregan nuevos tipos de archivo.",
    ])
    save_doc(doc, "10_Evidencias_Seguimientos_Version_Full.docx")


def db_doc() -> None:
    title = "Base de Datos - Version Full"
    doc = Document()
    base.style_doc(doc, title)
    base.add_title(doc, title, "Guia completa de ejecucion segura, scripts aprobados, orden, respaldos y modulos.")
    add_control_full(doc, "Base de Datos", "Tecnico DBA/desarrollador")
    base.add_heading(doc, "Regla principal", 1)
    base.add_para(doc, "No mezclar scripts experimentales con scripts aprobados. Toda ejecucion debe generar evidencia de salida y respaldo previo si afecta tablas existentes.")
    base.add_heading(doc, "Orden de ejecucion", 1)
    base.add_numbers(doc, [
        "Respaldar esquema o tablas afectadas.",
        "Ejecutar scripts maestros segun instalacion o actualizacion.",
        "Verificar manifiesto de scripts aprobados.",
        "Ejecutar validacion de IDs de modulos.",
        "Conservar logs de SQLPlus/spool.",
    ])
    base.add_heading(doc, "Endpoints backend", 1)
    base.add_para(doc, "No aplica como modulo HTTP. La base de datos soporta los endpoints del backend mediante tablas, vistas, secuencias, indices y permisos.")
    base.add_heading(doc, "Reglas de negocio y ejecucion", 1)
    base.add_bullets(doc, [
        "No ejecutar scripts experimentales en ambientes aprobados.",
        "No avanzar al siguiente modulo si falla 17_validate_module_ids.sql.",
        "Respaldar tablas afectadas antes de actualizaciones.",
        "Mantener IDs de RL_MODULOS alineados con Angular y ModuloAuthorize.",
        "Conservar evidencia de ejecucion y errores SQLPlus.",
    ])
    base.add_heading(doc, "Scripts aprobados", 1)
    base.add_table(doc, ["Script", "Uso"], [
        ("00_EJECUCION_PRIMERA_VEZ.sql", "Instalacion inicial."),
        ("00_EJECUCION_ACTUALIZACIONES_SEGURAS.sql", "Actualizacion segura."),
        ("00_MANIFIESTO_SCRIPTS_APROBADOS.md", "Control de alcance."),
        ("17_validate_module_ids.sql", "Validacion final obligatoria."),
    ], widths=[3.2, 3.3])
    base.add_heading(doc, "Tablas criticas", 1)
    base.add_bullets(doc, ["RL_USUARIOS", "RL_MODULOS", "RL_USUARIO_MODULOS", "RL_AUDITORIA", "RL_LISTA_POSITIVOS", "RL_DETALLE_LISTA", "RL_DETALLE_EVIDENCIA", "RL_CALIF_COINCIDENCIAS"])
    base.add_heading(doc, "Modulos registrados", 1)
    base.add_table(doc, ["ID", "Ruta", "Modulo"], [(m["id"], m["route"], m["name"]) for m in MODULES], widths=[0.7, 2.4, 3.4])
    base.add_heading(doc, "Pendientes conocidos", 1)
    base.add_bullets(doc, [
        "Crear scripts separados para Matrices de Riesgo cuando el modulo sea aprobado.",
        "Actualizar manifiesto al incorporar nuevos scripts aprobados.",
    ])
    save_doc(doc, "11_Base_Datos_Version_Full.docx")


def matrix_doc() -> None:
    title = "Matriz Full de Documentacion Modular"
    doc = Document()
    base.style_doc(doc, title)
    base.add_title(doc, title, "Control completo de documentos cliente, desarrollador, capturas y aprobacion.")
    add_control_full(doc, "Documentacion", "Matriz de control full")
    rows = []
    for m in MODULES:
        rows.append((m["name"], "Cliente", f"{m['id']:02d}_{m['slug']}_Version_Cliente.docx", "Generado con captura"))
        rows.append((m["name"], "Desarrollador", f"{m['id']:02d}_{m['slug']}_Version_Desarrollador.docx", "Generado con captura"))
    rows.extend([
        ("Evidencias y Seguimientos", "Transversal", "10_Evidencias_Seguimientos_Version_Full.docx", "Generado"),
        ("Base de Datos", "Transversal", "11_Base_Datos_Version_Full.docx", "Generado"),
    ])
    base.add_table(doc, ["Modulo", "Tipo", "Documento", "Estado"], rows, widths=[1.7, 1.1, 2.6, 1.1])
    base.add_heading(doc, "Reglas de aprobacion", 1)
    base.add_bullets(doc, [
        "Revisar cada Word antes de marcarlo como aprobado.",
        "No subir documentacion al repositorio remoto sin aprobacion.",
        "Si cambia un endpoint, ruta o tabla, actualizar cliente y desarrollador.",
        "Las capturas incluidas son actuales del build local con datos mock para documentacion visual.",
    ])
    save_doc(doc, "00_Matriz_Full_Documentacion_Modular.docx")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for m in MODULES:
        client_doc(m)
        developer_doc(m)
    evidence_doc()
    db_doc()
    matrix_doc()
    for p in sorted(OUT.glob("*.docx")):
        Document(p)
        print(p)


if __name__ == "__main__":
    main()
