from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_ROOT = ROOT / "docs"
MATRICES_DIR = next(DOCS_ROOT.rglob("3. Módulo Matrices de Riesgos"))
EVIDENCE_DIR = MATRICES_DIR / "Fase 3 - Modelo de datos y arquitectura Oracle" / "Evidencia_DBA"
OUT_FILE = EVIDENCE_DIR / "Evidencia_DBA_Correccion_Encoding_Fase_3_Matrices_Riesgos_SGRLA_IHSS.docx"

BLUE = RGBColor(46, 116, 181)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E2F0D9"
LIGHT_GOLD = "FFF2CC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=10.5, color=INK, bold=False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[int], fill: str = LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], fill)
        run = table.rows[0].cells[idx].paragraphs[0].add_run(header)
        set_run_font(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    return table


def build_document() -> None:
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size in [(1, 16), (2, 13)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE

    p = doc.add_paragraph()
    run = p.add_run("EVIDENCIA DBA DE CORRECCIÓN DE CODIFICACIÓN")
    set_run_font(run, size=10, color=BLUE, bold=True)
    p = doc.add_paragraph()
    run = p.add_run("Fase 3 - Matrices de Riesgos")
    set_run_font(run, size=22, bold=True)
    p = doc.add_paragraph()
    run = p.add_run("Sistema de Gestión de Riesgos LA/FT IHSS")
    set_run_font(run, size=13, color=MUTED)

    add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ("Proyecto", "RIESGO_LAVADO - IHSS"),
            ("Responsable", "Javier Mejía"),
            ("Fecha", date.today().strftime("%d/%m/%Y")),
            ("Alcance", "Corrección de tildes y caracteres especiales en comentarios Oracle y textos descriptivos RL_MR_*."),
            ("Estado", "Corrección ejecutada y validada."),
        ],
        [2200, 7160],
    )

    doc.add_heading("1. Motivo", level=1)
    doc.add_paragraph(
        "Durante la revisión de base de datos se identificaron textos con codificación incorrecta en comentarios y valores descriptivos del módulo Matrices de Riesgos. "
        "La corrección se realizó sin renombrar tablas, columnas, llaves, índices ni restricciones."
    )

    doc.add_heading("2. Corrección Ejecutada", level=1)
    add_table(
        doc,
        ["Elemento", "Resultado"],
        [
            ("Comentarios de tablas RL_MR_*", "Corregidos mediante COMMENT ON con codificación Windows-1252 controlada."),
            ("Comentarios de columnas RL_MR_*", "Corregidos mediante COMMENT ON con codificación Windows-1252 controlada."),
            ("RL_MODULOS", "Descripción del módulo corregida: Módulo, evaluación, cálculo y reportería."),
            ("RL_MR_MODELOS", "Nombre y descripción corregidos: Metodología y metodológicamente."),
            ("RL_MR_ESCALAS", "Niveles y descripciones corregidas: Crítico, Débil, acción, mitigación, exposición, entre otros."),
            ("Estructura técnica", "Sin cambios físicos en nombres de tablas, columnas, secuencias, índices o restricciones."),
        ],
        [3300, 6060],
        fill=LIGHT_GOLD,
    )

    doc.add_heading("3. Validación Final", level=1)
    add_table(
        doc,
        ["Validación", "Resultado"],
        [
            ("Comentarios de tabla dañados", "0"),
            ("Comentarios de columna dañados", "0"),
            ("Textos dañados en RL_MODULOS", "0"),
            ("Textos dañados en columnas RL_MR_*", "0"),
            ("Muestra MOD_DESCRIPCION", "M\\00F3dulo, evaluaci\\00F3n, c\\00E1lculo, reporter\\00EDa."),
            ("Muestra MRM_NOMBRE", "Metodolog\\00EDa base LA/FT IHSS."),
            ("Muestra comentarios", "m\\00F3dulo, espec\\00EDfica, metodol\\00F3gicas."),
        ],
        [3300, 6060],
        fill=LIGHT_GREEN,
    )

    doc.add_heading("4. Evidencia Técnica", level=1)
    add_table(
        doc,
        ["Archivo", "Descripción"],
        [
            ("04_F3_fix_encoding_textos_oracle.sql", "Script correctivo aprobado ubicado en Scripts Aprobables."),
            ("05_fix_encoding_textos_oracle_cp1252_*.log", "Log de ejecución exitosa del correctivo de codificación."),
            ("06_validacion_final_encoding_*.log", "Log de validación final sin caracteres dañados."),
        ],
        [3800, 5560],
    )

    props = doc.core_properties
    props.title = "Evidencia DBA - Corrección de Encoding - Fase 3"
    props.subject = "Sistema de Gestión de Riesgos LA/FT IHSS"
    props.author = "Javier Mejía"
    props.comments = "Evidencia de corrección de codificación en base de datos para Fase 3."
    doc.save(OUT_FILE)


if __name__ == "__main__":
    build_document()
    print(OUT_FILE)
