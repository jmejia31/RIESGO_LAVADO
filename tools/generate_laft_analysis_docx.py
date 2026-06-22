from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = r"C:\RIESGO_LAVADO\docs\Analisis_Propuesta_Catalogos_LAFT.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 30, 45)
MUTED = RGBColor(90, 102, 120)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
SOFT_GREEN = "EAF6EF"
SOFT_GOLD = "FFF5D6"
SOFT_RED = "FDECEC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    table.autofit = False


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_paragraph(doc, text="", bold_prefix=None, after=6, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, color=color, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        size, color = 16, BLUE
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        size, color = 13, BLUE
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        size, color = 12, DARK_BLUE
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=True)
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, bottom=130, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_matrix(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_width(hdr[i], widths[i])
        set_cell_shading(hdr[i], header_fill)
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ["List Bullet", "List Bullet 2"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = header.add_run("SGRLA-IHSS | Propuesta de mejora LAFT")
    set_run_font(hr, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Documento de analisis para decision interna")
    set_run_font(fr, size=9, color=MUTED)
    return doc


def build_doc():
    doc = setup_document()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Análisis y Propuesta")
    set_run_font(r, size=23, color=INK, bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(16)
    r2 = p2.add_run("Catálogos LAFT para estados, acciones de seguimiento y tipos de evidencia")
    set_run_font(r2, size=14, color=MUTED)

    meta = [
        ("Sistema:", "SGRLA-IHSS"),
        ("Tema:", "Separación de catálogos LAFT"),
        ("Objetivo:", "Explicar qué se gana, qué se pierde y qué áreas se verían afectadas"),
        ("Alcance:", "Monitoreo de Listas, Seguimientos, Evidencias, Auditoría y Reportería"),
    ]
    add_matrix(doc, ["Campo", "Detalle"], meta, [1.3, 5.2], header_fill=LIGHT_GRAY)

    add_callout(
        doc,
        "Resumen ejecutivo",
        "El sistema actualmente funciona y permite registrar seguimientos y evidencias. "
        "La propuesta no busca cambiar el acceso por Active Directory ni rehacer el sistema; "
        "busca ordenar mejor la información LAFT usando catálogos controlados para estados, "
        "acciones y tipos de evidencia. Esto daría más trazabilidad, mejores reportes y mayor control, "
        "pero debe manejarse como una mejora planificada porque implica cambios de base de datos, backend, frontend y pruebas.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "1. Situación actual", 1)
    add_paragraph(
        doc,
        "Hoy el módulo de Monitoreo de Listas permite registrar comentarios de seguimiento, adjuntar evidencias, "
        "consultar historial y dejar bitácora de acciones importantes. Es decir, el proceso ya opera.",
    )
    add_paragraph(
        doc,
        "La limitante es que parte de esa información queda como texto libre. Eso ayuda porque es flexible, "
        "pero complica medir y reportar cuántos casos están pendientes, en análisis, confirmados, cerrados o qué tipo de evidencia se adjuntó.",
    )

    add_heading(doc, "2. Qué se quiere lograr", 1)
    add_paragraph(
        doc,
        "La mejora propone que, al registrar un seguimiento, el usuario no solo escriba una nota, sino que también seleccione valores controlados.",
    )
    for item in [
        "Estado del caso: pendiente, en análisis, confirmado, falso positivo, cerrado.",
        "Acción realizada: revisión, validación documental, escalamiento, cierre, solicitud de soporte.",
        "Tipo de evidencia: oficio, informe, captura, acta, constancia o documento.",
    ]:
        add_bullet(doc, item)
    add_paragraph(
        doc,
        "La idea es mantener el comentario libre, pero acompañarlo de datos ordenados que permitan seguimiento, auditoría y reportes más claros.",
    )

    add_heading(doc, "3. Qué ganamos y qué perdemos", 1)
    add_matrix(
        doc,
        ["Aspecto", "Qué ganamos", "Qué perdemos o riesgo"],
        [
            (
                "Orden del proceso",
                "Los usuarios seleccionan valores estándar y se reduce la variación en textos.",
                "Se agregan campos nuevos que los usuarios deberán aprender a usar.",
            ),
            (
                "Trazabilidad",
                "Se entiende mejor qué acción se hizo, en qué estado quedó el caso y qué evidencia respalda la gestión.",
                "Requiere definir reglas claras para estados y acciones.",
            ),
            (
                "Reportería",
                "Permite reportes por estado, acción, tipo de evidencia, usuario y tiempos de atención.",
                "Los reportes existentes podrían requerir ajustes si se incorporan los nuevos datos.",
            ),
            (
                "Auditoría",
                "La bitácora tendría más contexto sobre cambios sensibles dentro del seguimiento.",
                "Se debe cuidar que todo cambio nuevo quede auditado correctamente.",
            ),
            (
                "Base de datos",
                "La información queda más estructurada y consultable.",
                "Implica nuevas tablas y posiblemente nuevas columnas.",
            ),
            (
                "Operación",
                "Facilita supervisión y cumplimiento.",
                "No conviene aplicarlo directo en producción sin pruebas previas.",
            ),
        ],
        [1.35, 2.55, 2.6],
        header_fill=LIGHT_GRAY,
    )

    add_heading(doc, "4. Información y bases de datos afectadas", 1)
    add_paragraph(doc, "La información afectada sería únicamente la relacionada con el seguimiento LAFT. No se tocarían claves ni autenticación de Active Directory.")
    add_matrix(
        doc,
        ["Elemento", "Uso actual", "Cambio propuesto"],
        [
            ("RL_LISTA_POSITIVOS", "Registro principal del caso o coincidencia monitoreada.", "Se mantiene como registro principal."),
            ("RL_DETALLE_LISTA", "Comentarios o seguimientos del caso.", "Podría guardar estado LAFT y acción seleccionada."),
            ("RL_DETALLE_EVIDENCIA", "Archivos adjuntos de evidencia.", "Podría guardar el tipo de evidencia."),
            ("RL_AUDITORIA", "Bitácora de acciones del sistema.", "Registrar cambios de estado, acciones y evidencias sensibles."),
            ("RL_CAT_ESTADOS_LAFT", "No existe actualmente.", "Nueva tabla para estados del caso."),
            ("RL_CAT_ACCIONES_SEGUIMIENTO", "No existe actualmente.", "Nueva tabla para acciones de seguimiento."),
            ("RL_CAT_TIPOS_EVIDENCIA", "No existe actualmente.", "Nueva tabla para tipos de evidencia."),
        ],
        [1.55, 2.35, 2.6],
        header_fill=LIGHT_GRAY,
    )

    add_heading(doc, "5. De dónde vendrá y a dónde irá la información", 1)
    add_paragraph(doc, "El acceso seguiría funcionando igual: el usuario entra con su cuenta institucional y Active Directory valida su identidad.")
    add_matrix(
        doc,
        ["Paso", "Origen", "Destino"],
        [
            ("Inicio de sesión", "Cuenta institucional del usuario", "Validación por Active Directory"),
            ("Permisos", "Usuario autenticado y permisos internos", "Módulos permitidos en el sistema"),
            ("Catálogos LAFT", "Tablas nuevas de catálogos", "Listas desplegables en Monitoreo de Listas"),
            ("Seguimiento", "Formulario de seguimiento", "RL_DETALLE_LISTA"),
            ("Evidencia", "Archivo cargado por el usuario", "RL_DETALLE_EVIDENCIA y repositorio de archivos"),
            ("Auditoría", "Acción realizada en pantalla o backend", "RL_AUDITORIA"),
        ],
        [1.4, 2.45, 2.65],
        header_fill=LIGHT_GRAY,
    )

    add_heading(doc, "6. Código que se vería afectado", 1)
    add_paragraph(doc, "El cambio debe separarse bien para no mezclar catálogos generales del sistema con catálogos propios del proceso LAFT.")
    add_matrix(
        doc,
        ["Capa", "Archivos o componentes", "Qué cambiaría"],
        [
            ("Backend", "CatalogosLaftController.cs, CatalogosLaftService.cs, CatalogosLaftRepository.cs", "Nuevos endpoints para consultar y mantener catálogos LAFT."),
            ("Backend existente", "ListasController.cs, ListasRepository.cs", "Guardar estado, acción y tipo de evidencia en los seguimientos."),
            ("Frontend", "monitoreo-listas.component.ts/html, listas.service.ts", "Agregar campos en la ventana de seguimiento e historial."),
            ("Frontend opcional", "catalogos-laft.component.ts/html", "Pantalla administrativa para mantener estados, acciones y tipos de evidencia."),
            ("Base de datos", "Scripts SQL", "Crear tablas nuevas y, si se aprueba, columnas relacionadas."),
            ("Auditoría", "AuditoriaRepository.cs y puntos de registro", "Registrar cambios sensibles asociados al seguimiento LAFT."),
        ],
        [1.35, 2.4, 2.75],
        header_fill=LIGHT_GRAY,
    )

    add_heading(doc, "7. Impacto en los tres ambientes", 1)
    add_matrix(
        doc,
        ["Ambiente", "Qué se haría", "Cuidado principal"],
        [
            ("Desarrollo", "Crear tablas, endpoints y cambios de pantalla. Validar que el flujo guarde correctamente.", "No romper el flujo actual de seguimiento y evidencia."),
            ("Pruebas", "Probar con usuarios y casos reales o controlados. Validar permisos, auditoría y reportes.", "Confirmar si los campos serán obligatorios u opcionales."),
            ("Producción", "Aplicar scripts y despliegues ya validados. Ejecutar prueba rápida posterior al cambio.", "Hacer respaldo, despliegue controlado y validación funcional."),
        ],
        [1.35, 2.75, 2.4],
        header_fill=LIGHT_GRAY,
    )

    add_heading(doc, "8. Riesgos y controles recomendados", 1)
    add_matrix(
        doc,
        ["Riesgo", "Control recomendado"],
        [
            ("Afectar seguimientos existentes", "Agregar campos como opcionales al inicio o usar un valor por defecto como 'No clasificado'."),
            ("Confusión de usuarios", "Capacitación breve y nombres claros en los catálogos."),
            ("Catálogos mal definidos", "Aprobar previamente estados, acciones y tipos de evidencia con el área responsable."),
            ("Impacto en reportes", "Revisar reportes después de validar el modelo de datos."),
            ("Cambios directos en producción", "Pasar por desarrollo y pruebas antes de producción."),
            ("Confundirlo con Active Directory", "Documentar que AD no se toca; solo se mejora el registro interno del proceso LAFT."),
        ],
        [2.4, 4.1],
        header_fill=LIGHT_GRAY,
    )

    add_heading(doc, "9. Recomendación", 1)
    add_callout(
        doc,
        "Recomendación para decisión",
        "Cerrar el punto actual de catálogos base como operativo con mejora futura identificada. "
        "La separación de catálogos LAFT sí aporta valor, pero debe manejarse como una fase planificada, "
        "porque implica cambios de base de datos, backend, frontend, auditoría, pruebas y posible capacitación.",
        fill=SOFT_GREEN,
    )
    add_paragraph(
        doc,
        "En palabras simples: hoy el sistema permite trabajar; la mejora haría que el trabajo quede mejor clasificado, más medible y más fácil de auditar.",
    )

    add_heading(doc, "10. Decisiones pendientes antes de implementar", 1)
    for item in [
        "Definir cuáles serán los estados oficiales del caso LAFT.",
        "Definir cuáles serán las acciones de seguimiento permitidas.",
        "Definir cuáles serán los tipos de evidencia.",
        "Decidir si los nuevos campos serán obligatorios u opcionales al inicio.",
        "Decidir si se migrará información histórica o si la mejora aplicará solo hacia adelante.",
        "Definir quién administrará los catálogos LAFT.",
    ]:
        add_bullet(doc, item)

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
