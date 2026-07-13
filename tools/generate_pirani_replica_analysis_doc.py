from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\RIESGO_LAVADO")
OUT_DIR = ROOT / "docs" / "3. Módulo Matrices de Riesgos" / "Revision Pirani Replica - Analisis Funcional y Diseno"
OUT_DOCX = OUT_DIR / "Analisis_Replica_Pirani_Matrices_Riesgos_SGRLA_IHSS.docx"


COLORS = {
    "ihss_blue": "1F3A8A",
    "heading": "2E74B5",
    "heading_dark": "1F4D78",
    "ink": "0B2545",
    "muted": "5B677A",
    "table_header": "F2F4F7",
    "callout": "EAF2F8",
    "warning": "FFF4E5",
    "success": "EAF7EF",
    "border": "D9E2EC",
}


def hex_to_rgb(value):
    value = value.strip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=COLORS["border"], size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_border(cell)
            set_cell_margins(cell)


def style_run(run, bold=False, italic=False, color=None, size=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = hex_to_rgb(color)
    if size:
        run.font.size = Pt(size)


def apply_paragraph_format(paragraph, before=0, after=6, line_spacing=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    if level == 1:
        apply_paragraph_format(paragraph, before=16, after=8)
        run = paragraph.add_run(text)
        style_run(run, bold=True, color=COLORS["heading"], size=16)
    elif level == 2:
        apply_paragraph_format(paragraph, before=12, after=6)
        run = paragraph.add_run(text)
        style_run(run, bold=True, color=COLORS["heading"], size=13)
    else:
        apply_paragraph_format(paragraph, before=8, after=4)
        run = paragraph.add_run(text)
        style_run(run, bold=True, color=COLORS["heading_dark"], size=12)
    return paragraph


def add_body(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    apply_paragraph_format(p)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        style_run(r1, bold=True, color=COLORS["ink"], size=11)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2, color=COLORS["ink"], size=11)
    else:
        r = p.add_run(text)
        style_run(r, color=COLORS["ink"], size=11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    apply_paragraph_format(p, after=8, line_spacing=1.167)
    for run in p.runs:
        style_run(run, color=COLORS["ink"], size=11)
    if not p.runs:
        r = p.add_run(text)
        style_run(r, color=COLORS["ink"], size=11)
    else:
        p.runs[0].text = text
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    apply_paragraph_format(p, after=8, line_spacing=1.167)
    if not p.runs:
        r = p.add_run(text)
        style_run(r, color=COLORS["ink"], size=11)
    else:
        p.runs[0].text = text
        for run in p.runs:
            style_run(run, color=COLORS["ink"], size=11)
    return p


def add_callout(doc, title, text, fill="callout"):
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS[fill])
    p = cell.paragraphs[0]
    apply_paragraph_format(p, after=4)
    r = p.add_run(title)
    style_run(r, bold=True, color=COLORS["ink"], size=11)
    p2 = cell.add_paragraph()
    apply_paragraph_format(p2, after=2)
    r2 = p2.add_run(text)
    style_run(r2, color=COLORS["ink"], size=10.5)
    doc.add_paragraph()
    return table


def add_table(doc, headers, rows, widths=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    if widths is None:
        widths = [int(9360 / len(headers))] * len(headers)
    set_table_fixed_width(table, widths)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = ""
        set_cell_shading(hdr[idx], COLORS["table_header"])
        p = hdr[idx].paragraphs[0]
        apply_paragraph_format(p, after=2, line_spacing=1.0)
        r = p.add_run(header)
        style_run(r, bold=True, color=COLORS["ink"], size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            apply_paragraph_format(p, after=2, line_spacing=1.0)
            r = p.add_run(str(value))
            style_run(r, color=COLORS["ink"], size=font_size)
            if idx < len(widths):
                tc_pr = cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:type"), "dxa")
                tc_w.set(qn("w:w"), str(widths[idx]))
            set_cell_border(cells[idx])
            set_cell_margins(cells[idx])
    doc.add_paragraph()
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    style_run(run, color=COLORS["muted"], size=9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = hex_to_rgb(COLORS["ink"])

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("SGRLA-IHSS | Módulo Matrices de Riesgos")
    style_run(run, bold=True, color=COLORS["muted"], size=9)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    return doc


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_paragraph_format(title, after=8)
    r = title.add_run("INFORME DE ANÁLISIS FUNCIONAL Y DISEÑO")
    style_run(r, bold=True, color=COLORS["ihss_blue"], size=18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_paragraph_format(subtitle, after=8)
    r = subtitle.add_run("Réplica funcional de Pirani Risks para el Módulo Matrices de Riesgos")
    style_run(r, bold=True, color=COLORS["heading_dark"], size=14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_paragraph_format(meta, after=14)
    r = meta.add_run("Sistema de Gestión de Riesgos LA/FT - IHSS | Documento de análisis, sin implementación de código")
    style_run(r, color=COLORS["muted"], size=10)

    add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ["Cliente", "Instituto Hondureño de Seguridad Social (IHSS)"],
            ["Sistema", "SGRLA-IHSS"],
            ["Módulo evaluado", "Matrices de Riesgos"],
            ["Referencia funcional", "Pirani Risks - entorno sandbox observado el 10/07/2026"],
            ["Autor del análisis", "Codex - Analista funcional, diseñador UX y programador senior"],
            ["Estado", "Análisis de replanteamiento funcional. No incluye programación."],
        ],
        [2500, 6860],
        font_size=10,
    )

    add_callout(
        doc,
        "Conclusión ejecutiva",
        "Sí es técnicamente posible replicar en el SGRLA-IHSS la experiencia funcional observada en Pirani Risks para Matrices de Riesgos. "
        "La réplica debe entenderse como equivalencia funcional y visual adaptada a nuestros colores, roles, auditoría, base de datos y reglas institucionales, no como copia de código ni dependencia del producto externo.",
        "success",
    )

    add_heading(doc, "0. Control de cambios", 1)
    add_table(
        doc,
        ["Versión", "Fecha", "Responsable", "Descripción"],
        [["1.0", "10/07/2026", "Codex", "Documento inicial de análisis para replantear Matrices de Riesgos con referencia Pirani Risks."]],
        [1000, 1400, 1800, 5160],
    )

    add_heading(doc, "1. Propósito del documento", 1)
    add_body(
        doc,
        "Este documento analiza cómo replicar dentro del SGRLA-IHSS el comportamiento funcional, la experiencia de usuario y la relación entre interfaces observadas en Pirani Risks para el módulo Matrices de Riesgos."
    )
    add_body(
        doc,
        "El objetivo no es programar en esta etapa. El objetivo es dejar clara la viabilidad, el alcance, los componentes, las relaciones de datos, la adaptación visual y el camino de construcción para que el cliente y la jefatura puedan validar el nuevo rumbo antes de modificar el sistema."
    )

    add_heading(doc, "2. Decisión de viabilidad", 1)
    add_table(
        doc,
        ["Pregunta", "Respuesta técnica"],
        [
            ["¿Se puede replicar?", "Sí. La arquitectura actual del SGRLA-IHSS permite construir una réplica funcional por módulos, manteniendo autenticación, permisos, auditoría y colores institucionales."],
            ["¿Es un ajuste pequeño?", "No. Es un replanteamiento funcional completo del módulo Matrices de Riesgos. Los avances actuales sirven como base técnica, pero la experiencia final debe reorganizarse alrededor de riesgos, procesos, factores, controles, evaluaciones y parametrización."],
            ["¿Debe depender de Monitoreo de Listas?", "No. Matrices de Riesgos y Monitoreo de Listas serán módulos independientes dentro del mismo sistema. Cualquier alerta de Matrices será una alerta interna de riesgo, no una coincidencia ni evento del módulo de listas."],
            ["¿Se deben cambiar colores?", "No. Se replica la lógica de navegación, jerarquía visual, formularios y experiencia; los colores se adaptan al diseño actual del SGRLA-IHSS."],
        ],
        [2600, 6760],
    )

    add_heading(doc, "3. Alcance observado en Pirani Risks", 1)
    add_body(
        doc,
        "La revisión se realizó sobre el entorno sandbox de Pirani Risks. Se observaron pantallas con datos limitados, por lo que algunos listados, reportes e indicadores mostraron estados vacíos o información mínima. Aun así, la navegación, formularios, relaciones funcionales y reglas de experiencia sí fueron suficientes para definir una réplica."
    )
    add_table(
        doc,
        ["Interfaz observada", "Función principal", "Decisión para SGRLA-IHSS"],
        [
            ["Dashboard", "Vista ejecutiva con mapa de calor, totales y accesos a parametrización/reportes.", "Debe ser la primera pantalla de Matrices de Riesgos."],
            ["Procesos", "Registro de procesos organizacionales y asociaciones con riesgos.", "Crear submódulo interno de procesos del módulo Matrices."],
            ["Factores de riesgo", "Clientes, contrapartes, canales, productos y jurisdicciones.", "Crear catálogo operativo de factores, independiente de Monitoreo de Listas."],
            ["Riesgos", "Entidad central: impacto, frecuencia, riesgo inherente, asociaciones y gestión.", "Debe reemplazar el enfoque actual de captura aislada por gestión integral del riesgo."],
            ["Causas y consecuencias", "Registro de causas/consecuencias asociables a riesgos y controles.", "Agregar catálogo relacional."],
            ["Controles", "Diseño, ejecución, solidez y mitigación del riesgo.", "Agregar evaluación de control y vínculo con riesgo residual."],
            ["Alertas", "Alertas internas de gestión de riesgos.", "Implementar solo como alerta del ecosistema de Matrices, sin conexión con Monitoreo de Listas."],
            ["Planes de acción", "Gestión de tratamientos y actividades.", "Agregar flujo de planes con responsables y fechas."],
            ["Evaluaciones", "Planes de evaluación periódica de riesgos y controles.", "Agregar evaluación periódica y trazabilidad."],
            ["Reportes e indicadores", "Consulta ejecutiva y tableros dinámicos.", "Implementar como capa de lectura sobre datos calculados."],
            ["Parametrización", "Mapa de calor, controles, pesos, variables y advertencias de recálculo.", "Debe ser submódulo clave antes de calcular riesgos."],
        ],
        [2100, 3650, 3610],
        font_size=8.8,
    )

    add_heading(doc, "4. Principios de réplica", 1)
    principles = [
        "Equivalencia funcional: el usuario debe sentir que opera el mismo modelo de gestión, aunque el sistema sea propio del IHSS.",
        "Diseño institucional: se mantiene la paleta visual, sidebar, tipografía y componentes base del SGRLA-IHSS.",
        "Separación estricta: Matrices de Riesgos no consume, no hereda y no depende de Monitoreo de Listas.",
        "Parametrización antes de cálculo: escalas, pesos, niveles, solidez y reglas deben vivir en base de datos.",
        "Trazabilidad total: crear, editar, recalcular, cambiar estado, exportar y descargar debe quedar auditado.",
        "UX guiada: las pantallas críticas deben usar formularios por pasos, estados vacíos claros, validaciones visibles y confirmaciones específicas.",
    ]
    for item in principles:
        add_bullet(doc, item)

    add_heading(doc, "5. Modelo conceptual de relación entre interfaces", 1)
    add_body(
        doc,
        "Pirani organiza la gestión alrededor del riesgo como entidad central. Los demás componentes alimentan, contextualizan, mitigan, monitorean o reportan ese riesgo. Esa misma relación debe replicarse en el SGRLA-IHSS."
    )
    add_table(
        doc,
        ["Componente", "Relación principal", "Impacto funcional"],
        [
            ["Parametrización", "Define escalas de impacto, frecuencia, niveles de riesgo, variables de control y rangos de solidez.", "Todo cálculo depende de la versión vigente de parametrización."],
            ["Procesos", "Se asocian a riesgos.", "Permiten ubicar el riesgo dentro de la operación institucional."],
            ["Factores de riesgo", "Se asocian a riesgos y pueden clasificar clientes, contrapartes, canales, productos y jurisdicciones.", "Aportan contexto LA/FT y segmentación."],
            ["Riesgos", "Entidad central que conecta procesos, factores, causas, consecuencias, controles, planes, alertas y evaluaciones.", "Concentra impacto, frecuencia, riesgo inherente y riesgo residual."],
            ["Causas y consecuencias", "Explican origen y efecto del riesgo.", "Mejoran análisis, controles y planes de acción."],
            ["Controles", "Mitigan riesgos mediante diseño y ejecución.", "Reducen o modifican el riesgo residual según solidez."],
            ["Planes de acción", "Tratan riesgos o brechas de control.", "Gestionan responsables, fechas y seguimiento."],
            ["Evaluaciones", "Revisan periódicamente riesgos y controles.", "Generan ciclos de mejora y actualización."],
            ["Reportes e indicadores", "Leen los resultados vigentes.", "No capturan; consolidan información para dirección y cumplimiento."],
        ],
        [2100, 4300, 2960],
        font_size=8.8,
    )

    add_heading(doc, "6. Interfaces a replicar", 1)
    interfaces = [
        (
            "6.1 Dashboard y mapa de calor",
            "Debe ser una vista ejecutiva con tarjetas de conteo, acceso directo a parametrización, mapa de calor inherente/residual, distribución por niveles y acceso a reportes. La experiencia observada prioriza lectura rápida antes que captura de datos.",
            [
                "Tarjetas: procesos, factores, riesgos, controles, alertas, planes de acción y evaluaciones.",
                "Mapa de calor 5x5 con impacto y frecuencia, colores institucionales para niveles bajo, medio, alto y crítico.",
                "Botón Parametrización para ajustar escalas y advertir recálculo.",
                "Botón Ver más o Reportes para navegar a la vista analítica.",
            ],
        ),
        (
            "6.2 Parametrización del sistema",
            "Es la base de la metodología. Permite configurar mapa de calor, niveles de riesgo, pesos de impacto/frecuencia y calificación de controles. Debe existir antes de que el usuario cree riesgos productivos.",
            [
                "Mapa de calor con variables de impacto y frecuencia.",
                "Niveles de riesgo con rangos y colores.",
                "Controles con peso de Diseño y Ejecución.",
                "Variables de diseño, variables de ejecución y solidez.",
                "Confirmación especial cuando un cambio recalifique riesgos existentes.",
            ],
        ),
        (
            "6.3 Procesos",
            "Gestiona procesos organizacionales. La pantalla debe tener listado, búsqueda, estado vacío, creación por formulario guiado y asociaciones con riesgos, responsables y documentos.",
            [
                "Formulario Crear proceso con Información general, Caracterización y Asociaciones.",
                "Campos mínimos: nombre, tipo, categoría, descripción, adjuntos y responsables.",
                "Relación hacia riesgos para saber en qué proceso vive cada exposición.",
            ],
        ),
        (
            "6.4 Factores de riesgo",
            "Agrupa los factores LA/FT que Pirani presenta como clientes, contrapartes, canales, productos y jurisdicciones. En nuestro sistema se manejarán como catálogos internos de Matrices.",
            [
                "Pestañas por tipo de factor.",
                "Formulario de creación con tipo de identificación, número, tipo de persona y datos transaccionales cuando aplique.",
                "Confirmación al salir con cambios sin guardar.",
            ],
        ),
        (
            "6.5 Riesgos",
            "Es la interfaz central. Permite crear el riesgo, calificar impacto por componentes, calificar frecuencia, asociar procesos/factores/controles y visualizar la ubicación inherente.",
            [
                "Formulario por pasos: General, Asociaciones y Gestión.",
                "Impacto descompuesto en riesgo reputacional, legal, operativo y contagio, con promedio o regla parametrizada.",
                "Frecuencia seleccionada según escala vigente.",
                "Riesgo inherente calculado por matriz de impacto/frecuencia.",
                "Adjuntos, responsables, descripción y switch de continuidad de negocio si se aprueba conservarlo.",
            ],
        ),
        (
            "6.6 Causas y consecuencias",
            "Permite explicar por qué puede ocurrir un riesgo y qué efecto tendría. Debe funcionar como catálogo asociable, no como texto suelto dentro del riesgo.",
            [
                "Campos: nombre, tipo, categoría y descripción.",
                "Asociación con riesgos, controles y responsables.",
                "Uso posterior para reportes, análisis causal y priorización de controles.",
            ],
        ),
        (
            "6.7 Controles",
            "Evalúa la capacidad de mitigación del control. Pirani separa diseño y ejecución, y combina ambos para obtener solidez. Esa lógica debe replicarse en backend.",
            [
                "Diseño del control: tipo de control, ejecución, frecuencia, documentación, evidencia y responsables.",
                "Ejecución del control: cumplimiento real o desempeño operativo.",
                "Solidez calculada y usada para riesgo residual.",
                "Asociación con riesgos y causas/consecuencias.",
            ],
        ),
        (
            "6.8 Alertas de Matrices",
            "Son alertas internas de gestión de riesgos. Deben permitir registrar eventos inusuales, investigaciones o señales operativas dentro del módulo Matrices, sin integrarse con Monitoreo de Listas.",
            [
                "Datos: tipo de alerta, estado, medio de detección, valor, moneda, tipo de transacción, responsable, periodo y descripción de hechos.",
                "Asociaciones con procesos, riesgos, controles y factores.",
                "Estados propios de investigación y cierre.",
            ],
        ),
        (
            "6.9 Planes de acción",
            "Gestiona tratamiento y mejora. Debe tener planes, actividades, responsables, fechas, estados y evidencias.",
            [
                "Formulario: información general, actividades, asociaciones e historial.",
                "Relación con riesgos, controles, procesos y ejecutores.",
                "Seguimiento por avance, vencimientos y responsables.",
            ],
        ),
        (
            "6.10 Evaluaciones",
            "Permite revisar periódicamente riesgos y controles. Funciona como ciclo de actualización metodológica y operativa.",
            [
                "Plan de evaluación con nombre, descripción, fechas, responsable y tipo.",
                "Entidades a evaluar: riesgos, controles o combinación aprobada.",
                "Resultado esperado: actualización de estado, historial y posibles recálculos.",
            ],
        ),
        (
            "6.11 Documentos, reportes e indicadores",
            "Documentos actúa como soporte documental y adjuntos; reportes e indicadores son capas de lectura ejecutiva sobre información ya registrada.",
            [
                "Documentos: evidencias y archivos asociados a objetos del módulo.",
                "Reportes: dashboard ejecutivo, exportaciones PDF/Excel y filtros.",
                "Indicadores: KRI/KPI calculados desde riesgos, controles, planes y evaluaciones.",
            ],
        ),
    ]
    for heading, paragraph, bullets in interfaces:
        add_heading(doc, heading, 2)
        add_body(doc, paragraph)
        for bullet in bullets:
            add_bullet(doc, bullet)

    add_heading(doc, "7. Adaptación sobre el avance actual del proyecto", 1)
    add_body(
        doc,
        "El módulo actual de Matrices de Riesgos ya tiene avances útiles, pero el nuevo requerimiento cambia el modelo mental del usuario. La pantalla actual de captura y cálculo no debe eliminarse sin análisis; debe transformarse o reutilizarse como parte del motor y de las evaluaciones, mientras la experiencia final adopta la estructura Pirani."
    )
    add_table(
        doc,
        ["Elemento actual", "Uso recomendado", "Cambio requerido"],
        [
            ["Layout, autenticación y permisos", "Conservar.", "Agregar permisos por submódulo interno de Matrices si se requiere granularidad."],
            ["Servicios backend de cálculo", "Reutilizar como base.", "Extender para impacto/frecuencia, controles, solidez y versiones de parametrización."],
            ["Catálogos y criterios", "Reutilizar parcialmente.", "Convertir a parametrización de mapa de calor, variables y controles."],
            ["Dashboard y mapa de calor actual", "Reutilizar como punto de partida.", "Rediseñar hacia vista ejecutiva tipo Pirani con más tarjetas y navegación interna."],
            ["Captura de matriz", "No usar como flujo final único.", "Replantear como creación/evaluación de riesgos con asociaciones."],
            ["Reportes/exportaciones", "Conservar patrón técnico.", "Ajustar a nuevos objetos, filtros y auditoría de exportación."],
        ],
        [2300, 2500, 4560],
        font_size=8.8,
    )

    add_heading(doc, "8. Modelo de datos propuesto", 1)
    add_body(
        doc,
        "La réplica requiere ampliar el modelo de datos para que el riesgo sea la entidad central. La nomenclatura final puede alinearse con las tablas existentes RL_MR_*, pero conceptualmente se recomienda cubrir las siguientes entidades."
    )
    add_table(
        doc,
        ["Entidad lógica", "Propósito", "Relaciones principales"],
        [
            ["MTR_PARAM_MAPA_CALOR", "Versionar configuración de impacto, frecuencia y niveles.", "Riesgos, recálculos, auditoría."],
            ["MTR_NIVEL_RIESGO", "Definir rangos, colores y etiquetas.", "Mapa de calor y resultados."],
            ["MTR_PROCESO", "Registrar procesos institucionales.", "Riesgos, responsables, documentos."],
            ["MTR_FACTOR_RIESGO", "Registrar clientes, contrapartes, canales, productos y jurisdicciones.", "Riesgos y reportes."],
            ["MTR_RIESGO", "Entidad central de gestión.", "Procesos, factores, causas, controles, alertas, planes, evaluaciones."],
            ["MTR_RIESGO_IMPACTO_DET", "Guardar calificación por tipo de impacto.", "Riesgo y parametrización."],
            ["MTR_CAUSA_CONSECUENCIA", "Catalogar causas y consecuencias.", "Riesgos, controles."],
            ["MTR_CONTROL", "Registrar controles.", "Riesgos y evaluaciones."],
            ["MTR_CONTROL_CALIFICACION", "Guardar diseño, ejecución y solidez.", "Control, riesgo residual."],
            ["MTR_PLAN_ACCION / ACTIVIDAD", "Gestionar tratamiento.", "Riesgos, controles, responsables."],
            ["MTR_ALERTA_RIESGO", "Registrar alertas internas de Matrices.", "Riesgos, procesos, factores, controles."],
            ["MTR_EVALUACION_PLAN", "Planificar revisiones periódicas.", "Riesgos, controles, responsables."],
            ["MTR_DOCUMENTO / HISTORIAL", "Soportar adjuntos, auditoría y trazabilidad.", "Todos los objetos principales."],
        ],
        [2500, 3300, 3560],
        font_size=8.2,
    )

    add_heading(doc, "9. Motor de cálculo esperado", 1)
    add_body(
        doc,
        "La réplica debe calcular en backend, no en Angular. Angular debe capturar, validar y mostrar; el backend debe resolver el cálculo con la parametrización vigente y guardar snapshot para auditoría."
    )
    calc_rows = [
        ["Impacto final", "Promedio o regla parametrizada de impactos reputacional, legal, operativo y contagio.", "Riesgo inherente."],
        ["Frecuencia", "Escala seleccionada según parametrización vigente.", "Riesgo inherente."],
        ["Riesgo inherente", "Cruce impacto/frecuencia en mapa de calor.", "Base antes de controles."],
        ["Solidez de control", "Combinación de diseño y ejecución.", "Mitigación."],
        ["Riesgo residual", "Riesgo inherente ajustado por controles asociados y solidez.", "Priorización, planes, reportes."],
        ["Recálculo", "Se dispara por cambio de parametrización, cambio de controles o actualización de evaluación.", "Historial y auditoría obligatoria."],
    ]
    add_table(doc, ["Resultado", "Regla funcional", "Uso"], calc_rows, [2100, 4700, 2560], font_size=8.8)

    add_heading(doc, "10. Experiencia visual y responsive", 1)
    add_body(
        doc,
        "La mejora estética debe replicar patrones de Pirani sin cambiar la identidad visual del SGRLA-IHSS. La prioridad es limpieza, jerarquía, consistencia y operación cómoda en escritorio, tablet y móvil."
    )
    ux = [
        "Sidebar institucional con módulo activo y grupos claros.",
        "Topbar estable con usuario, salida y contexto del sistema.",
        "Estados vacíos con mensaje claro y botón de acción principal.",
        "Formularios por pasos con secciones: General, Asociaciones, Gestión e Historial cuando aplique.",
        "Textareas con contador y límite institucional. Si se mantiene el estándar reciente del sistema, usar 1000 caracteres.",
        "Confirmaciones específicas por acción: guardar, descargar PDF, descargar Excel, salir sin guardar, recalcular o eliminar.",
        "Tablas con filtros compactos, paginación, badges de estado y acciones con iconos.",
        "Diseño responsive con filtros apilables, tarjetas adaptables y formularios con scroll controlado.",
    ]
    for item in ux:
        add_bullet(doc, item)

    add_heading(doc, "11. Seguridad, auditoría y gobierno", 1)
    add_table(
        doc,
        ["Tema", "Regla recomendada"],
        [
            ["Roles y permisos", "Mantener control por módulo y evaluar permisos internos para parametrización, creación, aprobación, cierre, exportación y administración."],
            ["Auditoría", "Registrar usuario, fecha, acción, objeto, estado anterior, estado nuevo, datos relevantes y motivo cuando aplique."],
            ["Exportaciones", "Auditar PDF/Excel con filtros usados, fecha, usuario y módulo."],
            ["Parametrización", "Requerir confirmación específica por impacto de recálculo y conservar versión anterior."],
            ["Eliminaciones", "Preferir inactivación/cierre lógico antes que borrado físico."],
            ["Adjuntos", "Validar tipo, tamaño, usuario, fecha y relación con objeto principal."],
        ],
        [2600, 6760],
        font_size=9,
    )

    add_heading(doc, "12. Plan de desarrollo propuesto", 1)
    phases = [
        ["Fase 1", "Validación de alcance", "Confirmar con cliente qué interfaces Pirani entran en primera entrega y cuáles quedan por etapa."],
        ["Fase 2", "Diseño UX funcional", "Wireframes internos con colores IHSS, navegación, formularios, estados, filtros y confirmaciones."],
        ["Fase 3", "Modelo de datos", "Diseñar migraciones/tablas RL_MR_* para entidades centrales, relaciones, parametrización y auditoría."],
        ["Fase 4", "Backend y contratos", "APIs para parametrización, riesgos, procesos, factores, controles, planes, evaluaciones, reportes y exportaciones."],
        ["Fase 5", "Frontend Angular", "Construir shell interno tipo Pirani, submódulos, formularios por pasos, dashboard y tablas responsive."],
        ["Fase 6", "Motor de cálculo", "Implementar impacto, frecuencia, inherente, solidez, residual, recálculo y snapshots."],
        ["Fase 7", "Reportería e indicadores", "PDF, Excel, dashboards, filtros y auditoría de descarga."],
        ["Fase 8", "QA y UAT", "Pruebas funcionales, responsive, permisos, auditoría, recálculo y validación con cliente/jefatura."],
    ]
    add_table(doc, ["Fase", "Nombre", "Resultado esperado"], phases, [1100, 2500, 5760], font_size=8.8)

    add_heading(doc, "13. Riesgos del replanteamiento", 1)
    risks = [
        ["Alcance demasiado amplio", "Pirani cubre varias capacidades. Si se intenta entregar todo en una sola etapa, puede crecer el tiempo de desarrollo.", "Definir MVP aprobado y fases posteriores."],
        ["Confusión con Monitoreo de Listas", "La palabra alerta puede hacer pensar que hay dependencia entre módulos.", "Documentar y validar que Alertas de Matrices son independientes."],
        ["Cálculo residual no definido por cliente", "Pirani muestra solidez de controles, pero la fórmula exacta institucional debe aprobarse.", "Aprobar metodología antes de programar cálculo final."],
        ["Parametrización cambia resultados", "Editar escalas o pesos puede recalificar riesgos existentes.", "Versionar configuración y exigir confirmación."],
        ["Avance actual vs nuevo diseño", "Parte de lo construido puede no calzar con la nueva experiencia.", "Reutilizar backend y patrones, rediseñar UI con control de alcance."],
    ]
    add_table(doc, ["Riesgo", "Descripción", "Mitigación"], risks, [2300, 3600, 3460], font_size=8.5)

    add_heading(doc, "14. Decisiones pendientes para aprobación", 1)
    decisions = [
        "Confirmar si la primera entrega debe incluir todos los submódulos observados en Pirani o solo Dashboard, Parametrización, Riesgos, Controles y Reportes.",
        "Confirmar si los factores de riesgo se manejarán como clientes/contrapartes/canales/productos/jurisdicciones desde la primera versión.",
        "Aprobar fórmula institucional de riesgo residual y efecto de controles.",
        "Definir estados finales para riesgos, controles, alertas, planes y evaluaciones.",
        "Aprobar si Documentos será solo adjuntos/evidencias dentro de Matrices o un centro documental completo.",
        "Confirmar límite de caracteres de textareas. Recomendación: mantener 1000 caracteres para consistencia con el sistema actual.",
    ]
    for idx, item in enumerate(decisions, start=1):
        add_number(doc, item)

    add_heading(doc, "15. Cierre del análisis", 1)
    add_body(
        doc,
        "La réplica de Pirani Risks en el módulo Matrices de Riesgos es viable y recomendable si se maneja como rediseño funcional controlado. La estructura observada aporta una experiencia más completa que la captura de matriz tradicional: inicia en parametrización, organiza los riesgos como entidad central, vincula procesos y factores, mide controles, calcula residual, gestiona planes y entrega reportes."
    )
    add_body(
        doc,
        "La condición crítica es no mezclar este rediseño con Monitoreo de Listas. Ambos módulos pueden convivir en el SGRLA-IHSS, compartir layout y seguridad, pero no deben depender entre sí ni compartir flujos funcionales."
    )
    add_callout(
        doc,
        "Recomendación final",
        "Presentar este documento al cliente y jefatura como base de aprobación del nuevo alcance. Una vez validado, se debe reemplazar la hoja de ruta anterior de Matrices por una nueva planificación por fases antes de iniciar programación.",
        "warning",
    )

    doc.core_properties.title = "Análisis de Réplica Pirani - Matrices de Riesgos SGRLA-IHSS"
    doc.core_properties.subject = "Documento funcional y diseño de replanteamiento"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Generado como análisis, sin implementación de código."
    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_document())
