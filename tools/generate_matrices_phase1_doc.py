from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_ROOT = ROOT / "docs"
MATRICES_DIR = next(DOCS_ROOT.rglob("3. Módulo Matrices de Riesgos"))
OUT_DIR = MATRICES_DIR / "Fase 1 - Gobierno y Control Documental"
OUT_FILE = OUT_DIR / "Fase_1_Acta_Inicio_Control_Documental_Matrices_Riesgos_SGRLA_IHSS_FINAL.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run("IHSS - SGRLA/FT | Módulo Matrices de Riesgos | Fase 1")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_after = Pt(0)
    run = footer_p.add_run("Documento de trabajo para revisión y aprobación")
    set_run_font(run, size=9, color=MUTED)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("DOCUMENTO DE FASE 1")
    set_run_font(run, size=10.5, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Acta de Inicio y Control Documental")
    set_run_font(run, size=22, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Módulo Matrices de Riesgos - Sistema de Gestión de Riesgos LA/FT IHSS")
    set_run_font(run, size=13, color=MUTED)

    meta = [
        ("Proyecto", "RIESGO_LAVADO - IHSS"),
        ("Módulo", "3. Matrices de Riesgos"),
        ("Fase", "Fase 1. Gobierno, alcance y control documental"),
        ("Versión", "1.3"),
        ("Fecha", date.today().strftime("%d/%m/%Y")),
        ("Estado", "Aprobado con ampliación de trazabilidad visual del requerimiento"),
        ("Fuente", "Plan de fases, análisis maestro del módulo y lineamientos cerrados del módulo Monitoreo de Listas"),
        ("Ubicación", "docs/3. Módulo Matrices de Riesgos/Fase 1 - Gobierno y Control Documental"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, [2300, 7060])
    for row, (label, value) in zip(table.rows, meta):
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)


def add_change_control_table(doc: Document) -> None:
    rows = [
        ("1.0", "01/07/2026", "Documento inicial de Fase 1 para revisión.", "Javier Mejía", "Revisión"),
        ("1.1", date.today().strftime("%d/%m/%Y"), "Pulido institucional: nombre uniforme, LA/FT, RACI ampliado, estructura documental y decisión de vigencia.", "Javier Mejía", "Vigente para aprobación"),
        ("1.2", date.today().strftime("%d/%m/%Y"), "Aprobación formal de Fase 1 y asignación de todos los roles a Javier Mejía.", "Javier Mejía", "Aprobado"),
        ("1.3", date.today().strftime("%d/%m/%Y"), "Versión final consolidada: incorporación de imágenes del requerimiento del cliente, módulo y submódulos, reglas institucionales y nombre funcional normalizado.", "Javier Mejía", "Aprobado"),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [900, 1350, 4300, 1500, 1310])
    for idx, text in enumerate(["Versión", "Fecha", "Descripción del cambio", "Responsable", "Estado"]):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
        table.rows[0].cells[idx].paragraphs[0].add_run(text).bold = True
    set_repeat_table_header(table.rows[0])
    for version, change_date, description, owner, status in rows:
        cells = table.add_row().cells
        cells[0].text = version
        cells[1].text = change_date
        cells[2].text = description
        cells[3].text = owner
        cells[4].text = status


def add_status_table(doc: Document) -> None:
    rows = [
        ("Acta de inicio del módulo", "Aprobada", "Aprobación formal registrada"),
        ("Matriz de responsables y aprobadores", "Definida y asignada", "Todos los roles asignados a Javier Mejía"),
        ("Inventario documental controlado", "Levantado sobre la carpeta actual", "Activo"),
        ("Lista de decisiones abiertas", "Incluida en este documento", "Activo"),
        ("Criterios de cierre de Fase 1", "Definidos y verificables", "Activo"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [3000, 3300, 3060])
    for idx, text in enumerate(["Entregable", "Estado actual", "Control requerido"]):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
        table.rows[0].cells[idx].paragraphs[0].add_run(text).bold = True
    set_repeat_table_header(table.rows[0])
    for item, status, control in rows:
        cells = table.add_row().cells
        cells[0].text = item
        cells[1].text = status
        cells[2].text = control


def add_matrix_table(doc: Document) -> None:
    rows = [
        ("Usuario / dueño funcional", "Javier Mejía", "Aprobar alcance, metodología y versión final", "Aprobador", "A"),
        ("Cumplimiento / Riesgo LA/FT", "Javier Mejía", "Definir criterios, escalas, variables y reglas de negocio", "Responsable funcional", "R"),
        ("Desarrollo backend", "Javier Mejía", "Diseñar servicios, motor de cálculo, endpoints y auditoría", "Responsable técnico", "R"),
        ("Desarrollo frontend", "Javier Mejía", "Diseñar pantallas, validaciones visibles y flujos por rol", "Responsable técnico", "R"),
        ("DBA / base de datos", "Javier Mejía", "Validar modelo RL_MR_*, scripts idempotentes y ejecución segura", "Responsable técnico", "R"),
        ("Seguridad / accesos", "Javier Mejía", "Validar módulo, permisos, roles, auditoría y controles de acceso", "Control técnico", "C"),
        ("Auditoría / documentación", "Javier Mejía", "Revisar trazabilidad, evidencias, reportes y control documental", "Consulta / control", "C"),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [1600, 2300, 3400, 1400, 660])
    for idx, text in enumerate(["Rol", "Nombre / área asignada", "Responsabilidad principal", "Tipo", "RACI"]):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
        table.rows[0].cells[idx].paragraphs[0].add_run(text).bold = True
    set_repeat_table_header(table.rows[0])
    for role, assigned, responsibility, kind, raci in rows:
        cells = table.add_row().cells
        cells[0].text = role
        cells[1].text = assigned
        cells[2].text = responsibility
        cells[3].text = kind
        cells[4].text = raci
        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_inventory_table(doc: Document) -> None:
    rows = [
        ("Analisis_Final_Maestro_Modulo_3_Matrices_Riesgos_SGRLA_IHSS.docx", "Técnico / maestro", "Base de análisis", "Vigente como fuente"),
        ("Plan_Fases_Modulo_3_Matrices_Riesgos_SGRLA_IHSS.docx", "Planificación", "Plan maestro por fases", "Vigente como ruta de trabajo"),
        ("Fase_1_Acta_Inicio_Control_Documental_Matrices_Riesgos_SGRLA_IHSS.docx", "Control de fase", "Acta e inventario Fase 1", "Generado para aprobación"),
        ("Próxima fase: Documento metodológico funcional", "Cliente / funcional", "Fase 2", "Habilitado después del cierre aprobado de Fase 1"),
        ("Próxima fase: Documento técnico desarrollador", "Desarrollador", "Fase 3-6", "Depende de metodología aprobada"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [3300, 1800, 2400, 1860])
    for idx, text in enumerate(["Documento", "Tipo", "Uso", "Estado"]):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
        table.rows[0].cells[idx].paragraphs[0].add_run(text).bold = True
    set_repeat_table_header(table.rows[0])
    for doc_name, kind, use, status in rows:
        cells = table.add_row().cells
        cells[0].text = doc_name
        cells[1].text = kind
        cells[2].text = use
        cells[3].text = status


def add_decisions_table(doc: Document) -> None:
    rows = [
            ("D-01", "El Módulo Matrices de Riesgos se construirá como motor metodológico y no como pantalla aislada.", "Aprobada para planificación", "Fase 1"),
        ("D-02", "Los cálculos de riesgo inherente, mitigación y residual vivirán en backend.", "Aprobada", "Fase 2-5"),
        ("D-03", "Las tablas físicas usarán prefijo RL_MR_* salvo ajuste DBA aprobado.", "Aprobada como recomendación", "Fase 3"),
        ("D-04", "No se crearán scripts experimentales dentro del flujo aprobado de base de datos.", "Aprobada por política", "Todas"),
        ("D-05", "La fórmula oficial, escalas y ponderaciones quedan pendientes de aprobación funcional.", "Abierta", "Fase 2"),
        ("D-06", "La integración futura con DNP no se desarrollará en Fase 1.", "Fuera de alcance inicial", "Fase futura"),
        ("D-07", "Una vez aprobado, este documento será la versión vigente de Fase 1; documentos previos quedarán como referencia o respaldo, no como fuente principal.", "Aprobada para control documental", "Fase 1"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [1000, 5000, 1900, 1460])
    for idx, text in enumerate(["ID", "Decisión", "Estado", "Fase"]):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
        table.rows[0].cells[idx].paragraphs[0].add_run(text).bold = True
    set_repeat_table_header(table.rows[0])
    for decision_id, decision, status, phase in rows:
        cells = table.add_row().cells
        cells[0].text = decision_id
        cells[1].text = decision
        cells[2].text = status
        cells[3].text = phase


def add_risks_table(doc: Document) -> None:
    rows = [
        ("Metodología no aprobada", "Alta", "No diseñar tablas definitivas ni pantallas finales sin criterios, escalas y ponderaciones aprobadas."),
        ("Cálculos duplicados en frontend", "Alta", "Centralizar el cálculo en backend y dejar Angular solo como captura, consulta y visualización."),
        ("Documentos duplicados o borradores mezclados", "Media", "Mantener carpeta de Fase 1 separada y versiones Word revisadas antes de repositorio."),
        ("Auditoría incompleta", "Alta", "Todo cálculo, aprobación, cambio de estado, exportación y recálculo debe auditarse."),
        ("Modelo de datos rígido", "Media", "Diseñar RL_MR_* versionado, parametrizable y con snapshot de resultados cerrados."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [2700, 1100, 5560])
    for idx, text in enumerate(["Riesgo", "Nivel", "Control desde Fase 1"]):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
        table.rows[0].cells[idx].paragraphs[0].add_run(text).bold = True
    set_repeat_table_header(table.rows[0])
    for risk, level, control in rows:
        cells = table.add_row().cells
        cells[0].text = risk
        cells[1].text = level
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[2].text = control


def add_checklist(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
        table.rows[0].cells[idx].paragraphs[0].add_run(header).bold = True
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_title_block(doc)

    doc.add_heading("0. Control de cambios", level=1)
    add_change_control_table(doc)

    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "Se da inicio formal a la Fase 1 del Módulo Matrices de Riesgos del SGRLA/FT. "
        "Esta fase no autoriza todavía programación funcional del motor de cálculo; su propósito es cerrar gobierno, alcance, responsables, control documental y reglas de trabajo antes de pasar al levantamiento metodológico."
    )
    doc.add_paragraph(
        "El módulo debe construirse como una base institucional, auditable, versionable y parametrizable para evaluar riesgo inherente, controles mitigantes y riesgo residual. "
        "La decisión técnica principal se mantiene: los cálculos sensibles vivirán en backend y la interfaz solo capturará, validará visualmente y presentará resultados calculados por el servidor."
    )

    doc.add_heading("2. Objetivo de la Fase 1", level=1)
    add_checklist(
        doc,
        [
            "Confirmar el alcance exclusivo del Módulo Matrices de Riesgos.",
            "Separar responsabilidades funcionales, técnicas, DBA, seguridad, auditoría y aprobación.",
            "Definir el inventario documental vigente del módulo y evitar documentos duplicados o temporales.",
            "Registrar decisiones abiertas antes de metodología, base de datos, backend, frontend o pruebas.",
            "Dejar criterios claros para cerrar Fase 1 y habilitar Fase 2.",
        ],
    )

    doc.add_heading("3. Alcance inicial aprobado para planificación", level=1)
    doc.add_paragraph(
        "El alcance de planificación incluye metodología, parametrización, cálculo, seguimiento, auditoría y reportes del Módulo Matrices de Riesgos. "
        "No se mezcla con Monitoreo de Listas, aunque sí debe integrarse con seguridad, auditoría, evidencias, reportes y usuarios ya cerrados."
    )
    add_checklist(
        doc,
        [
            "Modelos metodológicos versionados.",
            "Factores, criterios, escalas, rangos y ponderaciones.",
            "Matrices por patrono, proveedor, empleado, área, proceso, caso positivo o matriz institucional.",
            "Cálculo de riesgo inherente, controles mitigantes y riesgo residual.",
            "Planes de acción obligatorios cuando el residual sea alto o crítico.",
            "Historial funcional, auditoría transversal, reportes y exportaciones auditadas.",
        ],
    )

    doc.add_heading("4. Fuera de alcance de Fase 1", level=1)
    add_checklist(
        doc,
        [
            "Construcción de pantallas finales sin metodología aprobada.",
            "Creación de tablas definitivas sin modelo validado por DBA.",
            "Implementación de fórmulas oficiales sin aprobación funcional.",
            "Integración automática con sistemas externos o DNP.",
            "Subida al repositorio de documentos no revisados o versiones de borrador.",
        ],
    )

    doc.add_heading("5. Entregables de Fase 1", level=1)
    add_status_table(doc)

    doc.add_heading("6. Matriz de responsables y aprobadores", level=1)
    doc.add_paragraph(
        "La siguiente matriz define responsabilidades por rol y deja explícito que los nombres o áreas responsables quedan pendientes de designación institucional antes de declarar la fase formalmente aprobada."
    )
    add_matrix_table(doc)

    doc.add_heading("7. Estructura documental recomendada", level=1)
    doc.add_paragraph(
        "La Fase 1 debe conservar una ubicación única y reconocible dentro de la carpeta del módulo. La ruta vigente para este entregable es:"
    )
    route = doc.add_paragraph()
    route.paragraph_format.left_indent = Inches(0.25)
    route.paragraph_format.space_after = Pt(8)
    run = route.add_run("docs/3. Módulo Matrices de Riesgos/Fase 1 - Gobierno y Control Documental/")
    set_run_font(run, size=10.5, color=INK, bold=True)
    doc.add_paragraph(
        "Dentro de esta carpeta solo deben quedar documentos finales o en revisión formal de la fase. No deben guardarse capturas temporales, archivos de render, borradores duplicados ni documentos experimentales."
    )

    doc.add_heading("8. Inventario documental controlado", level=1)
    doc.add_paragraph(
        "La carpeta del módulo debe mantener documentos claros, sin borradores mezclados ni archivos temporales. Este inventario deja identificadas las fuentes vigentes y los documentos pendientes que nacerán en fases posteriores."
    )
    add_inventory_table(doc)

    doc.add_heading("9. Tratamiento de imágenes del requerimiento del cliente", level=1)
    doc.add_paragraph(
        "Las imágenes suministradas en el requerimiento del cliente se manejarán como insumo funcional obligatorio para entender la estructura esperada del módulo. "
        "No se tratarán como diseño visual definitivo ni como copia exacta de pantalla; se convertirán en reglas, campos, secciones, cálculos, reportes y flujos aprobables dentro de la documentación de cada fase."
    )
    add_checklist(
        doc,
        [
            "Cada imagen debe quedar asociada a un submódulo funcional de Matrices de Riesgos.",
            "Los campos visibles en las imágenes deben convertirse en campos funcionales o reglas metodológicas documentadas.",
            "La estructura de Excel del cliente se usará como referencia para identificar alimentación, calificación, mitigadores, residual, plan de acción, mapa de calor y reportería.",
            "Las pantallas de referencia de mercado se usarán únicamente como guía de experiencia funcional; el diseño final debe respetar la arquitectura y controles del sistema IHSS.",
            "Toda definición resultante debe quedar en Word antes de pasar a base de datos, backend o frontend.",
        ],
    )
    add_table(
        doc,
        ["Referencia visual", "Submódulo asociado", "Uso documental"],
        [
            ("Matriz factor / variables / ponderación / decisión / observaciones", "Perfilamiento de factores de riesgo", "Define la estructura de variables por factor institucional y su decisión funcional."),
            ("Identificación del riesgo", "Identificación de factores de riesgo", "Define campos de evento, probabilidad, impacto, puntaje y nivel de riesgo inherente."),
            ("Evaluación de mitigadores", "Controles mitigantes", "Define existencia de control, calidad, periodicidad, oportunidad, automatización, procedimientos y residual."),
            ("Plan de acción", "Planes de acción", "Define actividades, responsables, periodicidad, fechas, medios de prueba, observaciones y estado."),
            ("Crear riesgo / crear control", "Captura funcional", "Define navegación esperada, adjunto de evidencia, responsables, asociaciones y vista de calificación."),
            ("Mapa de calor y gráficos", "Mapa de calor y reportería dinámica", "Define filtros, riesgo inherente/residual, comparativos, descargas y gráficos estadísticos."),
        ],
        [2200, 3000, 4160],
    )

    doc.add_heading("10. Estructura por módulo y submódulo", level=1)
    doc.add_paragraph(
        "La estructura del Módulo Matrices de Riesgos debe respetar como regla funcional fija los tres factores institucionales definidos por el cliente: "
        "Proveedores 50%, Clientes/Patronos 25% y Empleados 25%. "
        "Cada factor será medido de forma independiente mediante variables de riesgo definidas por la Sección de Cumplimiento."
    )
    add_table(
        doc,
        ["Módulo", "Submódulo", "Resultado esperado"],
        [
            ("Matrices de Riesgos", "Identificación de factores de riesgo", "Registrar evento de riesgo, probabilidad, impacto, puntajes, nivel inherente y evidencia."),
            ("Matrices de Riesgos", "Perfilamiento / scoring", "Calificar Proveedores, Clientes/Patronos y Empleados con variables definidas por Cumplimiento."),
            ("Matrices de Riesgos", "Matriz resumen por factor", "Consolidar resultado individual por factor institucional y resumen institucional."),
            ("Matrices de Riesgos", "Mapa de calor", "Mostrar riesgo inherente y residual por factor, proceso, fecha, nivel y comparativo histórico."),
            ("Matrices de Riesgos", "Reportería estadística dinámica", "Generar gráficos, reportes, filtros y exportaciones auditadas."),
            ("Matrices de Riesgos", "Planes de acción", "Gestionar acciones, responsables, fechas, medios de prueba, observaciones y estado."),
        ],
        [2100, 3000, 4260],
    )

    doc.add_heading("11. Decisiones iniciales y decisiones abiertas", level=1)
    add_decisions_table(doc)

    doc.add_heading("12. Reglas obligatorias heredadas del cierre anterior", level=1)
    add_checklist(
        doc,
        [
            "Todo endpoint crítico debe tener autorización por módulo.",
            "Todo cálculo sensible debe ejecutarse en backend.",
            "Todo cambio de estado, aprobación, recálculo, exportación, impresión y consulta sensible debe auditarse.",
            "Toda eliminación lógica debe exigir motivo obligatorio y conservar trazabilidad.",
            "Todo documento debe estar en Word, en español latinoamericano, con versión, fecha, estado y fuente.",
            "No se suben borradores al repositorio sin revisión y aprobación.",
        ],
    )

    doc.add_heading("13. Riesgos de inicio y controles", level=1)
    add_risks_table(doc)

    doc.add_heading("14. Criterios de aceptación para cerrar Fase 1", level=1)
    add_checklist(
        doc,
        [
            "El alcance del módulo está confirmado y separado de Monitoreo de Listas.",
            "Los documentos vigentes están ubicados en la carpeta correcta del módulo.",
            "Existe acta de inicio con versión, fecha, estado, control de cambios y fuente.",
            "Existe inventario documental controlado.",
            "Todos los responsables por rol quedan asignados a Javier Mejía.",
            "Las decisiones abiertas para Fase 2 están registradas.",
            "No existen documentos temporales, capturas temporales o borradores mezclados en la carpeta de Fase 1.",
            "La decisión de vigencia documental está registrada para evitar fuentes duplicadas.",
            "Las imágenes del requerimiento del cliente quedan gobernadas como insumo funcional y trazadas hacia submódulos.",
        ],
    )

    doc.add_heading("15. Condiciones para iniciar Fase 2", level=1)
    doc.add_paragraph(
        "La Fase 2 puede iniciar cuando el usuario confirme la aprobación funcional del alcance de Fase 1 y autorice el levantamiento metodológico. "
        "En Fase 2 se deberá documentar la metodología LA/FT completa: factores, criterios, escalas, pesos, fórmulas, rangos, estados, aprobaciones y reglas de negocio."
    )

    doc.add_heading("16. Control de aprobación", level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [2500, 3430, 3430])
    for idx, text in enumerate(["Rol", "Nombre / área", "Firma / aprobación"]):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
        table.rows[0].cells[idx].paragraphs[0].add_run(text).bold = True
    for row_idx, row_data in enumerate(
        [
            ("Dueño funcional", "Javier Mejía", "Aprobado el 02/07/2026"),
            ("Responsable técnico", "Javier Mejía", "Aprobado el 02/07/2026"),
            ("Aprobación final", "Javier Mejía", "Aprobado el 02/07/2026"),
        ],
        start=1,
    ):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value

    props = doc.core_properties
    props.title = "Fase 1 - Acta de Inicio y Control Documental - Matrices de Riesgos"
    props.subject = "Sistema de Gestión de Riesgos LA/FT IHSS"
    props.keywords = "IHSS, SGRLA, Matrices de Riesgos, Fase 1, control documental"
    props.comments = "Documento elaborado desde el repositorio local RIESGO_LAVADO para revisión y aprobación."
    props.author = "Javier Mejía"

    doc.save(OUT_FILE)


if __name__ == "__main__":
    build_document()
    print(OUT_FILE)
