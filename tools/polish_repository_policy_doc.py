from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "7. Politica de Repositorio" / "Politica_Repositorio_SGRLA_IHSS.docx"


REPLACEMENTS = {
    "Paquete full vigente por módulo, separado cliente/desarrollador.": "Paquete completo vigente por módulo, separado entre cliente y desarrollador.",
    "Paquete full vigente por m\u00c3\u00b3dulo, separado cliente/desarrollador.": "Paquete completo vigente por módulo, separado entre cliente y desarrollador.",
    "9. Checklist previo a commit o subida": "9. Lista de verificación previa a commit o subida",
    "La versión final fue aprobada o esta marcada claramente como pendiente de aprobación.": "La versión final fue aprobada o está marcada claramente como pendiente de aprobación.",
    "Establecer una política clara para mantener el repositorio ordenado, evitar duplicados, controlar documentos finales y asegurar que cada entrega conserve separación entre información de cliente y documentación técnica para desarrolladores.": "Establecer una política clara para mantener el repositorio ordenado, evitar duplicados, controlar documentos finales y asegurar que cada entrega conserve separación entre la información para cliente y la documentación técnica para desarrolladores.",
    "capturas temporales, documentos sin revisar, archivos generados sin validación o mezcla de audiencia cliente/desarrollador.": "capturas temporales, documentos sin revisar, archivos generados sin validación o mezcla de audiencias cliente/desarrollador.",
    "asegurar que cada entrega conserve separacion entre información de cliente y documentación técnica para desarrolladores.": "asegurar que cada entrega conserve separación entre la información para cliente y la documentación técnica para desarrolladores.",
}


def replace_text(text):
    for source, target in REPLACEMENTS.items():
        text = text.replace(source, target)
    return text


def normalize_runs(runs):
    changed = 0
    for run in runs:
        new_text = replace_text(run.text)
        if new_text != run.text:
            run.text = new_text
            changed += 1
    if runs:
        joined = "".join(run.text for run in runs)
        normalized = replace_text(joined)
        if normalized != joined and len(normalized) == len(joined):
            pos = 0
            for run in runs:
                length = len(run.text)
                run.text = normalized[pos : pos + length]
                pos += length
            changed += 1
    return changed


def process_container(container):
    changed = 0
    for paragraph in container.paragraphs:
        changed += normalize_runs(paragraph.runs)
        current = paragraph.text
        replacement = replace_text(current)
        if replacement != current:
            style = paragraph.style
            alignment = paragraph.alignment
            paragraph.clear()
            paragraph.style = style
            paragraph.alignment = alignment
            paragraph.add_run(replacement)
            changed += 1
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                changed += process_container(cell)
    return changed


def main():
    doc = Document(DOC)
    changed = process_container(doc)
    for section in doc.sections:
        changed += process_container(section.header)
        changed += process_container(section.footer)
    if changed:
        doc.save(DOC)
    print(f"UPDATED_BLOCKS={changed}")


if __name__ == "__main__":
    main()
