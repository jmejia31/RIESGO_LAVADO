import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


REPLACEMENTS = {
    "Politica": "Política",
    "politica": "política",
    "Gestion": "Gestión",
    "gestion": "gestión",
    "Modulo": "Módulo",
    "modulo": "módulo",
    "Modulos": "Módulos",
    "modulos": "módulos",
    "Documentacion": "Documentación",
    "documentacion": "documentación",
    "Version": "Versión",
    "version": "versión",
    "Tecnica": "Técnica",
    "tecnica": "técnica",
    "Tecnico": "Técnico",
    "tecnico": "técnico",
    "Tecnicos": "Técnicos",
    "tecnicos": "técnicos",
    "Tecnicas": "Técnicas",
    "tecnicas": "técnicas",
    "Configuracion": "Configuración",
    "configuracion": "configuración",
    "Auditoria": "Auditoría",
    "auditoria": "auditoría",
    "Ejecucion": "Ejecución",
    "ejecucion": "ejecución",
    "Aprobacion": "Aprobación",
    "aprobacion": "aprobación",
    "Revision": "Revisión",
    "revision": "revisión",
    "Validacion": "Validación",
    "validacion": "validación",
    "Informacion": "Información",
    "informacion": "información",
    "Operacion": "Operación",
    "operacion": "operación",
    "Eliminacion": "Eliminación",
    "eliminacion": "eliminación",
    "Accion": "Acción",
    "accion": "acción",
    "Acciones": "Acciones",
    "acciones": "acciones",
    "Critica": "Crítica",
    "critica": "crítica",
    "Criticas": "Críticas",
    "criticas": "críticas",
    "Critico": "Crítico",
    "critico": "crítico",
    "Criticos": "Críticos",
    "criticos": "críticos",
    "Fisico": "Físico",
    "fisico": "físico",
    "Fisica": "Física",
    "fisica": "física",
    "Logica": "Lógica",
    "logica": "lógica",
    "Logico": "Lógico",
    "logico": "lógico",
    "Autorizacion": "Autorización",
    "autorizacion": "autorización",
    "Exportacion": "Exportación",
    "exportacion": "exportación",
    "Generacion": "Generación",
    "generacion": "generación",
    "Instalacion": "Instalación",
    "instalacion": "instalación",
    "Actualizacion": "Actualización",
    "actualizacion": "actualización",
    "Ubicacion": "Ubicación",
    "ubicacion": "ubicación",
    "Clasificacion": "Clasificación",
    "clasificacion": "clasificación",
    "Calificacion": "Calificación",
    "calificacion": "calificación",
    "Parametrizacion": "Parametrización",
    "parametrizacion": "parametrización",
    "Autenticacion": "Autenticación",
    "autenticacion": "autenticación",
    "Administracion": "Administración",
    "administracion": "administración",
    "Descripcion": "Descripción",
    "descripcion": "descripción",
    "Separacion": "Separación",
    "separacion": "separación",
    "Convencion": "Convención",
    "convencion": "convención",
    "Aplicacion": "Aplicación",
    "aplicacion": "aplicación",
    "Confusion": "Confusión",
    "confusion": "confusión",
    "Operacion": "Operación",
    "operacion": "operación",
    "Decision": "Decisión",
    "decision": "decisión",
    "Proteccion": "Protección",
    "proteccion": "protección",
    "Compilacion": "Compilación",
    "compilacion": "compilación",
    "Creacion": "Creación",
    "creacion": "creación",
    "Edicion": "Edición",
    "edicion": "edición",
    "Sesion": "Sesión",
    "sesion": "sesión",
    "contrasena": "contraseña",
    "Contrasena": "Contraseña",
    "pagina": "página",
    "Pagina": "Página",
    "Indice": "Índice",
    "indice": "índice",
    "Analisis": "Análisis",
    "analisis": "análisis",
    "Codigo": "Código",
    "codigo": "código",
    "Practica": "Práctica",
    "practica": "práctica",
    "Practicas": "Prácticas",
    "practicas": "prácticas",
    "Ambito": "Ámbito",
    "ambito": "ámbito",
    "Proposito": "Propósito",
    "proposito": "propósito",
    "Unica": "Única",
    "unica": "única",
    "Unico": "Único",
    "unico": "único",
    "Area": "Área",
    "area": "área",
    "Areas": "Áreas",
    "areas": "áreas",
    "Pais": "País",
    "pais": "país",
    "tambien": "también",
    "Tambien": "También",
    "despues": "después",
    "Despues": "Después",
    "mas": "más",
    "Mas": "Más",
    "periodo": "período",
    "Periodo": "Período",
    "diagnostico": "diagnóstico",
    "Diagnostico": "Diagnóstico",
    "metodo": "método",
    "Metodo": "Método",
    "dinamico": "dinámico",
    "Dinamico": "Dinámico",
    "logicamente": "lógicamente",
    "Logicamente": "Lógicamente",
    "numero": "número",
    "Numero": "Número",
    "automatico": "automático",
    "Automatico": "Automático",
    "manuales": "manuales",
    "juridicas": "jurídicas",
    "Juridicas": "Jurídicas",
    "natural": "natural",
    "naturales": "naturales",
    "inclusion": "inclusión",
    "Inclusion": "Inclusión",
    "exclusion": "exclusión",
    "Exclusion": "Exclusión",
    "impresion": "impresión",
    "Impresion": "Impresión",
    "visualizacion": "visualización",
    "Visualizacion": "Visualización",
    "imagenes": "imágenes",
    "Imagenes": "Imágenes",
    "genericos": "genéricos",
    "Genericos": "Genéricos",
    "numerico": "numérico",
    "Numerico": "Numérico",
    "estan": "están",
    "Estan": "Están",
    "esten": "estén",
    "Esten": "Estén",
    "Ningun": "Ningún",
    "ningun": "ningún",
    "util": "útil",
    "Util": "Útil",
    "Busqueda": "Búsqueda",
    "busqueda": "búsqueda",
    "especifica": "específica",
    "Especifica": "Específica",
    "Guias": "Guías",
    "guias": "guías",
    "Politicas": "Políticas",
    "politicas": "políticas",
    "implementacion": "implementación",
    "Implementacion": "Implementación",
    "descarga": "descarga",
    "obligatoria": "obligatoria",
    "Idempotente": "Idempotente",
    "idempotente": "idempotente",
    "secuencia": "secuencia",
    "Secuencia": "Secuencia",
    "migracion": "migración",
    "Migracion": "Migración",
}


TECHNICAL_HINTS = (
    "http://",
    "https://",
    "C:/",
    "C:\\",
    "/api/",
    "RL_",
    ".cs",
    ".ts",
    ".html",
    ".sql",
    ".json",
    ".docx",
    ".png",
    "Controller",
    "Service",
    "Repository",
    "ModuloAuthorize",
)


PATTERNS = [
    (re.compile(rf"(?<![A-Za-z0-9_]){re.escape(src)}(?![A-Za-z0-9_])"), dst)
    for src, dst in sorted(REPLACEMENTS.items(), key=lambda item: len(item[0]), reverse=True)
]


def replace_text(text):
    if not text:
        return text, 0
    original = text
    for pattern, replacement in PATTERNS:
        text = pattern.sub(replacement, text)
    return text, int(text != original)


def normalize_runs(runs):
    changes = 0
    for run in runs:
        text = run.text
        if not text:
            continue
        new_text, changed = replace_text(text)
        if changed:
            run.text = new_text
            changes += changed
    return changes


def normalize_paragraph(paragraph):
    changes = normalize_runs(paragraph.runs)
    if not paragraph.runs:
        return changes

    joined = "".join(run.text for run in paragraph.runs)
    normalized, changed = replace_text(joined)
    if changed and len(joined) == len(normalized):
        position = 0
        for run in paragraph.runs:
            length = len(run.text)
            run.text = normalized[position : position + length]
            position += length
        changes += changed
    return changes


def normalize_container(container):
    changes = 0
    for paragraph in container.paragraphs:
        changes += normalize_paragraph(paragraph)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                changes += normalize_container(cell)
    return changes


def normalize_docx(path):
    doc = Document(path)
    changes = normalize_container(doc)
    for section in doc.sections:
        changes += normalize_container(section.header)
        changes += normalize_container(section.footer)
    if changes:
        doc.save(path)
    return changes


def main():
    total = 0
    changed_docs = []
    for path in sorted(DOCS.rglob("*.docx")):
        if path.name.startswith("~$"):
            continue
        changes = normalize_docx(path)
        if changes:
            changed_docs.append((path, changes))
            total += changes
            print(f"UPDATED\t{path.relative_to(ROOT)}\t{changes}")
        else:
            print(f"OK\t{path.relative_to(ROOT)}\t0")
    print(f"TOTAL_DOCS_UPDATED={len(changed_docs)}")
    print(f"TOTAL_TEXT_BLOCKS_UPDATED={total}")


if __name__ == "__main__":
    main()
