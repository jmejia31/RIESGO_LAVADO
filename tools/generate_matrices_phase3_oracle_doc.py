from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_ROOT = ROOT / "docs"
MATRICES_DIR = next(DOCS_ROOT.rglob("3. Módulo Matrices de Riesgos"))
OUT_DIR = MATRICES_DIR / "Fase 3 - Modelo de datos y arquitectura Oracle"
OUT_FILE = OUT_DIR / "Fase_3_Modelo_Datos_Arquitectura_Oracle_Matrices_Riesgos_SGRLA_IHSS.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GOLD = "FFF2CC"
LIGHT_GREEN = "E2F0D9"
LIGHT_RED = "FCE4D6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_run_font(run, size=10, color=INK, bold=False, italic=False, name="Arial") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for level, size in [(1, 14), (2, 11.5), (3, 10.5)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = DARK_BLUE if level == 1 else BLUE
        style.paragraph_format.space_before = Pt(10 if level == 1 else 6)
        style.paragraph_format.space_after = Pt(4)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(9.5)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("IHSS - SGRLA/FT | Módulo Matrices de Riesgos | Fase 3")
    set_run_font(run, size=8, color=MUTED, bold=True)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Documento técnico para revisión DBA y aprobación funcional")
    set_run_font(run, size=8, color=MUTED)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[int], header_fill: str = LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], header_fill)
        run = table.rows[0].cells[idx].paragraphs[0].add_run(header)
        set_run_font(run, bold=True)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            if len(value) <= 10:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def add_checklist(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("DOCUMENTO TÉCNICO DE FASE 3")
    set_run_font(run, size=9, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Fase 3. Modelo de datos y arquitectura Oracle")
    set_run_font(run, size=22, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Módulo Matrices de Riesgos - Sistema de Gestión de Riesgos LA/FT IHSS")
    set_run_font(run, size=13, color=MUTED)

    meta = [
        ("Proyecto", "RIESGO_LAVADO - IHSS"),
        ("Módulo", "Matrices de Riesgos"),
        ("Fase", "Fase 3. Modelo de datos y arquitectura Oracle"),
        ("Versión", "1.9"),
        ("Fecha", date.today().strftime("%d/%m/%Y")),
        ("Estado", "Fase 3 ejecutada, validada y documentada con clasificación de evidencias"),
        ("Responsable", "Javier Mejía"),
        ("Fuente", "Plan de fases, análisis final maestro, Fase 1 aprobada y Fase 2 aprobada"),
        ("Ubicación", "docs/3. Módulo Matrices de Riesgos/Fase 3 - Modelo de datos y arquitectura Oracle"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [2450, 6910])
    for row, (label, value) in zip(table.rows, meta):
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_title_block(doc)

    doc.add_heading("0. Control de cambios", level=1)
    add_table(
        doc,
        ["Versión", "Fecha", "Descripción del cambio", "Responsable", "Estado"],
        [
            ("1.0", date.today().strftime("%d/%m/%Y"), "Creación de Fase 3 con modelo lógico Oracle, estructura RL_MR_*, scripts aprobables y guía DBA de ejecución segura.", "Javier Mejía", "Revisión técnica"),
            ("1.1", date.today().strftime("%d/%m/%Y"), "Ajustes técnicos: parametrización inicial, resultado vigente, motivo de recálculo, validación de permisos, validaciones de ponderación y exclusión de fuentes ajenas al proyecto.", "Javier Mejía", "Revisión técnica ajustada"),
            ("1.2", date.today().strftime("%d/%m/%Y"), "Alineación final del script 03 con Fase 2 aprobada: cinco niveles de riesgo, mitigación 0%, 10%, 25%, 40% y 55%, modelo base aprobado y variables iniciales con 100% interno por factor.", "Javier Mejía", "Versión final para aprobación funcional"),
            ("1.3", date.today().strftime("%d/%m/%Y"), "Revisión DBA estática sin ejecución: validación de nomenclatura, llaves, comentarios por tabla, comentarios por columna y restricción única lógica para escalas.", "Javier Mejía", "Revisión DBA estática completada"),
            ("1.4", date.today().strftime("%d/%m/%Y"), "Corrección funcional DBA: incorporación del estado OBSERVADA en la restricción física de estados de RL_MR_MATRICES, alineado con Fase 2 aprobada.", "Javier Mejía", "Observación corregida"),
            ("1.5", date.today().strftime("%d/%m/%Y"), "Aprobación formal de Fase 3 para cierre técnico del modelo de datos y arquitectura Oracle del módulo Matrices de Riesgos.", "Javier Mejía", "Aprobada"),
            ("1.6", date.today().strftime("%d/%m/%Y"), "Ejecución DBA controlada de scripts 01, 02 y 03 en el esquema indicado; corrección de compatibilidad Oracle 11g en script 02 y generación de evidencia técnica.", "Javier Mejía", "Ejecutada y validada"),
            ("1.7", date.today().strftime("%d/%m/%Y"), "Registro de cierre documental DBA transversal: esquema RIESGO_LAVADO validado con comentarios completos en tablas y columnas, sin comentarios faltantes ni codificación dañada.", "Javier Mejía", "Complementada"),
            ("1.8", date.today().strftime("%d/%m/%Y"), "Alineación DBA con Fase 4: incorporación del estado EN_EVALUACION en CK_RL_MR_MAT_ESTADO mediante script incremental controlado 05.", "Javier Mejía", "Alineada y validada"),
            ("1.9", date.today().strftime("%d/%m/%Y"), "Limpieza documental: clasificación de scripts 04 y 05 como validaciones/post-correctivos e identificación de logs fallidos como incidentes técnicos superados.", "Javier Mejía", "Documentada"),
        ],
        [900, 1350, 4650, 1500, 960],
    )

    doc.add_heading("1. Propósito de la fase", level=1)
    doc.add_paragraph(
        "La Fase 3 diseña la persistencia Oracle del módulo Matrices de Riesgos con trazabilidad, historial, versionamiento, auditoría y protección de matrices cerradas. "
        "Esta versión documenta el diseño aprobado y la ejecución DBA controlada realizada en el esquema indicado por el responsable funcional."
    )

    doc.add_heading("2. Dependencias aprobadas", level=1)
    add_table(
        doc,
        ["Dependencia", "Estado", "Uso en Fase 3"],
        [
            ("Fase 1", "Aprobada", "Gobierno documental, responsables, reglas heredadas y estructura de carpeta."),
            ("Fase 2", "Aprobada y cerrada", "Metodología LA/FT, pesos institucionales, estados, cálculo y reglas de cierre."),
            ("Plan de fases", "Vigente", "Define entregables: diagrama lógico, scripts RL_MR_*, registro del módulo y guía DBA."),
            ("Análisis final maestro", "Vigente", "Confirma prefijo RL_MR_*, backend calcula, Oracle conserva y DNP no recibe escrituras sin contrato."),
            ("Base RIESGO_LAVADO", "Referencia", "Contiene RL_AUDITORIA, RL_MODULOS, RL_USUARIO_MODULOS y tablas actuales del sistema."),
            ("Base DNP_IHSS", "Fuente externa controlada", "Aporta datos de patronos, proveedores, empleados, coincidencias y calificación futura por patrono."),
        ],
        [2200, 1800, 5360],
    )

    doc.add_heading("3. Decisiones técnicas rectoras", level=1)
    add_table(
        doc,
        ["Decisión", "Definición de Fase 3"],
        [
            ("Esquema destino", "La persistencia nueva del módulo se diseña en RIESGO_LAVADO con prefijo físico RL_MR_*."),
            ("Pesos institucionales", "Proveedores 50%, Clientes/Patronos 25% y Empleados 25%; no se modifican desde base de datos sin aprobación funcional."),
            ("Variables internas", "Cada factor institucional tendrá variables propias cuyo peso interno debe totalizar 100% por factor."),
            ("Variables iniciales", "El script 03 carga una base inicial revisable de variables por factor. Cualquier ajuste posterior debe versionarse sin alterar matrices cerradas."),
            ("Cálculo", "Backend calcula; Oracle guarda entradas, resultados y snapshots; Angular solo captura y muestra."),
            ("Matriz cerrada", "Una matriz cerrada no se recalcula ni se altera retroactivamente; conserva snapshot metodológico y de cálculo."),
            ("Auditoría", "RL_AUDITORIA registra eventos transversales; RL_MR_HISTORIAL conserva trazabilidad funcional específica del módulo."),
            ("DNP", "Queda como fuente y como integración futura obligatoria para calificación por patrono, sin escritura directa hasta contrato técnico aprobado."),
            ("Scripts", "Los scripts de esta fase son aprobables para revisión DBA; no forman parte del flujo oficial hasta aprobación."),
            ("Fuente de proveedores", "La fuente se definirá entre datos propios del módulo y fuentes autorizadas de DNP/IHSS cuando aplique; no se incorporarán fuentes ajenas al proyecto."),
        ],
        [2500, 6860],
        header_fill=LIGHT_GOLD,
    )

    doc.add_heading("4. Diagrama lógico de datos", level=1)
    diagram = doc.add_paragraph()
    run = diagram.add_run(
        "RL_MR_MODELOS\n"
        "  -> RL_MR_FACTORES\n"
        "      -> RL_MR_VARIABLES\n"
        "          -> RL_MR_CRITERIOS\n"
        "  -> RL_MR_ESCALAS\n"
        "  -> RL_MR_MATRICES\n"
        "      -> RL_MR_DETALLE\n"
        "      -> RL_MR_CONTROLES\n"
        "      -> RL_MR_RESULTADOS\n"
        "          -> RL_MR_PLANES_ACCION\n"
        "      -> RL_MR_EVIDENCIAS\n"
        "      -> RL_MR_HISTORIAL\n"
        "      -> RL_MR_INTEGRACION_DNP\n"
        "\n"
        "Integraciones transversales: RL_AUDITORIA, RL_MODULOS, RL_USUARIO_MODULOS, RL_USUARIOS.\n"
        "Fuentes externas controladas: DNP_IHSS.V_DATOS_EMPRESA, DNP_IHSS.V_ESTADO_DATOS, DNP_IHSS.V_EMPLEADOS_IHSS_PLANILLAS y DNP_IHSS.V_REPORTE_COINCIDENCIA."
    )
    set_run_font(run, size=8.5, name="Consolas")

    doc.add_heading("5. Inventario de tablas RL_MR_*", level=1)
    add_table(
        doc,
        ["Tabla", "Responsabilidad", "Regla crítica"],
        [
            ("RL_MR_MODELOS", "Versiona la metodología aprobable/aprobada.", "No modificar modelos usados por matrices cerradas; crear nueva versión."),
            ("RL_MR_FACTORES", "Guarda factores institucionales del modelo.", "Pesos fijos esperados: 50%, 25%, 25%."),
            ("RL_MR_VARIABLES", "Define variables internas por factor.", "Los pesos internos deben totalizar 100% por factor."),
            ("RL_MR_ESCALAS", "Define rangos y niveles.", "Cinco niveles para variable, inherente y residual; controles con mitigación 0%, 10%, 25%, 40% y 55%."),
            ("RL_MR_CRITERIOS", "Define criterios de calificación por variable.", "Criterios asociados a variables vigentes."),
            ("RL_MR_MATRICES", "Encabezado de cada evaluación.", "Estado cerrado protege la matriz contra edición retroactiva."),
            ("RL_MR_DETALLE", "Valores y puntajes por variable.", "Guarda peso snapshot usado al calcular."),
            ("RL_MR_CONTROLES", "Controles mitigantes.", "Solidez, efectividad y evidencia deben quedar trazables."),
            ("RL_MR_RESULTADOS", "Resultado inherente, mitigación y residual.", "Guarda versión de cálculo, resultado vigente, motivo de recálculo y snapshot del backend."),
            ("RL_MR_PLANES_ACCION", "Planes derivados del residual alto o crítico.", "Responsable, fechas, medio de prueba y estado obligatorios."),
            ("RL_MR_EVIDENCIAS", "Metadatos de evidencia documental.", "Eliminación lógica, motivo obligatorio y ruta protegida."),
            ("RL_MR_HISTORIAL", "Trazabilidad funcional específica.", "Complementa RL_AUDITORIA sin reemplazarla."),
            ("RL_MR_INTEGRACION_DNP", "Bandeja local para futura integración.", "No escribe en DNP sin contrato técnico y autorización institucional."),
        ],
        [2250, 3800, 3310],
    )

    doc.add_heading("6. Estados, restricciones e historial", level=1)
    add_table(
        doc,
        ["Objeto", "Estados propuestos", "Control"],
        [
            ("Modelo", "BORRADOR, EN_REVISION, APROBADO, CERRADO, INACTIVO", "Solo un modelo aprobado debe operar como vigente en backend."),
            ("Matriz", "BORRADOR, EN_EVALUACION, CALCULADA, EN_REVISION, OBSERVADA, APROBADA, CERRADA, INACTIVA", "EN_EVALUACION permite captura controlada antes del cálculo; OBSERVADA permite devolución con observaciones; CERRADA bloquea edición de metodología y resultados."),
            ("Control", "ACTIVO, OBSERVADO, INACTIVO", "La efectividad no reduce riesgo si no cumple evidencia o responsable."),
            ("Plan de acción", "PENDIENTE, EN_PROCESO, CERRADO, VENCIDO, INACTIVO", "Residual alto o crítico exige plan activo o justificación aprobada."),
            ("Integración DNP", "PENDIENTE, ENVIADO, ERROR, ANULADO", "Solo se habilita técnicamente cuando exista contrato de integración."),
        ],
        [2100, 3600, 3660],
    )

    doc.add_heading("7. Motivos, inactivación y recálculo", level=1)
    add_table(
        doc,
        ["Control", "Aplicación en el modelo"],
        [
            ("Motivo de inactivación", "Modelos, factores, variables, escalas, criterios, controles, evidencias y matrices conservan motivo directo o historial funcional obligatorio."),
            ("Motivo de recálculo", "RL_MR_RESULTADOS conserva MRR_MOTIVO_RECALCULO y referencia al resultado anterior cuando aplique."),
            ("Resultado vigente", "RL_MR_RESULTADOS conserva MRR_ES_VIGENTE y un índice único funcional para impedir más de un resultado vigente por matriz/factor/tipo."),
            ("Historial funcional", "RL_MR_HISTORIAL registra cambio de estado, motivo, datos anteriores, datos nuevos, usuario, IP y fecha."),
            ("Auditoría transversal", "RL_AUDITORIA seguirá registrando los eventos críticos desde backend."),
        ],
        [2600, 6760],
        header_fill=LIGHT_GOLD,
    )

    doc.add_heading("8. Compatibilidad con auditoría y módulos existentes", level=1)
    add_checklist(
        doc,
        [
            "Registrar el módulo en RL_MODULOS con ruta /matrices-riesgos y MOD_ID 10 sujeto a validación DBA.",
            "Asignar permisos iniciales solamente a usuarios administradores existentes y luego administrar accesos por RL_USUARIO_MODULOS.",
            "Validar antes de insertar permisos que RL_USUARIO_MODULOS mantenga la estructura real USM_USR_ID y USM_MOD_ID sin columnas obligatorias adicionales.",
            "Mantener RL_AUDITORIA como bitácora transversal para creación, edición, cálculo, aprobación, cierre, exportación e impresión.",
            "Usar RL_MR_HISTORIAL para conservar detalle funcional del cambio, motivo, estado anterior, estado nuevo y snapshot cuando aplique.",
            "No escribir en tablas DNP_IHSS desde esta fase; DNP queda como fuente autorizada y receptor futuro condicionado.",
        ],
    )

    doc.add_heading("9. Scripts entregados en Fase 3", level=1)
    add_table(
        doc,
        ["Script", "Clasificación", "Descripción", "Ejecución"],
        [
            ("01_F3_create_rl_mr_estructura.sql", "Ejecutado", "Crea secuencias, tablas, restricciones, índices y comentarios RL_MR_*.", "Ejecutado y validado."),
            ("02_F3_register_modulo_matrices_riesgos.sql", "Ejecutado", "Registra /matrices-riesgos en RL_MODULOS, valida RL_USUARIO_MODULOS y asigna usuarios 1 y 2 si existen. Se corrigió compatibilidad Oracle 11g usando nombre interno corto para el procedimiento de asignación.", "Ejecutado y validado."),
            ("03_F3_seed_metodologia_matrices_riesgos.sql", "Ejecutado", "Carga modelo base aprobado, factores institucionales 50/25/25, variables iniciales con 100% interno por factor, cinco niveles de riesgo y controles 0/10/25/40/55.", "Ejecutado y validado."),
            ("04_F3_fix_encoding_textos_oracle.sql", "Post-correctivo documental", "Corrige textos y comentarios con codificación dañada cuando aplique; no altera datos funcionales, cálculos ni estructura principal del módulo.", "Ejecutado como soporte documental y validado."),
            ("05_F3_align_estado_en_evaluacion_matrices.sql", "Incremental de alineación", "Alinea CK_RL_MR_MAT_ESTADO con Fase 4 incorporando EN_EVALUACION; valida estados existentes antes de recrear la restricción.", "Ejecutado y validado."),
        ],
        [2450, 1350, 3800, 1760],
    )

    doc.add_heading("9.1 Clasificación de evidencias e incidentes controlados", level=2)
    add_table(
        doc,
        ["Tipo de evidencia", "Tratamiento documental", "Criterio de uso"],
        [
            ("Logs finales correctos", "Se conservan como evidencia positiva de ejecución o validación.", "Pueden usarse para sustentar cierre técnico DBA."),
            ("Logs fallidos intermedios", "Se conservan como incidente técnico controlado y superado.", "No deben presentarse de forma aislada como evidencia final satisfactoria."),
            ("Script 04", "Se clasifica como post-correctivo documental de codificación y comentarios.", "Soporta limpieza de metadatos, no cambia la metodología funcional."),
            ("Script 05", "Se clasifica como incremental de alineación de estados.", "Respalda físicamente EN_EVALUACION para Fase 4."),
        ],
        [2200, 3900, 3260],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("10. Guía DBA de ejecución segura", level=1)
    add_checklist(
        doc,
        [
            "No ejecutar los scripts en producción mientras la Fase 3 no esté aprobada funcionalmente.",
            "Antes de ejecutar en pruebas, respaldar RL_MODULOS, RL_USUARIO_MODULOS y cualquier tabla RL_MR_* existente.",
            "Validar que el esquema activo sea RIESGO_LAVADO y que no se esté conectado a DNP_IHSS.",
            "Ejecutar primero 01_F3_create_rl_mr_estructura.sql, luego 02_F3_register_modulo_matrices_riesgos.sql y finalmente 03_F3_seed_metodologia_matrices_riesgos.sql.",
            "Confirmar que ningún script contiene DROP, TRUNCATE ni DELETE.",
            "Guardar salida completa de SQLPlus con fecha, usuario, ambiente y resultado.",
            "Después de ejecutar, validar conteo de tablas RL_MR_*, secuencias, índices, módulo 10 y permisos iniciales.",
            "No agregar estos scripts a database/00_EJECUCION_ACTUALIZACIONES_SEGURAS.sql hasta aprobación final.",
        ],
    )

    doc.add_heading("11. Validaciones SQL sugeridas", level=1)
    add_table(
        doc,
        ["Validación", "Consulta sugerida"],
        [
            ("Tablas RL_MR_*", "SELECT TABLE_NAME FROM USER_TABLES WHERE TABLE_NAME LIKE 'RL_MR_%' ORDER BY TABLE_NAME;"),
            ("Secuencias RL_MR_*", "SELECT SEQUENCE_NAME FROM USER_SEQUENCES WHERE SEQUENCE_NAME LIKE 'SEQ_RL_MR_%' ORDER BY SEQUENCE_NAME;"),
            ("Módulo registrado", "SELECT MOD_ID, MOD_NOMBRE, MOD_RUTA FROM RL_MODULOS WHERE MOD_RUTA = '/matrices-riesgos';"),
            ("Permisos iniciales", "SELECT * FROM RL_USUARIO_MODULOS WHERE USM_MOD_ID = 10;"),
            ("Estructura de permisos", "SELECT COLUMN_NAME, NULLABLE FROM USER_TAB_COLUMNS WHERE TABLE_NAME = 'RL_USUARIO_MODULOS' ORDER BY COLUMN_ID;"),
            ("Pesos institucionales", "SELECT SUM(MRF_PESO_INSTITUCIONAL) TOTAL FROM RL_MR_FACTORES WHERE MRF_ESTADO_REGISTRO = 1;"),
            ("Pesos internos por factor", "SELECT MRF.MRF_NOMBRE, SUM(MRV.MRV_PESO_INTERNO) TOTAL FROM RL_MR_FACTORES MRF JOIN RL_MR_VARIABLES MRV ON MRV.MRV_FACTOR_ID = MRF.MRF_ID WHERE MRV.MRV_ESTADO_REGISTRO = 1 GROUP BY MRF.MRF_NOMBRE;"),
            ("Escalas inherente/residual", "SELECT MRE_TIPO, MRE_VALOR_MIN, MRE_VALOR_MAX, MRE_NIVEL FROM RL_MR_ESCALAS WHERE MRE_TIPO IN ('INHERENTE','RESIDUAL') ORDER BY MRE_TIPO, MRE_ORDEN;"),
            ("Escala de controles", "SELECT MRE_VALOR_MIN, MRE_VALOR_MAX, MRE_NIVEL FROM RL_MR_ESCALAS WHERE MRE_TIPO = 'CONTROL' ORDER BY MRE_ORDEN;"),
            ("Resultado vigente", "SELECT MRR_MATRIZ_ID, NVL(MRR_FACTOR_ID, -1) FACTOR, MRR_TIPO_RESULTADO, COUNT(*) TOTAL FROM RL_MR_RESULTADOS WHERE MRR_ES_VIGENTE = 1 GROUP BY MRR_MATRIZ_ID, NVL(MRR_FACTOR_ID, -1), MRR_TIPO_RESULTADO HAVING COUNT(*) > 1;"),
            ("Restricción de auditoría", "Confirmar que las acciones del módulo usarán RL_AUDITORIA desde backend en Fase 5."),
        ],
        [2300, 7060],
    )

    doc.add_heading("12. Criterios de aceptación de Fase 3", level=1)
    add_table(
        doc,
        ["Criterio", "Estado en esta versión"],
        [
            ("Los scripts no eliminan información productiva.", "Cumplido: scripts sin DROP, TRUNCATE ni DELETE ejecutable."),
            ("Las tablas permiten conservar versiones históricas de metodología y resultados.", "Cumplido: modelos versionados, snapshots en matrices, detalle y resultados."),
            ("Toda matriz cerrada queda protegida contra edición metodológica retroactiva.", "Cumplido a nivel de diseño; se reforzará en backend con reglas de estado."),
            ("Existe diagrama lógico de datos.", "Cumplido dentro de este documento."),
            ("Existe script inicial de estructura RL_MR_*.", "Cumplido como script aprobable 01."),
            ("Existe script de registro del módulo en RL_MODULOS.", "Cumplido como script aprobable 02."),
            ("Existe script de parametrización inicial.", "Cumplido como script aprobable 03."),
            ("Existe comentario por tabla y por columna.", "Cumplido: 13 comentarios de tabla y 187 comentarios de columna en el script 01."),
            ("La nomenclatura DBA es consistente.", "Cumplido: cada tabla mantiene prefijo RL_MR_* y su llave primaria inicia con el prefijo funcional correspondiente."),
            ("Los estados físicos de matriz respetan el flujo aprobado.", "Cumplido: CK_RL_MR_MAT_ESTADO incluye EN_EVALUACION y OBSERVADA para respaldar el flujo funcional definido en Fase 4."),
            ("Existe validación de ponderaciones institucionales.", "Cumplido: script 03 valida total 100% y pesos 50/25/25."),
            ("Existe validación de pesos internos por factor.", "Cumplido: script 03 carga variables iniciales y valida 100% por cada factor institucional."),
            ("Las escalas coinciden con Fase 2 aprobada.", "Cumplido: inherente y residual usan cinco niveles; controles usan mitigación 0%, 10%, 25%, 40% y 55%."),
            ("El modelo base queda coherente con Fase 2 aprobada.", "Cumplido: el seed final registra el modelo como APROBADO, sujeto a aprobación DBA antes de ejecución."),
            ("Existe control de resultado vigente y recálculo.", "Cumplido: MRR_VERSION_CALCULO, MRR_ES_VIGENTE, MRR_MOTIVO_RECALCULO e índice único funcional."),
            ("Existe documento DBA de ejecución segura.", "Cumplido dentro de la sección 10 de este documento."),
        ],
        [4300, 5060],
        header_fill=LIGHT_GREEN,
    )

    doc.add_heading("13. Revisión DBA estática sin ejecución", level=1)
    doc.add_paragraph(
        "Se realizó una revisión DBA estática sobre los scripts de Fase 3 sin ejecutar instrucciones SQL contra ninguna base de datos. "
        "La revisión se concentró en nomenclatura, llaves primarias, llaves foráneas, restricciones, comentarios, idempotencia documental y orden seguro de ejecución."
    )
    add_table(
        doc,
        ["Control revisado", "Resultado"],
        [
            ("Nomenclatura de tablas", "Cumplido: todas las tablas nuevas usan el prefijo RL_MR_*."),
            ("Nomenclatura de llaves primarias", "Cumplido: cada tabla inicia con su campo ID correspondiente, por ejemplo MRM_ID, MRF_ID, MRV_ID, MRMAT_ID, MRCTRL_ID, MREV_ID y MRDNP_ID."),
            ("Prefijo de campos", "Cumplido: los campos conservan prefijo funcional por tabla; se usan prefijos extendidos cuando evitan ambigüedad."),
            ("Comentarios de tablas", "Cumplido: 13 tablas cuentan con COMMENT ON TABLE."),
            ("Comentarios de columnas", "Cumplido: 187 columnas cuentan con COMMENT ON COLUMN."),
            ("Escalas", "Mejorado: se agregó restricción única lógica UQ_RL_MR_ESC_RANGO para evitar duplicidad por modelo, tipo, rango y nivel."),
            ("Estado OBSERVADA", "Corregido: CK_RL_MR_MAT_ESTADO permite OBSERVADA para conservar el flujo funcional aprobado en Fase 2."),
            ("Estado EN_EVALUACION", "Corregido: CK_RL_MR_MAT_ESTADO permite EN_EVALUACION para conservar el flujo funcional definido en Fase 4."),
            ("Orden de ejecución", "Cumplido: 01 estructura, 02 registro del módulo y 03 parametrización inicial."),
            ("Seguridad de ejecución", "Cumplido a nivel documental: no se ejecutaron scripts y no hay DROP, TRUNCATE ni DELETE ejecutable."),
            ("Fuentes ajenas", "Cumplido: no se incorpora ni referencia una fuente ajena al proyecto."),
        ],
        [3100, 6260],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("14. Pendientes controlados antes de aprobar ejecución", level=1)
    add_table(
        doc,
        ["Pendiente", "Motivo", "Decisión requerida"],
        [
            ("Confirmar MOD_ID 10", "El manifiesto actual reserva la secuencia desde 10 para futuros módulos.", "DBA debe validar que 10 no esté ocupado."),
            ("Confirmar origen operativo de proveedores", "El módulo no utilizará fuentes ajenas al proyecto; la fuente debe quedar entre captura propia y datos autorizados de DNP/IHSS.", "Definir fuente oficial en Fase 4 o antes de carga inicial."),
            ("Confirmar tabla de evidencias final", "Matrices puede reutilizar reglas de evidencias, pero necesita metadatos propios o relación técnica.", "Validar si RL_MR_EVIDENCIAS queda propia o se integra con estructura común futura."),
            ("Confirmar contrato DNP", "La integración futura es obligatoria, pero depende de seguridad e interfaz.", "No habilitar escritura hacia DNP sin contrato técnico aprobado."),
            ("Confirmar cambios futuros de variables", "El script 03 deja variables iniciales para cerrar el modelo técnico.", "Cualquier cambio funcional posterior debe generar nueva versión metodológica y no alterar matrices cerradas."),
        ],
        [2450, 3550, 3360],
        header_fill=LIGHT_RED,
    )

    doc.add_heading("15. Aprobación formal de Fase 3", level=1)
    doc.add_paragraph(
        "Con base en la revisión técnica, la revisión DBA estática, la corrección de los estados OBSERVADA y EN_EVALUACION, y la validación final del paquete, "
        "la Fase 3 queda aprobada formalmente para cierre técnico y continuidad hacia la Fase 4 del módulo Matrices de Riesgos."
    )
    add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ("Estado de Fase 3", "Aprobada formalmente para cierre técnico."),
            ("Responsable de aprobación", "Javier Mejía."),
            ("Fecha de aprobación", date.today().strftime("%d/%m/%Y")),
            ("Alcance de la aprobación", "Documento técnico, modelo lógico Oracle, scripts aprobables 01, 02, 03, 04 y 05, validación DBA estática y evidencia de ejecución controlada."),
            ("Condición de ejecución", "La aprobación de Fase 3 no autoriza ejecución directa en producción; toda ejecución debe seguir protocolo institucional, respaldo, evidencia DBA y autorización correspondiente."),
            ("Siguiente fase autorizada", "Fase 4 del módulo Matrices de Riesgos."),
        ],
        [2600, 6760],
        header_fill=LIGHT_GREEN,
    )

    doc.add_heading("16. Cierre DBA con evidencia técnica", level=1)
    doc.add_paragraph(
        "Posterior a la aprobación funcional de Fase 3, se realizó ejecución DBA controlada de los scripts aprobables en el esquema RIESGO_LAVADO indicado por el responsable funcional. "
        "La ejecución dejó evidencia técnica en la carpeta de la fase y confirma que la estructura, parametrización inicial, módulo, permisos y validaciones principales quedaron correctamente aplicados."
    )
    add_table(
        doc,
        ["Control", "Resultado validado"],
        [
            ("Ambiente/esquema ejecutado", "Usuario y esquema RIESGO_LAVADO; base hpprod1; servidor desdb."),
            ("Orden de ejecución", "00 prevalidación, 01 estructura, 02 registro de módulo, 03 parametrización inicial, 04 validación/corrección documental cuando aplique y 05 alineación incremental de estado EN_EVALUACION."),
            ("Incidente controlado", "El primer intento del script 02 fue detenido por PLS-00114 debido a nombre interno de procedimiento superior a 30 caracteres. Se corrigió el identificador a asignar_mod_usuario y se reejecutó correctamente."),
            ("Logs fallidos intermedios", "Los intentos con error quedan archivados como incidentes técnicos superados y no como evidencia final positiva aislada."),
            ("Estructura creada", "13 tablas RL_MR_*, 13 secuencias SEQ_RL_MR_* y 33 índices relacionados."),
            ("Documentación física Oracle", "13 comentarios de tabla y 187 comentarios de columna aplicados."),
            ("Módulo y permisos", "Módulo /matrices-riesgos registrado con MOD_ID 10 y permisos iniciales asignados a usuarios existentes 1 y 2."),
            ("Metodología inicial", "Modelo base APROBADO, factores Proveedores 50%, Clientes/Patronos 25% y Empleados 25%."),
            ("Variables y escalas", "Siete variables iniciales por factor, peso interno 100% por factor y cinco escalas por tipo VARIABLE, INHERENTE, RESIDUAL y CONTROL."),
            ("Estados", "La restricción física de RL_MR_MATRICES incluye EN_EVALUACION y OBSERVADA."),
            ("Objetos inválidos", "No se detectaron objetos inválidos posteriores a la ejecución."),
            ("Evidencia", "Logs SQLPlus completos y documento de evidencia DBA en la carpeta Evidencia_DBA de la Fase 3."),
        ],
        [3000, 6360],
        header_fill=LIGHT_GREEN,
    )

    doc.add_heading("17. Cierre documental DBA transversal", level=1)
    doc.add_paragraph(
        "Como complemento posterior al cierre técnico de Fase 3, se revisó el esquema RIESGO_LAVADO completo a nivel de metadatos Oracle. "
        "La revisión confirmó la cobertura documental de tablas y columnas y corrigió comentarios faltantes o con codificación dañada, sin modificar estructura física, datos funcionales, llaves, índices ni restricciones."
    )
    add_table(
        doc,
        ["Control", "Resultado"],
        [
            ("Tablas revisadas", "29 tablas del esquema RIESGO_LAVADO."),
            ("Columnas revisadas", "314 columnas del esquema RIESGO_LAVADO."),
            ("Tablas sin comentario", "0 después del correctivo documental."),
            ("Columnas sin comentario", "0 después del correctivo documental."),
            ("Comentarios con codificación dañada", "0 después de la validación final."),
            ("Script aplicado", "database/18_add_missing_comments.sql."),
            ("Evidencia", "docs/1. Bases de Datos/Evidencia_DBA/Evidencia_DBA_Comentarios_Completos_RIESGO_LAVADO_SGRLA_IHSS.docx."),
        ],
        [3000, 6360],
        header_fill=LIGHT_BLUE,
    )

    props = doc.core_properties
    props.title = "Fase 3 - Modelo de datos y arquitectura Oracle - Matrices de Riesgos"
    props.subject = "Sistema de Gestión de Riesgos LA/FT IHSS"
    props.keywords = "IHSS, SGRLA, Matrices de Riesgos, Fase 3, Oracle, RL_MR"
    props.comments = "Documento ejecutado y validado con evidencia DBA para cierre técnico de Fase 3 del módulo Matrices de Riesgos, complementado con cierre documental DBA del esquema RIESGO_LAVADO."
    props.author = "Javier Mejía"

    doc.save(OUT_FILE)


if __name__ == "__main__":
    build_document()
    print(OUT_FILE)
