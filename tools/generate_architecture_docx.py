from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "Documentación chatGPT"
OUT_FILE = OUT_DIR / "Arquitectura_Proyecto_RIESGO_LAVADO.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9E2F3"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_inches: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(w * 1440) for w in widths_inches)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(w * 1440)))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_border_bottom(paragraph, color="A6A6A6", size="8", space="4") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_masthead(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    left = header.add_run("RIESGO_LAVADO")
    set_run_font(left, size=9, color=MUTED, bold=True)
    header.add_run("  |  ")
    right = header.add_run("Documentación de arquitectura")
    set_run_font(right, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Arquitectura técnica del proyecto")
    set_run_font(run, size=9, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("DOCUMENTACIÓN TÉCNICA")
    set_run_font(run, size=10, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    title = p.add_run("Arquitectura del Proyecto RIESGO_LAVADO")
    set_run_font(title, size=23, color=RGBColor(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    subtitle = p.add_run("Mapa técnico del repositorio, capas, módulos funcionales, seguridad, datos e integración.")
    set_run_font(subtitle, size=12, color=MUTED)

    metadata = [
        ("Proyecto", "RIESGO_LAVADO"),
        ("Fecha de generación", "26 de junio de 2026"),
        ("Alcance", "Backend ASP.NET Core, frontend Angular, scripts Oracle, documentación y herramientas"),
        ("Ubicación", str(ROOT)),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_run_font(r, bold=True)
        set_run_font(p.add_run(value))

    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, color="A6A6A6", size="8", space="4")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr_cells[idx], LIGHT_FILL)
        run = hdr_cells[idx].paragraphs[0].add_run(header)
        set_run_font(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ""
            for i, line in enumerate(value.split("\n")):
                if i:
                    cells[idx].paragraphs[0].add_run().add_break()
                run = cells[idx].paragraphs[0].add_run(line)
                set_run_font(run, size=10)
    set_table_width(table, widths)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=DARK_BLUE)
    p.add_run("\n")
    r = p.add_run(body)
    set_run_font(r, size=10.5)
    set_table_width(table, [6.5])
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    setup_styles(doc)
    add_masthead(doc)

    doc.add_heading("1. Resumen ejecutivo", level=1)
    add_callout(
        doc,
        "Arquitectura detectada",
        "El proyecto utiliza una arquitectura web cliente-servidor de tres capas: Angular como SPA, ASP.NET Core Web API como capa de servicios y Oracle como capa de datos. La API no usa ORM; centraliza el acceso a datos en repositorios con Oracle.ManagedDataAccess y SQL explícito.",
    )
    add_bullets(
        doc,
        [
            "Frontend: Angular 22 con componentes standalone, rutas protegidas, servicios HTTP, interceptor JWT y organización core/features/shared.",
            "Backend: ASP.NET Core net10.0 con Controllers, DTOs, Services, Repositories, Models, Infrastructure, Security y Middleware.",
            "Datos: Oracle con scripts SQL versionados en database y consumo de tablas propias RL_* más objetos del esquema DNP_IHSS.",
            "Seguridad: autenticación JWT, refresh tokens, BCrypt para usuarios locales, soporte de Active Directory y autorización por módulos.",
            "Operación: Serilog para logs, Swagger en desarrollo, CORS configurable, SMTP para notificaciones y manejo global de errores.",
        ]
    )

    doc.add_heading("2. Estructura del repositorio", level=1)
    add_table(
        doc,
        ["Carpeta", "Rol arquitectónico", "Contenido principal"],
        [
            ["backend", "Capa de API y lógica de negocio", "Solución ASP.NET Core RL.API y proyecto temporal TempSeqCheck."],
            ["frontend", "Capa de presentación", "Aplicación Angular rl-app con rutas, guards, services, features y layout compartido."],
            ["database", "Evolución de base de datos", "Scripts SQL numerados para crear tablas, registrar módulos, alterar estructuras y cargar semillas."],
            ["docs", "Documentación funcional/técnica", "Documentos Word/PDF existentes, bases SQL exportadas y documentación generada."],
            ["tools", "Automatización documental", "Scripts para generar documentos de análisis, diseño y esta documentación."],
            ["root", "Orquestación del proyecto", "RIESGO_LAVADO.sln, NuGet.Config, .gitignore y memoria técnica."],
        ],
        [1.25, 2.0, 3.25],
    )

    doc.add_heading("3. Vista de capas", level=1)
    add_table(
        doc,
        ["Capa", "Tecnología", "Responsabilidades"],
        [
            ["Presentación", "Angular 22, TypeScript, RxJS, Tailwind", "Interfaz administrativa, login, navegación por módulos, validaciones de UI, consumo de API y exportaciones PDF/Excel."],
            ["Aplicación/API", "ASP.NET Core net10.0", "Exposición REST, validación de autenticación/autorización, orquestación de servicios, manejo de archivos y respuestas JSON."],
            ["Dominio/servicios", "Services e interfaces C#", "Autenticación, Active Directory, catálogos, correo, configuración y reglas de operación."],
            ["Acceso a datos", "Repositories + Oracle.ManagedDataAccess", "Consultas y transacciones Oracle, auditoría, gestión de usuarios, listas, evidencias y coincidencias."],
            ["Persistencia", "Oracle", "Tablas RL_* propias del sistema y objetos DNP_IHSS para listas de cautela y coincidencias."],
        ],
        [1.25, 1.8, 3.45],
    )

    doc.add_heading("4. Backend: backend/RL.API", level=1)
    add_table(
        doc,
        ["Área", "Archivos/carpetas", "Descripción"],
        [
            ["Entrada", "Program.cs", "Configura Serilog, MVC/Newtonsoft, Swagger, JWT, CORS, DI, middleware, archivos estáticos y controllers."],
            ["Controllers", "Auth, Catalogos, Configuracion, Auditoria, Listas", "Endpoints REST agrupados por dominio funcional."],
            ["Services", "AuthService, ActiveDirectorioService, CatalogoService, EmailService", "Reglas de autenticación, AD, catálogos y envío de correos."],
            ["Repositories", "Usuario, Catalogo, Configuracion, Auditoria, Listas", "Acceso directo a Oracle con comandos parametrizados y transacciones."],
            ["Infrastructure", "OracleDbContext", "Fábrica de conexiones Oracle basada en ConnectionStrings:OracleDB."],
            ["Security", "ModuloAuthorizeAttribute", "Filtro backend que valida el claim modulos del JWT contra los módulos requeridos por endpoint."],
            ["Middleware", "ErrorHandlingMiddleware", "Captura excepciones no controladas y responde JSON estándar."],
        ],
        [1.35, 2.05, 3.1],
    )

    doc.add_heading("5. Frontend: frontend/rl-app", level=1)
    add_table(
        doc,
        ["Área", "Archivos/carpetas", "Descripción"],
        [
            ["Arranque", "main.ts, app.config.ts, app.routes.ts", "Aplicación Angular standalone con router y HttpClient configurado con interceptor."],
            ["core/services", "auth, catalogo, configuracion, listas, auditoria, active-directorio", "Cliente HTTP centralizado por dominio de API."],
            ["core/guards", "auth.guard, modulo.guard, role.guard", "Control de acceso por sesión, módulos y roles en rutas."],
            ["core/interceptors", "auth.interceptor.ts", "Adjunta Bearer token, renueva token en 401 y redirige a sin-acceso en 403."],
            ["features/auth", "login", "Pantalla de autenticación y arranque de sesión."],
            ["features/admin", "usuarios, configuracion, monitoreo-listas, bitacora, tipo-listas, cargar-listas, coincidencias", "Módulos funcionales administrativos y de listas de cautela."],
            ["shared", "main-layout, sin-acceso", "Layout autenticado y pantalla de acceso denegado."],
        ],
        [1.35, 2.05, 3.1],
    )

    doc.add_heading("6. Módulos funcionales", level=1)
    add_table(
        doc,
        ["ID lógico", "Ruta frontend", "Módulo", "Control de acceso"],
        [
            ["2", "/usuarios", "Usuarios del Sistema", "Rol ADMINISTRADOR + ModuloAuthorize(2)"],
            ["3", "/configuracion", "Configuración del Sistema", "Rol ADMINISTRADOR + ModuloAuthorize(3)"],
            ["4", "/monitoreo-listas", "Monitoreo de Listas", "ModuloAuthorize(4)"],
            ["5", "/bitacora", "Bitácora de Sistema", "ModuloAuthorize(5)"],
            ["6", "/tipo-listas", "Tipo Listas", "ModuloAuthorize(6)"],
            ["7", "/cargar-listas", "Cargar Listas", "ModuloAuthorize(7)"],
            ["8", "/coincidencias-patrono", "Coincidencias Patrono", "ModuloAuthorize(8)"],
            ["9", "/coincidencias-empleado", "Coincidencias Empleado", "ModuloAuthorize(9)"],
        ],
        [0.75, 1.5, 2.0, 2.25],
    )

    doc.add_heading("7. Endpoints principales", level=1)
    add_table(
        doc,
        ["Controller", "Base", "Responsabilidad"],
        [
            ["AuthController", "/api/auth", "Login, refresh, logout, perfil, usuarios, estado, cambio y recuperación de contraseña."],
            ["CatalogosController", "/api/catalogos", "Roles, dominios y módulos activos."],
            ["ConfiguracionController", "/api/configuracion", "Configuración general del sistema y slides del login."],
            ["AuditoriaController", "/api/auditoria", "Consulta paginada de bitácora y registro de exportaciones."],
            ["ListasController", "/api/listas", "Monitoreo, positivos, seguimientos, evidencias, tipos de listas, cargas y calificación de coincidencias."],
        ],
        [1.55, 1.45, 3.5],
    )

    doc.add_heading("8. Flujo de autenticación y autorización", level=1)
    add_numbered(
        doc,
        [
            "El usuario inicia sesión desde Angular contra POST /api/auth/login.",
            "AuthService busca el usuario por correo o usuario de dominio y valida contraseña local con BCrypt o credenciales contra Active Directory.",
            "La API genera JWT con claims de identidad, rol, uid, dominio, modulos y debe_cambiar_pass; además persiste refresh token en RL_REFRESH_TOKENS.",
            "Angular guarda access_token, refresh_token y token_expira en localStorage y el interceptor adjunta Authorization: Bearer.",
            "Las rutas Angular validan sesión con authGuard y módulo con moduloGuard; el backend vuelve a validar permisos con ModuloAuthorizeAttribute.",
            "Si una petición recibe 401, el interceptor intenta renovar el token; si recibe 403, redirige a /sin-acceso manteniendo la sesión activa.",
        ]
    )

    doc.add_heading("9. Base de datos", level=1)
    add_table(
        doc,
        ["Grupo", "Objetos detectados", "Uso"],
        [
            ["Seguridad", "RL_USUARIOS, RL_ROLES, RL_DOMINIO, RL_REFRESH_TOKENS, RL_PASSWORD_RESET_TOKENS", "Usuarios locales/dominio, roles, dominios AD, sesiones y recuperación de acceso."],
            ["Autorización modular", "RL_MODULOS, RL_USUARIO_MODULOS", "Catálogo de módulos y relación usuario-módulo usada por frontend y backend."],
            ["Configuración", "RL_CONFIG_SISTEMA, RL_LOGIN_SLIDES", "Parámetros visuales, legales, timeout, intentos y slides del login."],
            ["Auditoría", "RL_AUDITORIA", "Registro de acciones, exportaciones, login/logout y cambios relevantes."],
            ["Listas y evidencias", "RL_LISTA_POSITIVOS, RL_DETALLE_LISTA, RL_DETALLE_EVIDENCIA, RL_TIPOS_DOCUMENTO", "Positivos manuales, seguimientos, motivos, evidencia documental y trazabilidad."],
            ["Coincidencias", "RL_CALIF_COINCIDENCIAS, DNP_IHSS.REPORTE_COINCIDENCIAS, vistas DNP_IHSS", "Calificación, resumen y consulta de coincidencias de patronos, naturales y empleados."],
            ["Listas cautela", "DNP_IHSS.TIPO_LISTAS_CAUTELA, DNP_IHSS.LISTA_CAUTELA", "Mantenimiento y carga de archivos CSV/XML/Excel para listas OFAC, ONU, Engel y PEPS."],
        ],
        [1.45, 2.65, 2.4],
    )

    doc.add_heading("10. Configuración y dependencias", level=1)
    add_table(
        doc,
        ["Componente", "Archivo", "Detalle"],
        [
            ["API", "backend/RL.API/appsettings.example.json", "OracleDB, Jwt, Cors, ActiveDirectory, Smtp, Evidencias y AllowedHosts."],
            ["API", "backend/RL.API/RL.API.csproj", "net10.0, Oracle.ManagedDataAccess.Core, JwtBearer, BCrypt, MailKit, Serilog, Swagger y ExcelDataReader."],
            ["Frontend", "frontend/rl-app/package.json", "Angular 22, RxJS, jwt-decode, SweetAlert2, jsPDF, xlsx, Tailwind, Vitest y TypeScript 6."],
            ["Frontend", "src/environments/environment.ts", "apiUrl http://localhost:5043/api y hubUrl http://localhost:5043/hubs."],
            ["Solución", "RIESGO_LAVADO.sln", "Orquesta la solución .NET del backend."],
        ],
        [1.35, 2.35, 2.8],
    )

    doc.add_heading("11. Flujos funcionales clave", level=1)
    add_table(
        doc,
        ["Flujo", "Frontend", "Backend/Datos"],
        [
            ["Administración de usuarios", "features/admin/usuarios + AuthService/CatalogoService", "AuthController + UsuarioRepository + RL_USUARIOS/RL_USUARIO_MODULOS/RL_ROLES/RL_DOMINIO."],
            ["Configuración visual/login", "features/admin/configuracion + ConfiguracionService", "ConfiguracionController + ConfiguracionRepository + RL_CONFIG_SISTEMA/RL_LOGIN_SLIDES."],
            ["Monitoreo de listas", "features/admin/monitoreo-listas + ListasService", "ListasController + ListasRepository + DNP_IHSS.REPORTE_COINCIDENCIAS/vistas + RL_LISTA_POSITIVOS."],
            ["Seguimientos y evidencias", "ListasService y pantallas de coincidencias", "Endpoints positivos/seguimientos/evidencias + RL_DETALLE_LISTA/RL_DETALLE_EVIDENCIA."],
            ["Carga de listas cautela", "features/admin/cargar-listas", "Procesamiento de CSV/XML/Excel con ExcelDataReader/XML y reemplazo transaccional en DNP_IHSS.LISTA_CAUTELA."],
            ["Bitácora", "features/admin/bitacora + AuditoriaService", "AuditoriaController/Repository + RL_AUDITORIA."],
        ],
        [1.6, 2.1, 2.8],
    )

    doc.add_heading("12. Seguridad transversal", level=1)
    add_bullets(
        doc,
        [
            "JWT valida issuer, audience, firma, expiración y clock skew cero desde Program.cs.",
            "El claim modulos es el contrato de autorización fina entre base de datos, backend y frontend.",
            "Las contraseñas locales se almacenan con BCrypt; los usuarios de dominio se validan contra Active Directory.",
            "Refresh tokens se almacenan y revocan en Oracle para renovación de sesión y logout.",
            "Los endpoints sensibles combinan [Authorize], roles y ModuloAuthorize según el caso.",
            "ErrorHandlingMiddleware normaliza errores no controlados, aunque expone exception.Message en detalle.",
        ]
    )

    doc.add_heading("13. Observaciones a validar", level=1)
    add_table(
        doc,
        ["Tema", "Observación", "Sugerencia"],
        [
            ["SignalR/hubs", "Angular define hubUrl, pero Program.cs no mapea hubs ni registra SignalR.", "Confirmar si es deuda técnica futura o retirar configuración no usada."],
            ["Proyecto temporal", "backend/TempSeqCheck aparece junto al backend principal.", "Validar si debe permanecer en la solución/documentación o moverse a tools/tests."],
            ["Documentación generada", "Existen documentos en docs/Documentación chatGPT y herramientas en tools.", "Mantener esta documentación actualizada con el mismo generador cuando cambie la arquitectura."],
            ["Acceso a datos", "Se usan consultas SQL extensas en repositorios.", "Considerar separar consultas complejas o documentar contratos SQL críticos para facilitar mantenimiento."],
        ],
        [1.25, 2.75, 2.5],
    )

    doc.add_heading("14. Referencias internas", level=1)
    add_table(
        doc,
        ["Archivo", "Uso documental"],
        [
            ["backend/RL.API/Program.cs", "Pipeline, DI, JWT, CORS, Swagger, Serilog y middleware."],
            ["backend/RL.API/RL.API.csproj", "Framework y paquetes backend."],
            ["frontend/rl-app/package.json", "Framework y paquetes frontend."],
            ["frontend/rl-app/src/app/app.routes.ts", "Mapa de rutas y módulos."],
            ["frontend/rl-app/src/app/core/interceptors/auth.interceptor.ts", "Flujo de token, refresh y manejo de 401/403."],
            ["backend/RL.API/Security/ModuloAuthorizeAttribute.cs", "Autorización backend por módulos."],
            ["database/*.sql", "Evolución de esquema, módulos y tablas funcionales."],
        ],
        [2.9, 3.6],
    )

    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    main()
