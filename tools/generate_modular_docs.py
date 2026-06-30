from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\RIESGO_LAVADO")
OUT = ROOT / "docs" / "5. Documentacion Modular"
DATE = "30/06/2026"
VERSION = "1.0"
PRESET = "compact_reference_guide"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "SGRLA-IHSS | Documentacion modular"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer.paragraphs[0]
    footer.text = f"{title} | Version {VERSION} | {DATE}"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(100, 100, 100)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(subtitle)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(90, 90, 90)


def add_control(doc: Document, modulo: str, tipo: str, estado: str = "Word generado para revision y aprobacion") -> None:
    rows = [
        ("Modulo", modulo),
        ("Tipo de documento", tipo),
        ("Version", VERSION),
        ("Fecha", DATE),
        ("Preset", PRESET),
        ("Estado", estado),
        ("Fuente", "Repositorio local C:/RIESGO_LAVADO"),
    ]
    add_table(doc, ["Campo", "Valor"], rows, widths=[1.7, 4.8])


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[tuple], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def add_capture_section(doc: Document, rows: list[tuple[str, str]]) -> None:
    add_heading(doc, "Capturas actualizadas", 1)
    add_para(
        doc,
        "No se insertaron capturas nuevas porque el sistema no quedo disponible en navegador durante esta ejecucion. "
        "Para mantener evidencia confiable, las capturas deben tomarse desde el ambiente local o de pruebas con sesion valida."
    )
    add_table(
        doc,
        ["Pantalla", "Estado / accion requerida"],
        rows,
        widths=[2.4, 4.1],
    )


FRONT_ROUTES = [
    ("/usuarios", "Modulo 2", "Gestion de usuarios", "moduloGuard(2)"),
    ("/configuracion", "Modulo 3", "Configuracion", "moduloGuard(3)"),
    ("/monitoreo-listas", "Modulo 4", "Monitoreo de listas", "moduloGuard(4)"),
    ("/bitacora", "Modulo 5", "Auditoria / bitacora", "moduloGuard(5)"),
    ("/tipo-listas", "Modulo 6", "Tipos de listas", "moduloGuard(6)"),
    ("/cargar-listas", "Modulo 7", "Carga de listas", "moduloGuard(7)"),
    ("/coincidencias-patrono", "Modulo 8", "Coincidencias patrono", "moduloGuard(8)"),
    ("/coincidencias-empleado", "Modulo 9", "Coincidencias empleado", "moduloGuard(9)"),
]


MONITOREO_ENDPOINTS = [
    ("GET", "/api/listas/juridicas", "Consulta coincidencias juridicas y positivos manuales", "Modulo 4"),
    ("GET", "/api/listas/naturales", "Consulta coincidencias naturales y positivos manuales", "Modulo 4"),
    ("GET", "/api/listas/empleados", "Consulta coincidencias de empleados y positivos manuales", "Modulo 4"),
    ("GET", "/api/listas/naturales/{numeroIdentificacion}/detalle", "Detalle de coincidencia natural", "Modulo 4"),
    ("GET", "/api/listas/empleados/{numeroIdentificacion}/detalle", "Detalle de coincidencia empleado", "Modulo 4"),
    ("GET", "/api/listas/tipos-documento", "Catalogo de tipos de documento", "Modulo 4"),
    ("GET", "/api/listas/tipos-listas-cautela", "Catalogo de listas de cautela", "Modulo 4/6/7"),
    ("POST", "/api/listas/positivos", "Registra o actualiza motivo de ingreso positivo", "Modulo 4 + auditoria"),
    ("GET", "/api/listas/positivos/{noDocumento}", "Consulta positivo existente", "Modulo 4"),
    ("GET", "/api/listas/positivos/{noDocumento}/seguimientos", "Historial de seguimientos con filtros de fecha", "Modulo 4"),
    ("POST", "/api/listas/positivos/{noDocumento}/seguimientos", "Crea seguimiento y evidencias", "Modulo 4 + auditoria"),
    ("PUT", "/api/listas/seguimientos/{detalleId}", "Actualiza seguimiento y agrega evidencias", "Modulo 4 + auditoria"),
    ("DELETE", "/api/listas/seguimientos/{detalleId}", "Eliminacion logica de seguimiento con motivo", "Modulo 4 + auditoria"),
    ("GET", "/api/listas/evidencias/{evidenciaId}", "Descarga/visualizacion auditada de evidencia", "Modulo 4 + auditoria"),
    ("DELETE", "/api/listas/evidencias/{evidenciaId}", "Eliminacion logica de evidencia con motivo", "Modulo 4 + auditoria"),
    ("POST", "/api/listas/positivos/{noDocumento}/reporte-impreso", "Audita impresion/generacion de reporte", "Modulo 4 + auditoria"),
    ("POST", "/api/auditoria/exportacion", "Audita exportaciones Excel/PDF", "Modulo 4/5/7/8/9"),
]


SECURITY_ENDPOINTS = [
    ("POST", "/api/auth/login", "Inicio de sesion local o dominio", "AuditRequired Login"),
    ("POST", "/api/auth/refresh", "Renovacion de access token", "Sin auditoria directa"),
    ("POST", "/api/auth/logout", "Cierre de sesion", "AuditRequired Logout"),
    ("PUT", "/api/auth/password", "Cambio de contrasena", "AuditRequired"),
    ("GET", "/api/auth/perfil", "Perfil, modulos y bandera de cambio de clave", "Authorize"),
    ("GET", "/api/auth/usuarios", "Listado de usuarios", "Modulo 2"),
    ("POST", "/api/auth/usuarios", "Creacion de usuario", "Modulo 2 + AuditRequired"),
    ("PUT", "/api/auth/usuarios/{uid}", "Edicion de usuario", "Modulo 2 + AuditRequired"),
    ("PUT", "/api/auth/usuarios/{uid}/estado", "Cambio de estado", "Modulo 2 + AuditRequired"),
    ("GET", "/api/auth/validar-dominio", "Validacion de usuario AD", "Modulo 2"),
    ("POST", "/api/auth/recuperar-password", "Recuperacion con clave provisional", "AuditRequired"),
]


AUDIT_ENDPOINTS = [
    ("GET", "/api/auditoria", "Bitacora paginada con filtros", "Modulo 5"),
    ("POST", "/api/auditoria/exportacion", "Registro obligatorio de exportacion/reporte", "Modulos 4,5,7,8,9"),
]


DB_SCRIPTS = [
    ("00_EJECUCION_PRIMERA_VEZ.sql", "Primera instalacion", "Ejecuta base completa y validacion final"),
    ("00_EJECUCION_ACTUALIZACIONES_SEGURAS.sql", "Actualizacion", "Ejecuta alteraciones y modulos aprobados"),
    ("00_MANIFIESTO_SCRIPTS_APROBADOS.md", "Control", "Lista de scripts aprobados y alcance"),
    ("03_create_modules_table.sql", "Modulos", "Crea RL_MODULOS, RL_USUARIO_MODULOS y modulo 2"),
    ("05_register_monitoreo_listas.sql", "Modulo 4", "Registra Monitoreo de Listas"),
    ("08_register_bitacora.sql", "Modulo 5", "Registra Bitacora"),
    ("09_create_detalle_evidencia.sql", "Evidencias", "Crea RL_DETALLE_EVIDENCIA"),
    ("13_create_calificaciones_coincidencias.sql", "Coincidencias", "Crea RL_CALIF_COINCIDENCIAS"),
    ("15_update_detalle_evidencia_soft_delete.sql", "Evidencias", "Campos de eliminacion logica"),
    ("16_alter_lista_positivos_origen_registro.sql", "Positivos", "Origen de registro en RL_LISTA_POSITIVOS"),
    ("17_validate_module_ids.sql", "Validacion", "Verifica IDs de modulos alineados con Angular/backend"),
]


def monitoreo_cliente() -> None:
    title = "Monitoreo de Listas - Version Cliente"
    doc = Document()
    style_doc(doc, title)
    add_title(doc, title, "Guia funcional para usuarios del modulo de monitoreo, positivos, evidencias y reportes.")
    add_control(doc, "Monitoreo de Listas", "Cliente funcional")

    add_heading(doc, "Objetivo funcional", 1)
    add_para(doc, "Permitir al usuario de cumplimiento revisar coincidencias juridicas, naturales y empleados contra listas de cautela, registrar motivos de inclusion, dar seguimiento, adjuntar evidencias y generar reportes auditados.")

    add_heading(doc, "Flujo funcional", 1)
    add_numbers(doc, [
        "Ingresar al sistema con usuario autorizado al modulo Monitoreo de Listas.",
        "Seleccionar vista Juridicas, Naturales, Empleados o registros manuales segun aplique.",
        "Filtrar por estado, fecha o texto para ubicar coincidencias especificas.",
        "Revisar detalle de coincidencia cuando el registro proviene de DNP/listas.",
        "Registrar motivo de inclusion antes de habilitar seguimiento.",
        "Agregar seguimientos con comentario obligatorio y evidencias opcionales.",
        "Visualizar o descargar evidencias desde el historial, quedando auditada la accion.",
        "Generar ficha Excel o PDF con seguimientos y rango de fechas seleccionado.",
    ])

    add_heading(doc, "Reglas de negocio visibles para el usuario", 1)
    add_bullets(doc, [
        "No se permite dar seguimiento a un registro sin motivo inicial.",
        "Los seguimientos requieren comentario obligatorio.",
        "La eliminacion de evidencias o seguimientos exige motivo obligatorio.",
        "Las evidencias no se borran fisicamente al eliminarse; quedan inactivas logicamente.",
        "Los reportes y exportaciones se cancelan si no se registra auditoria.",
        "Los filtros de fecha no aceptan rango invertido.",
    ])

    add_heading(doc, "Rutas del frontend", 1)
    add_table(doc, ["Ruta", "Modulo", "Uso", "Proteccion"], [r for r in FRONT_ROUTES if r[0] == "/monitoreo-listas"], widths=[1.5, 1.1, 2.5, 1.4])

    add_heading(doc, "Auditorias generadas", 1)
    add_table(doc, ["Accion", "Momento"], [
        ("INSERT/UPDATE en RL_LISTA_POSITIVOS", "Registro o actualizacion de motivo de ingreso."),
        ("INSERT/UPDATE/DELETE en RL_DETALLE_LISTA", "Seguimientos creados, editados o inactivados."),
        ("INSERT/VER/DELETE en RL_DETALLE_EVIDENCIA", "Carga, visualizacion/descarga o eliminacion logica."),
        ("VER en RL_LISTA_POSITIVOS", "Impresion o generacion de reporte."),
        ("VER por exportacion", "Exportacion Excel/PDF desde el modulo."),
    ], widths=[2.2, 4.3])

    add_capture_section(doc, [
        ("Pantalla principal /monitoreo-listas", "Capturar con filtros visibles y tabla con registros."),
        ("Modal motivo de ingreso", "Capturar campos obligatorios, origen y motivo."),
        ("Historial de seguimientos", "Capturar filtros de fecha, evidencias y acciones."),
        ("Reporte PDF/Excel generado", "Capturar salida desde ambiente con datos de prueba."),
    ])

    add_heading(doc, "Pendientes conocidos", 1)
    add_bullets(doc, [
        "Tomar capturas finales desde ambiente disponible con datos de prueba.",
        "Validar con usuario funcional que los nombres de reportes coinciden con la nomenclatura institucional.",
    ])
    doc.save(OUT / "Monitoreo_Listas_Version_Cliente.docx")


def monitoreo_dev() -> None:
    title = "Monitoreo de Listas - Version Desarrollador"
    doc = Document()
    style_doc(doc, title)
    add_title(doc, title, "Documento tecnico de rutas, endpoints, tablas, servicios y auditoria del modulo.")
    add_control(doc, "Monitoreo de Listas", "Tecnico desarrollador")

    add_heading(doc, "Componentes frontend", 1)
    add_table(doc, ["Elemento", "Archivo"], [
        ("Componente principal", "frontend/rl-app/src/app/features/admin/monitoreo-listas/monitoreo-listas.component.ts"),
        ("Servicio HTTP", "frontend/rl-app/src/app/core/services/listas.service.ts"),
        ("Ruta protegida", "/monitoreo-listas con moduloGuard(4)"),
        ("Guard global", "authGuard + ModuloAuthorize equivalente en backend"),
    ], widths=[2.0, 4.5])

    add_heading(doc, "Endpoints backend", 1)
    add_table(doc, ["Metodo", "Endpoint", "Uso", "Seguridad"], MONITOREO_ENDPOINTS, widths=[0.7, 2.6, 2.5, 0.9])

    add_heading(doc, "Servicios backend involucrados", 1)
    add_table(doc, ["Servicio", "Responsabilidad"], [
        ("ListasService", "Entrada funcional del modulo; valida positivos, rangos de seguimiento, tipos y carga de cautela."),
        ("EvidenciasService", "Valida, guarda, descarga y elimina logicamente evidencias y seguimientos."),
        ("CoincidenciasService", "Se usa para modulos de coincidencias relacionados, manteniendo reglas por patrono/empleado."),
        ("AuditoriaRepository", "Registra auditorias con IP, usuario, tabla, accion, datos anteriores y nuevos."),
    ], widths=[2.0, 4.5])

    add_heading(doc, "Tablas y vistas involucradas", 1)
    add_table(doc, ["Objeto", "Uso"], [
        ("DNP_IHSS.V_REPORTE_COINCIDENCIA", "Fuente de coincidencias juridicas, naturales y empleados."),
        ("DNP_IHSS.LISTA_CAUTELA", "Registros cargados de listas de cautela."),
        ("DNP_IHSS.TIPO_LISTA_CAUTELA", "Catalogo de listas."),
        ("DNP_IHSS.TIPO_DOCUMENTO", "Catalogo de documentos."),
        ("RL_LISTA_POSITIVOS", "Motivos de inclusion y positivos manuales."),
        ("RL_DETALLE_LISTA", "Seguimientos y acciones posteriores."),
        ("RL_DETALLE_EVIDENCIA", "Metadatos de evidencias y eliminacion logica."),
        ("RL_AUDITORIA", "Bitacora transversal."),
    ], widths=[2.4, 4.1])

    add_heading(doc, "Reglas tecnicas", 1)
    add_bullets(doc, [
        "Los controladores solo exponen HTTP y delegan reglas a servicios.",
        "Toda accion critica debe tener ModuloAuthorize y auditoria cuando aplique.",
        "La descarga de evidencia pasa por endpoint protegido; no se exponen rutas publicas directas.",
        "Las evidencias se guardan con nombre fisico seguro tipo GUID y nombre original en metadatos.",
        "La eliminacion logica exige motivo y conserva archivo fisico.",
        "Las exportaciones registran auditoria antes de generar el archivo.",
    ])

    add_capture_section(doc, [
        ("Listado juridicas/naturales/empleados", "Capturar despues de levantar backend y frontend."),
        ("Modal seguimiento y evidencias", "Capturar validacion de comentario obligatorio."),
        ("Bitacora despues de descarga", "Capturar registro VER generado por evidencia."),
    ])

    add_heading(doc, "Pendientes conocidos", 1)
    add_bullets(doc, [
        "Completar capturas en ambiente funcional.",
        "Mantener documentacion versionada antes de iniciar Matrices de Riesgo.",
    ])
    doc.save(OUT / "Monitoreo_Listas_Version_Desarrollador.docx")


def seguridad_tecnica() -> None:
    title = "Seguridad y Usuarios - Version Tecnica"
    doc = Document()
    style_doc(doc, title)
    add_title(doc, title, "Documento tecnico de autenticacion, roles, modulos, guards y usuarios.")
    add_control(doc, "Seguridad y Usuarios", "Tecnico desarrollador")

    add_heading(doc, "Alcance", 1)
    add_bullets(doc, [
        "Login/logout con JWT y refresh token.",
        "Usuarios locales y usuarios de dominio.",
        "Cambio de contrasena y clave provisional.",
        "Roles, modulos asignados y menu filtrado por permisos.",
        "Autorizacion backend por ModuloAuthorize y guards Angular.",
    ])

    add_heading(doc, "Rutas frontend", 1)
    add_table(doc, ["Ruta", "Modulo", "Uso", "Proteccion"], [
        ("/login", "Publico", "Inicio de sesion y cambio de clave obligatoria", "Sin modulo"),
        ("/usuarios", "Modulo 2", "Gestion de usuarios", "authGuard + moduloGuard(2)"),
        ("/sin-acceso", "Protegida", "Pantalla de acceso denegado", "authGuard"),
        ("/home", "Protegida", "Redireccion al primer modulo autorizado", "authGuard"),
    ], widths=[1.5, 1.1, 2.6, 1.3])

    add_heading(doc, "Endpoints backend", 1)
    add_table(doc, ["Metodo", "Endpoint", "Uso", "Seguridad/Auditoria"], SECURITY_ENDPOINTS, widths=[0.7, 2.5, 2.4, 0.9])

    add_heading(doc, "Tablas involucradas", 1)
    add_table(doc, ["Tabla", "Uso"], [
        ("RL_USUARIOS", "Cuentas, estado, UID, dominio, cambio de clave y metadatos."),
        ("RL_ROLES", "Roles de autorizacion."),
        ("RL_DOMINIO", "Catalogo de dominios institucionales."),
        ("RL_MODULOS", "Catalogo de modulos y rutas."),
        ("RL_USUARIO_MODULOS", "Relacion usuario-modulo."),
        ("RL_REFRESH_TOKENS", "Persistencia y revocacion de refresh tokens."),
        ("RL_PASSWORD_RESET_TOKENS", "Recuperacion de contrasena y claves provisionales."),
        ("RL_AUDITORIA", "Registro de acciones criticas."),
    ], widths=[2.1, 4.4])

    add_heading(doc, "Reglas de negocio y seguridad", 1)
    add_bullets(doc, [
        "El refresh token no debe renovar sesion de usuario inactivo.",
        "El usuario que requiere cambio de contrasena no puede navegar a modulos internos.",
        "Los modulos del JWT gobiernan menu, guards y rutas del backend.",
        "Usuarios de dominio se validan contra Active Directory antes de guardar.",
        "Usuarios locales pueden cambiar o recuperar contrasena mediante flujo controlado.",
        "Cambio de estado de usuario queda auditado.",
    ])

    add_heading(doc, "Auditorias generadas", 1)
    add_bullets(doc, [
        "LOGIN y LOGOUT.",
        "Creacion, edicion y cambio de estado de usuarios.",
        "Cambio o solicitud de recuperacion de contrasena.",
        "La IP se resuelve desde el contexto HTTP y encabezados configurados.",
    ])

    add_capture_section(doc, [
        ("Login", "Capturar pantalla con validaciones visibles."),
        ("Usuarios", "Capturar listado y modal de creacion/edicion."),
        ("Sin acceso", "Capturar redireccion por falta de modulo."),
    ])
    doc.save(OUT / "Seguridad_Usuarios_Version_Tecnica.docx")


def auditoria_tecnica() -> None:
    title = "Auditoria - Version Tecnica"
    doc = Document()
    style_doc(doc, title)
    add_title(doc, title, "Documento tecnico de bitacora, reglas obligatorias y eventos auditables.")
    add_control(doc, "Auditoria / Bitacora", "Tecnico desarrollador")

    add_heading(doc, "Objetivo", 1)
    add_para(doc, "Centralizar la trazabilidad de acciones criticas del sistema, incluyendo cambios de datos, visualizaciones sensibles, descargas, exportaciones, reportes y eventos de autenticacion.")

    add_heading(doc, "Ruta frontend", 1)
    add_table(doc, ["Ruta", "Modulo", "Uso", "Proteccion"], [("/bitacora", "Modulo 5", "Consulta paginada de auditoria", "authGuard + moduloGuard(5)")], widths=[1.5, 1.1, 2.6, 1.3])

    add_heading(doc, "Endpoints", 1)
    add_table(doc, ["Metodo", "Endpoint", "Uso", "Seguridad"], AUDIT_ENDPOINTS, widths=[0.8, 2.5, 2.4, 0.8])

    add_heading(doc, "Estructura de datos", 1)
    add_table(doc, ["Campo logico", "Descripcion"], [
        ("Tabla", "Objeto funcional afectado."),
        ("RegistroId", "Identificador del registro auditado."),
        ("Accion", "INSERT, UPDATE, DELETE, LOGIN, LOGOUT o VER segun aplique."),
        ("DatosAnt/DatosNvo", "JSON con estado anterior y/o nuevo."),
        ("Usuario/IP", "Usuario autenticado e IP de la computadora que ejecuta la accion."),
        ("Modulo", "Modulo funcional que origina la auditoria."),
    ], widths=[2.0, 4.5])

    add_heading(doc, "Eventos obligatorios", 1)
    add_bullets(doc, [
        "Login/logout.",
        "Creacion, edicion y cambio de estado de usuarios.",
        "Carga de listas.",
        "Exportaciones Excel/PDF y generacion de reportes.",
        "Visualizacion o descarga de evidencias.",
        "Eliminacion logica de evidencias y seguimientos.",
        "Calificacion de coincidencias.",
        "Cambios de configuracion.",
    ])

    add_heading(doc, "Regla de aprobacion tecnica", 1)
    add_para(doc, "Ningun endpoint critico debe considerarse cerrado si no registra auditoria o si no existe justificacion documentada de por que no aplica.")

    add_capture_section(doc, [
        ("Bitacora general", "Capturar filtros, IP, usuario, accion y detalle."),
        ("Detalle auditado", "Capturar modal con valores anteriores/nuevos."),
        ("Documentos eliminados", "Capturar filtro rapido de evidencias eliminadas."),
    ])
    doc.save(OUT / "Auditoria_Version_Tecnica.docx")


def evidencias_flujo() -> None:
    title = "Evidencias y Seguimientos - Flujo Funcional"
    doc = Document()
    style_doc(doc, title)
    add_title(doc, title, "Flujo funcional y reglas de seguridad documental.")
    add_control(doc, "Evidencias y Seguimientos", "Flujo funcional")

    add_heading(doc, "Flujo funcional", 1)
    add_numbers(doc, [
        "El usuario registra motivo de ingreso positivo.",
        "El sistema habilita seguimiento para el documento.",
        "El usuario ingresa comentario obligatorio y adjunta evidencias si aplica.",
        "El backend valida extension, MIME, tamano y firma real del archivo si esta activa.",
        "El sistema guarda nombre original y nombre fisico seguro con GUID.",
        "La descarga o visualizacion pasa por endpoint protegido y genera auditoria.",
        "La eliminacion exige motivo; el registro queda inactivo y el archivo fisico se conserva.",
    ])

    add_heading(doc, "Endpoints principales", 1)
    add_table(doc, ["Metodo", "Endpoint", "Uso", "Auditoria"], [
        ("GET", "/api/listas/evidencias/politica", "Politica de tamano, MIME y extensiones", "No directa"),
        ("POST", "/api/listas/positivos/{noDocumento}/seguimientos", "Crear seguimiento y evidencias", "INSERT"),
        ("PUT", "/api/listas/seguimientos/{detalleId}", "Editar seguimiento y adjuntar evidencias", "UPDATE/INSERT"),
        ("GET", "/api/listas/evidencias/{evidenciaId}", "Descargar/visualizar evidencia", "VER"),
        ("DELETE", "/api/listas/evidencias/{evidenciaId}", "Eliminar evidencia logicamente", "DELETE"),
        ("DELETE", "/api/listas/seguimientos/{detalleId}", "Eliminar seguimiento logicamente", "DELETE"),
    ], widths=[0.7, 2.7, 2.4, 0.7])

    add_heading(doc, "Reglas de seguridad documental", 1)
    add_bullets(doc, [
        "Extensiones permitidas segun configuracion: PDF, DOC, DOCX, XLS, XLSX, JPG, JPEG, PNG.",
        "Tamano maximo configurable, por defecto 10 MB.",
        "MIME permitido y firma real del archivo verificable.",
        "Ruta final configurable en Evidencias:StoragePath.",
        "Ruta legacy solo para compatibilidad: Evidencias:LegacyStoragePath.",
        "No exponer carpetas de evidencias como ruta publica directa.",
        "Mantener archivo fisico al eliminar logicamente.",
    ])

    add_heading(doc, "Tablas involucradas", 1)
    add_table(doc, ["Tabla", "Uso"], [
        ("RL_LISTA_POSITIVOS", "Registro principal de positivo y motivo inicial."),
        ("RL_DETALLE_LISTA", "Seguimientos por positivo."),
        ("RL_DETALLE_EVIDENCIA", "Metadatos, estado y motivo de eliminacion."),
        ("RL_AUDITORIA", "Trazabilidad de carga, descarga, edicion y eliminacion."),
    ], widths=[2.2, 4.3])

    add_capture_section(doc, [
        ("Modal seguimiento", "Capturar comentario obligatorio y selector de archivos."),
        ("Historial con evidencias", "Capturar evidencias existentes y botones de descarga/eliminacion."),
        ("Motivo de eliminacion", "Capturar SweetAlert de motivo obligatorio."),
    ])
    doc.save(OUT / "Evidencias_Seguimientos_Flujo_Funcional.docx")


def base_datos_guia() -> None:
    title = "Base de Datos - Guia de Ejecucion"
    doc = Document()
    style_doc(doc, title)
    add_title(doc, title, "Guia operativa para primera instalacion, actualizacion segura y control de scripts.")
    add_control(doc, "Base de Datos", "Guia de ejecucion segura")

    add_heading(doc, "Regla principal", 1)
    add_para(doc, "No mezclar scripts experimentales con scripts aprobados. Los scripts aprobados se ejecutan desde la raiz database; los utilitarios y experimentales se mantienen separados.")

    add_heading(doc, "Orden recomendado", 1)
    add_numbers(doc, [
        "Respaldar tablas afectadas o esquema completo segun el tipo de ejecucion.",
        "Validar conexion, esquema y permisos Oracle.",
        "Ejecutar 00_EJECUCION_PRIMERA_VEZ.sql solo en instalacion nueva.",
        "Ejecutar 00_EJECUCION_ACTUALIZACIONES_SEGURAS.sql en bases existentes.",
        "Ejecutar 17_validate_module_ids.sql al finalizar.",
        "Conservar spool/log de salida como evidencia de ejecucion.",
    ])

    add_heading(doc, "Scripts aprobados clave", 1)
    add_table(doc, ["Script", "Tipo", "Proposito"], DB_SCRIPTS, widths=[2.6, 1.4, 2.5])

    add_heading(doc, "Tablas criticas", 1)
    add_table(doc, ["Tabla", "Criterio de cuidado"], [
        ("RL_USUARIOS", "Respaldar antes de cambios de seguridad, dominio o contrasenas."),
        ("RL_MODULOS", "IDs deben coincidir con frontend y ModuloAuthorize."),
        ("RL_USUARIO_MODULOS", "Controla accesos iniciales."),
        ("RL_AUDITORIA", "No truncar en ambientes de control."),
        ("RL_LISTA_POSITIVOS", "Contiene positivos manuales y motivos."),
        ("RL_DETALLE_LISTA", "Seguimientos operativos."),
        ("RL_DETALLE_EVIDENCIA", "Metadatos y eliminacion logica de evidencias."),
        ("RL_CALIF_COINCIDENCIAS", "Calificaciones de patronos y empleados."),
    ], widths=[2.2, 4.3])

    add_heading(doc, "Modulos registrados", 1)
    add_table(doc, ["ID", "Ruta", "Modulo"], [
        ("2", "/usuarios", "Usuarios"),
        ("3", "/configuracion", "Configuracion"),
        ("4", "/monitoreo-listas", "Monitoreo de Listas"),
        ("5", "/bitacora", "Bitacora"),
        ("6", "/tipo-listas", "Tipo Listas"),
        ("7", "/cargar-listas", "Cargar Listas"),
        ("8", "/coincidencias-patrono", "Coincidencias Patrono"),
        ("9", "/coincidencias-empleado", "Coincidencias Empleado"),
    ], widths=[0.8, 2.4, 3.3])

    add_heading(doc, "Pendientes conocidos", 1)
    add_bullets(doc, [
        "Antes de Matrices de Riesgo, crear scripts propios del modulo en archivo nuevo y aprobado.",
        "No usar carpeta _experimental_no_ejecutar en ambientes productivos.",
        "Actualizar manifiesto cuando se apruebe un nuevo script.",
    ])
    doc.save(OUT / "Base_Datos_Guia_Ejecucion.docx")


def matriz_control() -> None:
    title = "Matriz de Control de Documentacion Modular"
    doc = Document()
    style_doc(doc, title)
    add_title(doc, title, "Inventario de documentos Word requeridos antes de avanzar al siguiente modulo.")
    add_control(doc, "Documentacion", "Matriz de control")

    add_heading(doc, "Estado de cierre documental minimo", 1)
    add_table(doc, ["Documento", "Modulo", "Estado"], [
        ("Monitoreo_Listas_Version_Cliente.docx", "Monitoreo de Listas", "Generado para revision"),
        ("Monitoreo_Listas_Version_Desarrollador.docx", "Monitoreo de Listas", "Generado para revision"),
        ("Seguridad_Usuarios_Version_Tecnica.docx", "Seguridad y usuarios", "Generado para revision"),
        ("Auditoria_Version_Tecnica.docx", "Auditoria", "Generado para revision"),
        ("Evidencias_Seguimientos_Flujo_Funcional.docx", "Evidencias y seguimientos", "Generado para revision"),
        ("Base_Datos_Guia_Ejecucion.docx", "Base de datos", "Generado para revision"),
    ], widths=[3.2, 1.8, 1.5])

    add_heading(doc, "Regla de versionamiento", 1)
    add_bullets(doc, [
        "Toda documentacion debe revisarse primero en Word.",
        "Cuando el usuario apruebe, se puede considerar version aprobada.",
        "No subir cambios documentales a repositorio remoto sin aprobacion.",
        "Las capturas deben tomarse desde ambiente real o de pruebas, nunca simularse.",
    ])
    doc.save(OUT / "Matriz_Control_Documentacion_Modular.docx")


def validate_docx(path: Path) -> None:
    Document(path)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    builders = [
        monitoreo_cliente,
        monitoreo_dev,
        seguridad_tecnica,
        auditoria_tecnica,
        evidencias_flujo,
        base_datos_guia,
        matriz_control,
    ]
    for builder in builders:
        builder()
    for path in sorted(OUT.glob("*.docx")):
        validate_docx(path)
        print(path)


if __name__ == "__main__":
    main()
