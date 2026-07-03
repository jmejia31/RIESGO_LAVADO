from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "1. Bases de Datos" / "Evidencia_DBA"
OUT_FILE = OUT_DIR / "Evidencia_DBA_Comentarios_Completos_RIESGO_LAVADO_SGRLA_IHSS.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "E2F0D9"
LIGHT_GOLD = "FFF2CC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_run_font(run, size=10.5, color=INK, bold=False, italic=False, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, before, after, color in [
        (1, 16, 16, 8, BLUE),
        (2, 13, 12, 6, BLUE),
        (3, 12, 8, 4, DARK_BLUE),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(doc: Document) -> None:
    header_p = doc.sections[0].header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("IHSS - SGRLA/FT | Evidencia DBA Base de Datos")
    set_run_font(run, size=8, color=MUTED, bold=True)

    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Comentarios completos del esquema RIESGO_LAVADO")
    set_run_font(run, size=8, color=MUTED)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[int], fill: str = LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], fill)
        run = table.rows[0].cells[idx].paragraphs[0].add_run(header)
        set_run_font(run, bold=True)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    return table


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("EVIDENCIA DBA")
    set_run_font(run, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Comentarios Completos del Esquema RIESGO_LAVADO")
    set_run_font(run, size=22, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Sistema de Gestión de Riesgos LA/FT IHSS")
    set_run_font(run, size=13, color=MUTED)

    add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ("Proyecto", "RIESGO_LAVADO - IHSS"),
            ("Documento", "Evidencia DBA de comentarios completos en base de datos"),
            ("Versión", "1.0"),
            ("Fecha", date.today().strftime("%d/%m/%Y")),
            ("Responsable", "Javier Mejía"),
            ("Estado", "Ejecutado y validado"),
            ("Ubicación", "docs/1. Bases de Datos/Evidencia_DBA"),
        ],
        [2200, 7160],
        fill=LIGHT_BLUE,
    )

    doc.add_heading("1. Propósito", level=1)
    doc.add_paragraph(
        "Dejar evidencia formal de que el esquema RIESGO_LAVADO fue revisado a nivel de metadatos Oracle y que todas sus tablas y columnas cuentan con comentarios descriptivos. "
        "La corrección se ejecutó sin cambios estructurales y sin alterar información funcional del sistema."
    )

    doc.add_heading("2. Alcance Técnico", level=1)
    add_table(
        doc,
        ["Elemento", "Resultado"],
        [
            ("Esquema revisado", "RIESGO_LAVADO"),
            ("Tablas revisadas", "29"),
            ("Columnas revisadas", "314"),
            ("Tablas sin comentario detectadas inicialmente", "2"),
            ("Columnas sin comentario detectadas inicialmente", "52"),
            ("Comentarios existentes con codificación dañada", "2"),
            ("Script aplicado", "database/18_add_missing_comments.sql"),
        ],
        [3600, 5760],
        fill=LIGHT_GOLD,
    )

    doc.add_heading("3. Correcciones Aplicadas", level=1)
    add_table(
        doc,
        ["Tipo", "Detalle"],
        [
            ("Tablas", "Se agregaron comentarios a RL_TIPOS_DOCUMENTO y RL_TIPOS_POSITIVO."),
            ("Columnas", "Se agregaron comentarios a 52 columnas sin comentario."),
            ("Codificación", "Se corrigieron comentarios existentes en RL_CALIF_COINCIDENCIAS.CAL_FECHA y RL_CALIF_COINCIDENCIAS.CAL_USUARIO_ID."),
            ("Seguridad", "No se ejecutaron cambios estructurales, DML funcional ni renombrado de objetos."),
            ("Codificación de ejecución", "El script fue ejecutado con NLS_LANG=AMERICAN_AMERICA.WE8MSWIN1252 para conservar tildes correctamente en Oracle."),
        ],
        [2600, 6760],
    )

    doc.add_heading("4. Validación Final", level=1)
    add_table(
        doc,
        ["Validación", "Resultado"],
        [
            ("Tablas totales", "29"),
            ("Tablas sin comentario", "0"),
            ("Columnas totales", "314"),
            ("Columnas sin comentario", "0"),
            ("Comentarios de tabla con codificación dañada", "0"),
            ("Comentarios de columna con codificación dañada", "0"),
        ],
        [4300, 5060],
        fill=LIGHT_GREEN,
    )

    doc.add_heading("5. Evidencia Generada", level=1)
    add_table(
        doc,
        ["Archivo", "Descripción"],
        [
            ("18_add_missing_comments_final_20260703_142409.log", "Log de ejecución final de comentarios."),
            ("18_add_missing_comments_final_20260703_142409.sql", "Copia del script ejecutado en formato compatible con SQLPlus."),
            ("20_validate_comments_final_20260703_142547.log", "Log de validación final con cero comentarios faltantes y cero errores de codificación."),
            ("20_validate_comments_final_20260703_142547.sql", "Consulta de validación final."),
        ],
        [4500, 4860],
    )

    doc.add_heading("6. Conclusión", level=1)
    doc.add_paragraph(
        "El esquema RIESGO_LAVADO queda documentalmente completo a nivel de comentarios Oracle para tablas y columnas. "
        "La validación final confirma que no existen comentarios faltantes ni caracteres dañados en los comentarios revisados."
    )

    props = doc.core_properties
    props.title = "Evidencia DBA - Comentarios Completos RIESGO_LAVADO"
    props.subject = "Sistema de Gestión de Riesgos LA/FT IHSS"
    props.keywords = "IHSS, SGRLA, RIESGO_LAVADO, DBA, comentarios Oracle"
    props.comments = "Evidencia de cierre DBA para comentarios completos del esquema RIESGO_LAVADO."
    props.author = "Javier Mejía"
    doc.save(OUT_FILE)


if __name__ == "__main__":
    build_document()
    print(OUT_FILE)
