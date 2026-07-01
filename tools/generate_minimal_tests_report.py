import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "8. Pruebas Mínimas"
JSON_RESULT = OUT_DIR / "resultado_pruebas_minimas_readonly.json"
OUT_DOCX = OUT_DIR / "Informe_Pruebas_Minimas_Pre_Matrices_SGRLA_IHSS.docx"


BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
HEADER = "E8EEF5"
WARN = "FFF2CC"
OK = "E2F0D9"
PENDING = "FCE4D6"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": 80, "bottom": 80, "start": 120, "end": 120}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2
    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Informe de Pruebas Mínimas")
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = DARK
    p = doc.add_paragraph()
    r = p.add_run("Validación previa al módulo Matrices de Riesgo")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED


def add_meta(doc):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Campo"
    table.rows[0].cells[1].text = "Detalle"
    rows = [
        ("Proyecto", "Sistema de Gestión de Riesgos LAFT - IHSS"),
        ("Fecha", "01/07/2026"),
        ("Modo de ejecución", "No destructivo / solo lectura"),
        ("API evaluada", "http://localhost:5043/api"),
        ("Estado", "Pruebas de consulta ejecutadas; pruebas de escritura pendientes de ambiente seguro o aprobación explícita."),
        ("Fuente", "Repositorio local C:/RIESGO_LAVADO"),
    ]
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    for cell in table.rows[0].cells:
        shade(cell, HEADER)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    table_geometry(table, [2600, 6760])
    doc.add_paragraph()


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, WARN)
    cell.text = text
    table_geometry(table, [9360])
    doc.add_paragraph()


def add_test_matrix(doc, results):
    result_map = {r["name"]: r for r in results}
    rows = [
        ("Login", "No ejecutada", "Requiere credencial confirmada. No se intentó para evitar bloqueo de cuenta o escritura de auditoría fallida."),
        ("Acceso por módulo", "Ejecutada", "Validado con token JWT de prueba con módulos 2 al 9; /api/auth/perfil y /api/auth/usuarios respondieron 200."),
        ("Crear usuario", "Pendiente controlado", "Escribe en RL_USUARIOS y auditoría. Debe ejecutarse solo con usuario de prueba y limpieza aprobada."),
        ("Editar usuario", "Pendiente controlado", "Escribe en RL_USUARIOS y auditoría. Debe ejecutarse solo sobre usuario de prueba."),
        ("Cargar lista", "Pendiente controlado", "Escribe registros de lista y auditoría. Requiere archivo de prueba y tipo de lista aprobado."),
        ("Consultar monitoreo", "Ejecutada", "Jurídicas, naturales y empleados respondieron 200."),
        ("Registrar positivo manual", "Pendiente controlado", "Escribe en lista de positivos. Requiere documento de prueba identificable."),
        ("Agregar seguimiento", "Pendiente controlado", "Escribe seguimiento. Debe asociarse a positivo de prueba."),
        ("Subir evidencia", "Pendiente controlado", "Escribe metadata y archivo físico. Debe usarse evidencia de prueba."),
        ("Descargar evidencia", "Pendiente controlado", "Debe ejecutarse sobre evidencia de prueba para no visualizar datos reales."),
        ("Eliminar evidencia con motivo", "Pendiente controlado", "Hace eliminación lógica. Debe ejecutarse solo sobre evidencia de prueba."),
        ("Eliminar seguimiento con motivo", "Pendiente controlado", "Hace eliminación lógica. Debe ejecutarse solo sobre seguimiento de prueba."),
        ("Exportar reportes", "Pendiente controlado", "La exportación registra auditoría. Debe ejecutarse con datos de prueba o autorización explícita."),
        ("Revisar bitácora", "Ejecutada", "Endpoint /api/auditoria respondió 200 con paginación."),
        ("Calificar coincidencia", "Pendiente controlado", "Actualiza calificación y auditoría. Debe ejecutarse sobre coincidencia de prueba."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Prueba mínima", "Estado", "Resultado / criterio"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        shade(cell, HEADER)
    for name, status, detail in rows:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = status
        cells[2].text = detail
        if status == "Ejecutada":
            shade(cells[1], OK)
        elif status.startswith("Pendiente") or status == "No ejecutada":
            shade(cells[1], PENDING)
    table_geometry(table, [2300, 1900, 5160])
    doc.add_paragraph()

    doc.add_heading("Resultados de endpoints no destructivos", 2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for idx, header in enumerate(["Prueba", "Método", "Ruta", "Estado HTTP", "Tiempo"]):
        cell = table.rows[0].cells[idx]
        cell.text = header
        shade(cell, HEADER)
    for item in results:
        cells = table.add_row().cells
        cells[0].text = item["name"]
        cells[1].text = item["method"]
        cells[2].text = item["url"]
        cells[3].text = str(item["status"])
        cells[4].text = f'{item["elapsedMs"]} ms'
        shade(cells[3], OK if item["ok"] else PENDING)
    table_geometry(table, [2850, 950, 3050, 1250, 1260])
    doc.add_paragraph()


def add_findings(doc):
    doc.add_heading("Hallazgos", 1)
    bullets = [
        "No se modificó información existente durante la ejecución de pruebas no destructivas.",
        "La API local estaba activa y respondió correctamente en los endpoints de consulta evaluados.",
        "El build normal no pudo ejecutarse porque la DLL de la API está tomada por el proceso activo de dotnet.",
        "El intento de build con salida temporal evidenció duplicidad de atributos de ensamblado por archivos generados incluidos desde obj/temporales; debe revisarse antes de cierre formal de pruebas técnicas.",
        "Las pruebas de escritura quedan pendientes hasta confirmar ambiente de pruebas o autorizar una ejecución controlada con datos identificables y limpieza posterior.",
    ]
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = json.loads(JSON_RESULT.read_text(encoding="utf-8"))
    doc = Document()
    styles(doc)
    title(doc)
    add_meta(doc)
    add_callout(
        doc,
        "Regla aplicada: no se ejecutaron operaciones POST, PUT ni DELETE contra la base conectada para evitar alterar información existente. Las pruebas de escritura deben ejecutarse únicamente con datos de prueba y autorización del responsable.",
    )
    doc.add_heading("Resumen ejecutivo", 1)
    doc.add_paragraph(
        "Se ejecutó una validación no destructiva del backend local antes de avanzar al módulo Matrices de Riesgo. "
        "La revisión confirmó disponibilidad de API, autorización por módulo mediante token de prueba, consultas de monitoreo, bitácora, política de evidencias y resúmenes de coincidencias."
    )
    add_test_matrix(doc, data["results"])
    add_findings(doc)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Pruebas mínimas SGRLA-IHSS | 01/07/2026")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
