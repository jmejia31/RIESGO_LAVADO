from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "2. Módulo Monitoreo de Listas" / "Cierre_Final_Aprobado"
OUT = OUT_DIR / "Acta_Cierre_Formal_Modulo_Monitoreo_Listas_SGRLA_IHSS.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
HEADER = "E8EEF5"
OK = "E2F0D9"
WARN = "FFF2CC"
PENDING = "FCE4D6"
INFO = "F4F6F9"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Acta Interna de Cierre Formal")
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    r = p.add_run("Módulo Monitoreo de Listas - Sistema de Gestión de Riesgos LAFT IHSS")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED


def add_table(doc, headers, rows, widths, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        shade(cell, HEADER)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
        if status_col is not None:
            value = row[status_col].lower()
            if "terminado" in value or "aplicable" in value or "vigente" in value:
                shade(cells[status_col], OK)
            elif "pendiente" in value or "parcial" in value:
                shade(cells[status_col], PENDING)
            elif "aprobación" in value or "no tocar" in value:
                shade(cells[status_col], WARN)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_callout(doc, text, fill=INFO):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = text
    shade(cell, fill)
    set_table_geometry(table, [9360])
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_styles(doc)
    add_title(doc)

    add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ("Proyecto", "Sistema de Gestión de Riesgos LAFT - IHSS"),
            ("Módulo", "Monitoreo de Listas"),
            ("Fecha de cierre documental", "01/07/2026"),
            ("Estado", "Cierre formal documentado; pendiente únicamente de validación final del usuario responsable."),
            ("Carpeta de cierre", "docs/2. Módulo Monitoreo de Listas/Cierre_Final_Aprobado"),
            ("Fuente", "Repositorio local C:/RIESGO_LAVADO"),
        ],
        [2600, 6760],
    )

    add_callout(
        doc,
        "Regla de cierre: desde este punto, el módulo Monitoreo de Listas no debe modificarse funcional, documental ni estructuralmente sin aprobación expresa del usuario responsable.",
        WARN,
    )

    doc.add_heading("1. Qué quedó terminado", 1)
    add_table(
        doc,
        ["Área", "Estado", "Detalle"],
        [
            ("Backend del módulo", "Terminado", "Controladores reducidos a entrada/salida HTTP, servicios separados y validaciones críticas movidas a servicios."),
            ("Servicios", "Terminado", "ListasService, EvidenciasService y CoincidenciasService quedan como base técnica; MatricesRiesgoService queda disponible como patrón futuro."),
            ("Auditoría", "Terminado", "Login/logout, usuarios, estados, carga de listas, exportaciones, evidencias, seguimientos, configuración, reportes y calificaciones registran auditoría cuando aplica."),
            ("Evidencias", "Terminado", "Validación de extensión, MIME, tamaño, nombre físico seguro con GUID, eliminación lógica con motivo, descarga auditada y almacenamiento controlado."),
            ("Monitoreo funcional", "Terminado", "Jurídicas, naturales, empleados, positivos manuales, origen, motivo, seguimientos, evidencias, filtros, exportaciones, reportes, estados y acciones principales quedaron cerrados funcionalmente."),
            ("Coincidencias", "Terminado", "Patrono y empleado cuentan con búsqueda, detalle, paginación/resumen, exportación, calificación, auditoría y permisos por módulo."),
            ("Seguridad y accesos", "Terminado", "Roles, módulos asignados, ModuloAuthorize, guards Angular, menú filtrado, pantalla sin acceso, sesión, cambio de contraseña y usuarios locales/dominio quedan integrados."),
            ("Base de datos", "Terminado", "Scripts de instalación, actualización, manifiesto, guía de ejecución segura, validación de módulos y separación de experimentales quedaron ordenados."),
            ("Documentación", "Terminado", "Documentación full por módulo, política de repositorio, guía de base de datos y pruebas mínimas no destructivas quedan en Word."),
        ],
        [2100, 1600, 5660],
        status_col=1,
    )

    doc.add_heading("2. Qué quedó parcialmente terminado", 1)
    add_table(
        doc,
        ["Elemento", "Estado", "Motivo"],
        [
            ("Pruebas de escritura", "Parcial", "No se ejecutaron POST, PUT ni DELETE contra la base conectada para no alterar información existente."),
            ("Render visual con LibreOffice", "Parcial", "El entorno no tiene soffice disponible; se aplicó validación estructural con python-docx."),
            ("Aprobación formal del usuario", "Pendiente", "La aprobación final depende de revisión del usuario responsable en Word."),
            ("Matrices de Riesgo", "Pendiente", "No se inicia desarrollo funcional hasta cerrar aprobación del módulo actual."),
        ],
        [2700, 1700, 4960],
        status_col=1,
    )

    doc.add_heading("3. Qué queda pendiente", 1)
    add_bullets(
        doc,
        [
            "Ejecutar pruebas controladas de escritura en ambiente de pruebas o con autorización explícita: crear/editar usuario, cargar lista, positivo manual, seguimiento, evidencia, eliminación lógica, exportación y calificación.",
            "Confirmar visualmente los Word finales si se habilita LibreOffice o revisión manual en Microsoft Word.",
            "Recibir aprobación final del usuario responsable antes de subir cambios o declarar cierre aprobado en repositorio remoto.",
            "Mantener pendiente cualquier ajuste de Matrices hasta validar que no afecta reglas de Monitoreo de Listas.",
        ],
    )

    doc.add_heading("4. Qué requiere validación del usuario", 1)
    add_table(
        doc,
        ["Validación", "Requiere aprobación"],
        [
            ("Contenido final de la documentación Word", "Sí"),
            ("Ejecución de pruebas que escriben en base de datos", "Sí"),
            ("Carpeta Cierre_Final_Aprobado como paquete oficial del módulo", "Sí"),
            ("Subida de cambios al repositorio remoto", "Sí"),
            ("Inicio del módulo Matrices de Riesgo", "Sí"),
        ],
        [4200, 5160],
    )

    doc.add_heading("5. Qué no debe tocarse sin aprobación", 1)
    add_bullets(
        doc,
        [
            "Scripts aprobados de base de datos y manifiesto de ejecución segura.",
            "IDs de módulos, rutas Angular y reglas de ModuloAuthorize.",
            "Política de auditoría obligatoria para endpoints críticos.",
            "Reglas de evidencias: validación, almacenamiento, descarga auditada y eliminación lógica con motivo.",
            "Documentos finales dentro de Cierre_Final_Aprobado.",
            "Configuración de conexión, SMTP o rutas de almacenamiento sin revisión previa.",
            "Estados y reglas de calificación de coincidencias.",
        ],
    )

    doc.add_heading("6. Qué se puede reutilizar para Matrices de Riesgo", 1)
    add_table(
        doc,
        ["Componente reutilizable", "Aplicación en Matrices"],
        [
            ("Patrón controlador + servicio", "Mantener controladores como HTTP y mover reglas de negocio/cálculos a MatricesRiesgoService."),
            ("ModuloAuthorize y guards Angular", "Aplicar autorización backend y frontend al nuevo módulo desde el primer endpoint/pantalla."),
            ("AuditRequired y AuditoriaRepository", "Registrar cambios de configuración, cálculos, aprobaciones, visualizaciones sensibles y exportaciones."),
            ("Guía de base de datos", "Crear scripts idempotentes, manifiesto aprobado, respaldos y orden de ejecución antes de tocar datos reales."),
            ("Documentación full por módulo", "Crear versión cliente, versión desarrollador, flujo funcional, endpoints, tablas, reglas, auditorías y pendientes."),
            ("Política de repositorio", "Revisar en Word, no subir borradores, no subir duplicados ni temporales, y separar cliente/desarrollador."),
            ("Pruebas mínimas", "Repetir esquema no destructivo primero y luego pruebas de escritura con datos controlados."),
            ("Evidencias y seguimientos", "Reutilizar criterios de seguridad documental si Matrices adjunta soportes o aprobaciones."),
        ],
        [3100, 6260],
    )

    doc.add_heading("7. Documentos incluidos en el paquete de cierre", 1)
    add_table(
        doc,
        ["Documento", "Uso"],
        [
            ("Acta_Cierre_Formal_Modulo_Monitoreo_Listas_SGRLA_IHSS.docx", "Acta interna de cierre formal."),
            ("Documentacion_Modulo_Monitoreo_Listas_Terminado.docx", "Documento final funcional del módulo."),
            ("Cierre_Maestro_Final_Fase_1_Modulos_Base_Monitoreo_Listas.docx", "Cierre maestro de fase base."),
            ("Monitoreo_Listas_Version_Cliente_Final.docx", "Versión cliente de documentación modular."),
            ("Monitoreo_Listas_Version_Desarrollador_Final.docx", "Versión desarrollador de documentación modular."),
            ("Evidencias_Seguimientos_Version_Full_Final.docx", "Flujo funcional y técnico de evidencias y seguimientos."),
            ("Informe_Pruebas_Minimas_Pre_Matrices_SGRLA_IHSS.docx", "Evidencia de pruebas mínimas no destructivas previas a Matrices."),
        ],
        [4300, 5060],
    )

    doc.add_heading("8. Criterio de avance a Matrices de Riesgo", 1)
    add_callout(
        doc,
        "Se puede iniciar Matrices de Riesgo cuando el usuario responsable confirme este paquete de cierre, autorice los pendientes controlados que correspondan y mantenga congeladas las reglas críticas del módulo Monitoreo de Listas.",
        OK,
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Cierre formal Monitoreo de Listas | SGRLA-IHSS | 01/07/2026")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
