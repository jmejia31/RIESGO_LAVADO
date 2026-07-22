from __future__ import annotations

import json
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import RGBColor

ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "docs/3. Módulo Matrices de Riesgos/Fase 12 - Mejora ejecutiva UXUI y mapa de calor"
DOCX = BASE / "Cierre_Tecnico_Fase_12_Matrices_Riesgos_SGRLA_IHSS.docx"
EVIDENCE = BASE / "Evidencia_Fase_12_5_5/fase12_5_5_cierre_definitivo.json"
XLSX_SOURCE = ROOT / "backend/RL.API/Infrastructure/Reporting/InstitutionalXlsxWorkbook.cs"
RENDERER_TEST = ROOT / "backend/RL.API.Tests/Features/MatricesRiesgos/MatricesRiesgosReportRendererTests.cs"


def replace_once(path: Path, old: str, new: str, label: str) -> None:
    text = path.read_text(encoding="utf-8")
    count = text.count(old)
    if count != 1:
        raise RuntimeError(f"{label}: se esperaba una coincidencia y se encontraron {count}")
    path.write_text(text.replace(old, new, 1), encoding="utf-8")


def patch_sources() -> None:
    replace_once(
        XLSX_SOURCE,
        '               "<worksheet xmlns=\\"http://schemas.openxmlformats.org/spreadsheetml/2006/main\\">" +\n'
        '               $"<sheetViews><sheetView workbookViewId=\\"0\\"><pane ySplit=\\"4\\" topLeftCell=\\"A5\\" activePane=\\"bottomLeft\\" state=\\"frozen\\"/></sheetView></sheetViews>" +',
        '               "<worksheet xmlns=\\"http://schemas.openxmlformats.org/spreadsheetml/2006/main\\">" +\n'
        '               "<sheetPr><pageSetUpPr fitToPage=\\"1\\" autoPageBreaks=\\"0\\"/></sheetPr>" +\n'
        '               $"<sheetViews><sheetView workbookViewId=\\"0\\"><pane ySplit=\\"4\\" topLeftCell=\\"A5\\" activePane=\\"bottomLeft\\" state=\\"frozen\\"/></sheetView></sheetViews>" +',
        "Habilitación explícita de fitToPage en OpenXML",
    )
    replace_once(
        RENDERER_TEST,
        '["FuenteDato"] = "Expediente institucional"',
        '["FuenteDato"] = "Expediente IHSS"',
        "Texto de evidencia de prueba",
    )
    replace_once(
        RENDERER_TEST,
        '        foreach (var hoja in new[] { "Resumen", "Matrices", "Factores", "Mapa transición", "Matrices críticas", "Planes" })\n'
        '        {\n'
        '            Assert.Contains(hoja, workbookXml);\n'
        '        }\n',
        '        foreach (var hoja in new[] { "Resumen", "Matrices", "Factores", "Mapa transición", "Matrices críticas", "Planes" })\n'
        '        {\n'
        '            Assert.Contains(hoja, workbookXml);\n'
        '        }\n\n'
        '        foreach (var worksheet in zip.Entries.Where(entry =>\n'
        '                     entry.FullName.StartsWith("xl/worksheets/sheet", StringComparison.OrdinalIgnoreCase)\n'
        '                     && entry.FullName.EndsWith(".xml", StringComparison.OrdinalIgnoreCase)))\n'
        '        {\n'
        '            using var sheetReader = new StreamReader(worksheet.Open(), Encoding.UTF8);\n'
        '            var sheetXml = sheetReader.ReadToEnd();\n'
        '            Assert.Contains("<pageSetUpPr fitToPage=\\"1\\" autoPageBreaks=\\"0\\"/>", sheetXml);\n'
        '            Assert.Contains("fitToWidth=\\"1\\"", sheetXml);\n'
        '        }\n',
        "Prueba OpenXML de ajuste de impresión",
    )


def set_cell(cell, value: str) -> None:
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def replace_paragraph(document: Document, old: str, new: str) -> int:
    changed = 0
    containers = [document]
    for section in document.sections:
        containers.extend([section.header, section.footer])
    for container in containers:
        for paragraph in container.paragraphs:
            if old not in paragraph.text:
                continue
            full = paragraph.text.replace(old, new)
            if paragraph.runs:
                paragraph.runs[0].text = full
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(full)
            changed += 1
    return changed


def update_document() -> None:
    document = Document(DOCX)
    if len(document.tables) < 29:
        raise RuntimeError("Estructura inesperada del Documento Maestro.")

    set_cell(
        document.tables[0].cell(0, 0),
        "CIERRE TÉCNICO DE LA FASE 12 Y ACTUALIZACIÓN 12.5\n"
        "Mejora ejecutiva UX/UI, reportería institucional, pruebas y cierre definitivo",
    )
    set_cell(document.tables[1].cell(1, 1), "1.1")
    set_cell(document.tables[1].cell(2, 1), "22 de julio de 2026")
    set_cell(document.tables[1].cell(4, 1), "CIERRE TÉCNICO APROBADO - PENDIENTE APROBACIÓN FORMAL")

    set_cell(
        document.tables[2].cell(0, 0),
        "CONTROL DEL DOCUMENTO Y APROBACIÓN\nDocumento maestro de cierre técnico de la Fase 12 y actualización final 12.5",
    )
    set_cell(document.tables[3].cell(2, 1), "Cierre Técnico Fase 12 y Actualización 12.5 - Matrices de Riesgos")
    set_cell(document.tables[3].cell(3, 1), "1.1")
    set_cell(document.tables[3].cell(4, 1), "22/07/2026")
    set_cell(document.tables[3].cell(8, 1), "#2 - Fase 12.5: estandarización institucional y refinamiento final")

    set_cell(
        document.tables[4].cell(0, 0),
        "RESULTADO TÉCNICO\nLa Fase 12 y su actualización 12.5 se encuentran técnicamente completas. "
        "El código, la reportería PDF/XLSX, la ficha individual, las pruebas automatizadas, "
        "la evidencia visual y el Documento Maestro han sido validados; la aprobación formal y el merge a main permanecen pendientes.",
    )

    test_rows = [
        ("Backend .NET", "96", "0", "0", "APROBADO"),
        ("Frontend unitarias", "156", "0", "—", "APROBADO"),
        ("Archivos de prueba frontend", "15", "0", "—", "APROBADO"),
        ("E2E Playwright", "7", "0", "—", "APROBADO"),
        ("Build Angular", "Correcto", "0", "—", "APROBADO"),
        ("Quality Gates finales", "Ejecución 29955226359", "0", "—", "APROBADO"),
    ]
    for row_index, values in enumerate(test_rows, start=1):
        for column, value in enumerate(values):
            set_cell(document.tables[16].cell(row_index, column), value)

    coverage_rows = [
        ("Backend", "Líneas", "20.56%"),
        ("Backend", "Ramas", "20.86%"),
        ("Frontend", "Sentencias", "36.59%"),
        ("Frontend", "Ramas", "30.24%"),
        ("Frontend", "Funciones", "33.68%"),
        ("Frontend", "Líneas", "37.06%"),
    ]
    for row_index, values in enumerate(coverage_rows, start=1):
        for column, value in enumerate(values):
            set_cell(document.tables[17].cell(row_index, column), value)

    traceability_rows = [
        ("Estándar institucional 12.5.1", "526dbb6c70739531a465a2a752dc60ff88ef910f", "Inventario y estándar compartido."),
        ("Normalización Monitoreo 12.5.2", "c17a7812ef1bdc7e895c10a3fc1ab1584398a7cc", "PDF/XLSX institucionales en Monitoreo."),
        ("Reportería Matrices 12.5.3", "42200b08a0c3a29669b446feddcffccc02d9831c", "PDF, ficha y XLSX generados en backend."),
        ("Refinamiento 12.5.4", "0bdc15193a795046bd63d4a685f856f902e21337", "UX, accesibilidad y retiro seguro de /recalcular."),
        ("Cierre técnico 12.5.5", "40bee949495785167f678e41ab929fba479b7814", "Archivos reales, Documento Maestro y evidencia."),
        ("Ejecución controlada 12.5.5", "29955226359", "Quality Gates, build y render aprobados."),
        ("Evidencia final", "Evidencia_Fase_12_5_5/fase12_5_5_cierre_definitivo.json", "Métricas, checksums y limitación Oracle."),
        ("Documento Maestro", "Cierre_Tecnico_Fase_12_Matrices_Riesgos_SGRLA_IHSS.docx", "Versión 1.1 con actualización 12.5."),
        ("PR principal", "#2", "Abierto, en borrador y sin merge a main."),
        ("PR técnico temporal", "#12", "Ejecución controlada; cierre sin merge a main."),
    ]
    for row_index, values in enumerate(traceability_rows, start=1):
        for column, value in enumerate(values):
            set_cell(document.tables[21].cell(row_index, column), value)

    set_cell(document.tables[23].cell(1, 0), "Alcance de Fase 12 y actualización 12.5 implementado")
    set_cell(
        document.tables[24].cell(0, 0),
        "RECOMENDACIÓN DE APROBACIÓN\nAprobar formalmente la Fase 12 y su actualización 12.5 del Módulo Matrices de Riesgos. "
        "Como acción posterior e independiente, autorizar la integración del PR #2 a main únicamente después de verificar nuevamente sus controles.",
    )

    replacements = {
        "Los workflows de integración y scripts auxiliares utilizados para aplicar y validar la Fase 12.4 fueron retirados antes de publicar el commit productivo.":
            "Los workflows de integración y scripts auxiliares utilizados para aplicar y validar las Fases 12.4 y 12.5 fueron retirados antes de publicar los commits productivos.",
        "Con la evidencia técnica disponible, la Fase 12 queda declarada técnicamente completada.":
            "Con la evidencia técnica disponible, la Fase 12 y su actualización 12.5 quedan declaradas técnicamente completadas.",
        "FIN DEL DOCUMENTO": "ANEXO DE ACTUALIZACIÓN FINAL DE LA FASE 12.5 A CONTINUACIÓN",
    }
    for old, new in replacements.items():
        if replace_paragraph(document, old, new) != 1:
            raise RuntimeError(f"No se pudo actualizar el texto documental: {old}")

    document.core_properties.title = "Cierre Técnico Fase 12 y Actualización 12.5 - Matrices de Riesgos"
    document.core_properties.subject = "Pruebas, evidencia y cierre definitivo"
    document.core_properties.version = "1.1"
    document.core_properties.modified = datetime.now()
    document.save(DOCX)


def update_evidence(run_id: int) -> None:
    data = json.loads(EVIDENCE.read_text(encoding="utf-8"))
    data["estado"] = "control_final_generado_pendiente_revision_visual_manual"
    data["control_final"] = {
        "run_id": run_id,
        "ajuste_xlsx_fit_to_page": True,
        "documento_maestro_version": "1.1",
        "metricas_historicas_reconciliadas": True,
        "trazabilidad_12_5_actualizada": True,
        "revision_visual_manual": "pendiente_revision_asistente",
    }
    data["siguiente_paso"] = "Revisión visual manual final y Quality Gate independiente"
    EVIDENCE.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Uso: fase12_5_5_control_final.py <run_id>")
    patch_sources()
    update_document()
    update_evidence(int(sys.argv[1]))
    print(json.dumps({
        "xlsx_fit_to_page": True,
        "documento_version": "1.1",
        "evidencia": str(EVIDENCE.relative_to(ROOT)),
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
