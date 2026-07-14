from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_ROOT = ROOT / "docs"
MATRICES_DIR = next(DOCS_ROOT.rglob("3. Módulo Matrices de Riesgos"))
FASE_DIR = MATRICES_DIR / "Fase 3 - Modelo de datos y arquitectura Oracle"
EVIDENCE_DIR = FASE_DIR / "Evidencia_DBA"
OUT_FILE = EVIDENCE_DIR / "Evidencia_DBA_Cierre_Tecnico_Fase_3_Matrices_Riesgos_SGRLA_IHSS.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "E2F0D9"
LIGHT_GOLD = "FFF2CC"
LIGHT_RED = "FCE4D6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_run_font(run, size=10.5, color=INK, bold=False, italic=False, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, before, after, color in [
        (1, 16, 16, 8, BLUE),
        (2, 13, 12, 6, BLUE),
        (3, 12, 8, 4, DARK_BLUE),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(doc: Document) -> None:
    header_p = doc.sections[0].header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("IHSS - SGRLA/FT | Evidencia DBA Fase 3")
    set_run_font(run, size=8, color=MUTED, bold=True)

    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Documento de evidencia técnica - Javier Mejía")
    set_run_font(run, size=8, color=MUTED)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[int], header_fill: str = LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, bold=True)
    set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=9.5)
            if len(value) <= 12:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("EVIDENCIA DBA DE EJECUCIÓN CONTROLADA")
    set_run_font(run, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Cierre Técnico de Fase 3")
    set_run_font(run, size=22, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Módulo Matrices de Riesgos - Sistema de Gestión de Riesgos LA/FT IHSS")
    set_run_font(run, size=13, color=MUTED)

    add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ("Proyecto", "RIESGO_LAVADO - IHSS"),
            ("Módulo", "Matrices de Riesgos"),
            ("Fase", "Fase 3. Modelo de datos y arquitectura Oracle"),
            ("Documento", "Evidencia DBA de ejecución controlada"),
            ("Versión", "1.0"),
            ("Fecha", date.today().strftime("%d/%m/%Y")),
            ("Responsable", "Javier Mejía"),
            ("Estado", "Ejecución completada y validada"),
            ("Ubicación", "docs/3. Módulo Matrices de Riesgos/Fase 3 - Modelo de datos y arquitectura Oracle/Evidencia_DBA"),
        ],
        [2200, 7160],
        header_fill=LIGHT_BLUE,
    )


def log_rows() -> list[tuple[str, str, str]]:
    rows: list[tuple[str, str, str]] = []
    for path in sorted(EVIDENCE_DIR.glob("*.log")):
        status = "Evidencia"
        name = path.name
        if "fallo" in name:
            status = "Incidente controlado"
        elif "validacion" in name:
            status = "Validación"
        elif "ejecucion" in name:
            status = "Ejecución"
        elif "preflight" in name:
            status = "Prevalidación"
        rows.append((name, status, f"{path.stat().st_size:,} bytes".replace(",", ".")))
    return rows


def build_document() -> None:
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_title_block(doc)

    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "Se ejecutó el paquete técnico de Fase 3 del módulo Matrices de Riesgos en el esquema RIESGO_LAVADO indicado por el responsable funcional. "
        "La ejecución siguió el orden definido en el documento de fase: prevalidación, estructura, registro de módulo, parametrización inicial y validación posterior."
    )
    doc.add_paragraph(
        "El cierre confirma que el modelo físico RL_MR_* quedó creado, el módulo fue registrado, los permisos iniciales fueron asignados a usuarios existentes, "
        "la metodología base quedó parametrizada y las validaciones posteriores no reportaron objetos inválidos."
    )

    doc.add_heading("2. Ambiente validado", level=1)
    add_table(
        doc,
        ["Elemento", "Resultado"],
        [
            ("Usuario de sesión", "RIESGO_LAVADO"),
            ("Esquema actual", "RIESGO_LAVADO"),
            ("Base de datos", "hpprod1"),
            ("Servicio", "hpprod1"),
            ("Servidor", "desdb"),
            ("Fecha de ejecución", date.today().strftime("%d/%m/%Y")),
            ("Alcance", "Scripts 01, 02 y 03 de Fase 3, más validación posterior."),
        ],
        [2500, 6860],
    )

    doc.add_heading("3. Orden de ejecución", level=1)
    add_table(
        doc,
        ["Paso", "Archivo / acción", "Resultado"],
        [
            ("00", "Prevalidación de ambiente, esquema, colisiones y permisos", "Completado."),
            ("01", "database/19_matrices_riesgos/01_create_rl_mr_estructura.sql", "Completado. Estructura RL_MR_* creada."),
            ("02", "database/19_matrices_riesgos/02_register_modulo_matrices_riesgos.sql - primer intento", "Detenido por PLS-00114 antes de registrar el módulo; incidente controlado."),
            ("02R", "database/19_matrices_riesgos/02_register_modulo_matrices_riesgos.sql - reintento", "Completado después de ajustar el nombre interno del procedimiento a asignar_mod_usuario."),
            ("03", "database/19_matrices_riesgos/03_seed_metodologia_matrices_riesgos.sql", "Completado. Metodología inicial cargada."),
            ("04", "Validación posterior de estructura, parametrización, módulo y objetos", "Completado sin objetos inválidos."),
        ],
        [900, 4660, 3800],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("4. Incidente controlado y corrección aplicada", level=1)
    add_table(
        doc,
        ["Control", "Detalle"],
        [
            ("Incidente", "El primer intento del script 02 presentó PLS-00114 por un identificador PL/SQL mayor a 30 caracteres."),
            ("Impacto", "El error ocurrió en compilación del bloque PL/SQL; no se registró el módulo en ese intento."),
            ("Corrección", "Se redujo el nombre interno del procedimiento a asignar_mod_usuario, manteniendo la misma lógica funcional."),
            ("Reintento", "El script 02 corregido se ejecutó correctamente y confirmó el registro del módulo y permisos iniciales."),
            ("Evidencia", "Se conservaron el log y el SQL del intento fallido, además del log y SQL del reintento exitoso."),
        ],
        [2300, 7060],
        header_fill=LIGHT_GOLD,
    )

    doc.add_heading("5. Validaciones finales", level=1)
    add_table(
        doc,
        ["Validación", "Resultado"],
        [
            ("Tablas RL_MR_*", "13 tablas creadas."),
            ("Secuencias SEQ_RL_MR_*", "13 secuencias creadas."),
            ("Índices RL_MR_*", "33 índices y restricciones indexadas relacionadas."),
            ("Comentarios de tablas", "13 comentarios de tabla aplicados."),
            ("Comentarios de columnas", "187 comentarios de columna aplicados."),
            ("Módulo", "Ruta /matrices-riesgos registrada con MOD_ID 10."),
            ("Permisos iniciales", "Usuarios existentes 1 y 2 asociados al MOD_ID 10."),
            ("Factores institucionales", "Proveedores 50%, Clientes/Patronos 25% y Empleados 25%."),
            ("Variables internas", "Siete variables por factor, con suma interna de 100% por cada factor."),
            ("Escalas", "Cinco registros por tipo VARIABLE, INHERENTE, RESIDUAL y CONTROL."),
            ("Estado OBSERVADA", "Incluido en la restricción física de estados de RL_MR_MATRICES."),
            ("Objetos inválidos", "No se detectaron objetos inválidos posteriores a la ejecución."),
        ],
        [3000, 6360],
        header_fill=LIGHT_GREEN,
    )

    doc.add_heading("6. Archivos de evidencia", level=1)
    rows = log_rows()
    if rows:
        add_table(doc, ["Archivo", "Tipo", "Tamaño"], rows, [5200, 2300, 1860])
    else:
        doc.add_paragraph("No se encontraron archivos de log en la carpeta de evidencia al momento de generar este documento.")

    doc.add_heading("7. Conclusión DBA", level=1)
    doc.add_paragraph(
        "Con base en la ejecución y validación posterior, la Fase 3 queda cerrada técnicamente desde el rol DBA para el esquema indicado. "
        "El paquete aplicado respeta el prefijo RL_MR_*, conserva trazabilidad, mantiene la metodología aprobada de Fase 2 y deja evidencia completa para revisión institucional."
    )
    doc.add_paragraph(
        "Cualquier ajuste solicitado posteriormente por el cliente deberá gestionarse mediante control de cambios, nueva versión metodológica o script incremental, sin alterar retroactivamente matrices cerradas."
    )

    props = doc.core_properties
    props.title = "Evidencia DBA - Cierre Técnico Fase 3 - Matrices de Riesgos"
    props.subject = "Sistema de Gestión de Riesgos LA/FT IHSS"
    props.keywords = "IHSS, SGRLA, Matrices de Riesgos, Fase 3, DBA, Oracle"
    props.comments = "Evidencia DBA de ejecución controlada y validación posterior de Fase 3."
    props.author = "Javier Mejía"

    doc.save(OUT_FILE)


if __name__ == "__main__":
    build_document()
    print(OUT_FILE)
