from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_ROOT = ROOT / "docs"
MATRICES_DIR = next(DOCS_ROOT.rglob("3. Módulo Matrices de Riesgos"))
OUT_DIR = MATRICES_DIR / "Fase 2 - Metodología LA-FT"
OUT_FILE = OUT_DIR / "Fase_2_Metodologia_LAFT_Matrices_Riesgos_SGRLA_IHSS_FINAL.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GOLD = "FFF2CC"
LIGHT_GREEN = "E2F0D9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run("IHSS - SGRLA/FT | Módulo Matrices de Riesgos | Fase 2")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_after = Pt(0)
    run = footer_p.add_run("Documento metodológico aprobado y cerrado")
    set_run_font(run, size=9, color=MUTED)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("DOCUMENTO METODOLÓGICO")
    set_run_font(run, size=10.5, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Fase 2. Levantamiento Metodológico LA/FT")
    set_run_font(run, size=22, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Módulo Matrices de Riesgos - Sistema de Gestión de Riesgos LA/FT IHSS")
    set_run_font(run, size=13, color=MUTED)

    meta = [
        ("Proyecto", "RIESGO_LAVADO - IHSS"),
        ("Módulo", "Matrices de Riesgos"),
        ("Fase", "Fase 2. Levantamiento metodológico LA/FT"),
        ("Versión", "1.4"),
        ("Fecha", date.today().strftime("%d/%m/%Y")),
        ("Estado", "Fase 2 aprobada y cerrada"),
        ("Responsable", "Javier Mejía"),
        ("Fuente", "Fase 1 aprobada, Fase 2 aprobada y criterios de control documental vigentes"),
        ("Ubicación", "docs/3. Módulo Matrices de Riesgos/Fase 2 - Metodología LA-FT"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, [2300, 7060])
    for row, (label, value) in zip(table.rows, meta):
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[int], header_fill: str = LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], header_fill)
        table.rows[0].cells[idx].paragraphs[0].add_run(header).bold = True
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            if len(value) <= 8:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def add_checklist(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_title_block(doc)

    doc.add_heading("0. Control de cambios", level=1)
    add_table(
        doc,
        ["Versión", "Fecha", "Descripción del cambio", "Responsable", "Estado"],
        [
            ("1.0", date.today().strftime("%d/%m/%Y"), "Creación de la metodología base de Fase 2 para revisión.", "Javier Mejía", "Revisión"),
            ("1.1", date.today().strftime("%d/%m/%Y"), "Ajuste metodológico al requerimiento del cliente: factores institucionales Proveedores 50%, Clientes/Patronos 25% y Empleados 25%; variables definidas por Cumplimiento.", "Javier Mejía", "Revisión"),
            ("1.2", date.today().strftime("%d/%m/%Y"), "Incorporación de lectura funcional de imágenes del cliente, campos mínimos por submódulo y trazabilidad módulo/submódulo.", "Javier Mejía", "Revisión"),
            ("1.3", date.today().strftime("%d/%m/%Y"), "Ajustes finales: ponderaciones internas por variable, DNP como integración obligatoria futura, reportería de solidez de controles y normalización del nombre del módulo.", "Javier Mejía", "Final para aprobación funcional"),
            ("1.4", date.today().strftime("%d/%m/%Y"), "Aprobación formal y cierre de Fase 2 como base metodológica para iniciar Fase 3.", "Javier Mejía", "Aprobado y cerrado"),
        ],
        [900, 1350, 4400, 1500, 1210],
    )

    doc.add_heading("1. Propósito de la fase", level=1)
    doc.add_paragraph(
        "La Fase 2 convierte el requerimiento funcional del Módulo Matrices de Riesgos en una metodología operativa, medible, versionable y auditable. "
        "Este documento queda como base funcional aprobada para factores, criterios, escalas, pesos, reglas de cálculo, estados y controles antes de diseñar tablas definitivas, endpoints o pantallas."
    )

    doc.add_heading("2. Alcance metodológico", level=1)
    add_checklist(
        doc,
        [
            "Definir los tres factores institucionales obligatorios: Proveedores, Clientes/Patronos y Empleados.",
            "Mantener la ponderación institucional establecida: Proveedores 50%, Clientes/Patronos 25% y Empleados 25%.",
            "Definir criterios y variables medibles por cada factor institucional.",
            "Establecer escalas de calificación de 1 a 5.",
            "Definir ponderaciones internas por variable dentro de cada factor institucional, asegurando que totalicen 100% por factor.",
            "Separar riesgo inherente, controles mitigantes, riesgo residual y planes de acción.",
            "Definir reglas de obligatoriedad, aprobación, vigencia, versionamiento y auditoría.",
            "Mantener la fórmula y ponderaciones como base metodológica aprobada para el diseño técnico de Fase 3.",
        ],
    )

    doc.add_heading("3. Principios metodológicos", level=1)
    add_table(
        doc,
        ["Principio", "Regla de aplicación"],
        [
            ("Trazabilidad", "Toda matriz debe conservar quién la creó, quién la revisó, quién la aprobó, fecha, versión metodológica y justificación."),
            ("Versionamiento", "Los cambios de factores, escalas o pesos no deben alterar matrices cerradas; cada cierre debe guardar snapshot metodológico."),
            ("Cálculo en backend", "El frontend no debe calcular riesgo inherente, mitigación ni residual; solo debe capturar datos y mostrar resultados."),
            ("Auditoría", "Toda creación, edición, aprobación, recalculo, cierre, inactivación, exportación o impresión debe registrar auditoría."),
            ("Evidencia", "Toda justificación relevante debe permitir soporte documental cuando aplique, bajo reglas de evidencias protegidas."),
            ("Aprobación", "Ninguna metodología nueva entra en vigencia sin revisión y aprobación funcional."),
        ],
        [2300, 7060],
    )

    doc.add_heading("4. Submódulos requeridos por el cliente", level=1)
    add_table(
        doc,
        ["Submódulo", "Regla metodológica requerida", "Cobertura en esta versión"],
        [
            ("Identificación de factores de riesgo", "Debe alimentar riesgos identificados, identificación/calificación, mitigadores y plan de acción.", "Cubierto como estructura funcional obligatoria por secciones."),
            ("Perfilamiento de factores de riesgo", "Debe calificar Clientes/Patronos, Proveedores y Empleados con variables definidas por Cumplimiento.", "Cubierto con factores institucionales 50/25/25 y variables independientes."),
            ("Matriz resumen por factor", "Debe mostrar resumen individual por factor y resumen institucional.", "Cubierto como salida mínima obligatoria."),
            ("Mapa de calor", "Debe mostrar información por factor y consolidado institucional.", "Cubierto para riesgo inherente y residual."),
            ("Reportería estadística dinámica", "Debe permitir reportes dinámicos, filtros, comparativos y explotación de resultados.", "Cubierto como catálogo inicial sujeto a ampliación durante el desarrollo."),
        ],
        [2500, 4300, 2560],
    )

    doc.add_heading("4.1 Lectura funcional de imágenes del cliente", level=2)
    doc.add_paragraph(
        "Las imágenes del requerimiento se interpretan como referencia funcional de las hojas, pantallas, cálculos y reportes que debe cubrir el módulo. "
        "El sistema no copiará literalmente las imágenes; las convertirá en formularios, catálogos, reglas de cálculo, vistas de resumen, mapas y reportes auditados."
    )
    add_table(
        doc,
        ["Referencia", "Lectura funcional", "Submódulo destino"],
        [
            ("Matriz factor / variables / ponderación / decisión / observaciones", "Permite configurar o revisar variables de cada factor institucional, su ponderación interna, decisión y observaciones.", "Perfilamiento / scoring."),
            ("Evaluación del riesgo del cliente", "Define rangos de valor mínimo, valor máximo, escala y criterio de supervisión para clasificar el resultado.", "Perfilamiento / scoring y matriz resumen."),
            ("Detalle de clientes y ponderación", "Muestra cálculo individual, total del factor, porcentaje individual y ponderación institucional del 25% para Clientes/Patronos.", "Matriz resumen por factor."),
            ("Identificación del riesgo", "Registra factor, subfactor, variable, evento, probabilidad, impacto, puntaje total y nivel de riesgo inherente.", "Identificación de factores de riesgo."),
            ("Evaluación de mitigadores", "Registra controles, periodicidad, oportunidad, automatización, procedimientos, calidad del control y riesgo residual.", "Controles mitigantes y riesgo residual."),
            ("Plan de acción", "Registra actividades, responsables, periodicidad, fechas, medios de prueba, observaciones y estado.", "Planes de acción."),
            ("Crear riesgo", "Representa formulario de captura de riesgo, impacto, frecuencia, descripción, evidencia, asociaciones, procesos, controles y responsables.", "Captura funcional de riesgos."),
            ("Crear control", "Representa formulario de control, diseño, ejecución, evidencia, responsables y resultado de calificación.", "Captura funcional de controles."),
            ("Mapa de calor", "Representa consulta por fecha, período, riesgo inherente/residual, sistema, proceso, comparativo y descarga.", "Mapa de calor."),
            ("Gráficos estadísticos", "Representa distribución por niveles, sujetos activos, actualización de expedientes, debida diligencia y solidez de controles.", "Reportería estadística dinámica."),
        ],
        [2100, 5200, 2060],
    )

    doc.add_heading("4.2 Campos mínimos por submódulo", level=2)
    add_table(
        doc,
        ["Submódulo", "Campos mínimos", "Regla de negocio"],
        [
            ("Identificación de factores de riesgo", "Factor institucional, subfactor, variable, evento de riesgo, probabilidad, puntos de probabilidad, impacto, puntos de impacto, total, nivel inherente.", "Probabilidad e impacto deben generar el nivel inherente desde backend."),
            ("Perfilamiento / scoring", "Factor institucional, variable, ponderación interna, decisión, observaciones, valor mínimo, valor máximo, escala y criterio de supervisión.", "Cada factor institucional se mide de manera independiente con variables definidas por Cumplimiento."),
            ("Controles mitigantes", "Existencia de control, descripción, periodicidad, oportunidad, automatización, procedimientos, calidad, promedio, calidad total, puntos residuales y riesgo residual.", "El control solo mitiga si tiene evidencia, responsable y evaluación suficiente."),
            ("Planes de acción", "Actividades, responsables, periodicidad, fecha de inicio, fecha de finalización, medios de prueba, observaciones y estado.", "Riesgo residual alto o crítico debe exigir plan de acción."),
            ("Matriz resumen", "Detalle por sujeto, riesgo individual, porcentaje individual, ponderado por factor, cantidad, promedio residual, escala y criterio de supervisión.", "Debe mostrar resumen por factor institucional y consolidado institucional."),
            ("Mapa de calor", "Fecha, rango, período, inherente/residual, sistema, proceso, nivel, comparativo y descarga.", "Debe permitir vista actual, histórica, por factor e institucional."),
            ("Reportería dinámica", "Gráficos por nivel, factor, sujeto, actualización de expedientes, debida diligencia, controles y distribución institucional.", "Toda exportación o descarga debe auditarse."),
        ],
        [2200, 5000, 2160],
    )

    doc.add_heading("5. Factores institucionales obligatorios", level=1)
    doc.add_paragraph(
        "Esta regla se toma como base fija del requerimiento del cliente y no debe modificarse en la metodología de la versión inicial. "
        "El riesgo institucional se compone de tres factores generales: Proveedores, Clientes/Patronos y Empleados. "
        "Cada factor se mide de manera independiente mediante variables de riesgo definidas por la Sección de Cumplimiento."
    )
    institutional_factor_rows = [
        ("FI01", "Proveedores", "50%", "Factor institucional con mayor peso dentro del riesgo total."),
        ("FI02", "Clientes/Patronos", "25%", "Factor institucional asociado a clientes o patronos evaluados por variables aprobadas."),
        ("FI03", "Empleados", "25%", "Factor institucional asociado a empleados y exposición funcional interna."),
    ]
    add_table(doc, ["Código", "Factor institucional", "Peso fijo", "Regla"], institutional_factor_rows, [900, 2600, 1100, 4760])
    doc.add_paragraph(
        "La suma de los tres factores institucionales es 100% del riesgo institucional. "
        "Las variables internas de cada factor podrán parametrizarse, pero no sustituyen esta ponderación institucional."
    )

    doc.add_heading("6. Sujetos o alcances evaluables", level=1)
    add_table(
        doc,
        ["Alcance", "Uso dentro de la matriz", "Observación"],
        [
            ("Patrono", "Evaluación de riesgo asociado a empleador, aportaciones, coincidencias y comportamiento registrado.", "Puede reutilizar información de monitoreo y coincidencias."),
            ("Proveedor", "Evaluación de riesgo asociado a contratación, rubro, ubicación, antecedentes y controles.", "Debe quedar preparado para futuras integraciones."),
            ("Empleado", "Evaluación de riesgo asociado a cargo, exposición, funciones sensibles y hallazgos internos.", "Debe respetar permisos y confidencialidad."),
            ("Área", "Evaluación de procesos o unidades internas con exposición LA/FT.", "Útil para matriz institucional."),
            ("Proceso", "Evaluación de procesos críticos del sistema o de operación institucional.", "Puede asociarse a controles internos."),
            ("Caso positivo", "Evaluación derivada de hallazgos de monitoreo o seguimiento.", "Debe conservar relación con evidencias y seguimientos."),
            ("Institucional", "Matriz global consolidada para visión ejecutiva.", "Debe alimentarse de resultados agregados."),
        ],
        [1600, 5000, 2760],
    )

    doc.add_heading("7. Variables metodológicas por factor institucional", level=1)
    variable_rows = [
        ("V01", "Perfil del sujeto evaluado", "Identidad, tipo de sujeto, naturaleza, exposición y condición institucional.", "Cumplimiento define si aplica y su peso interno por factor."),
        ("V02", "Actividad, rubro o función", "Rubro económico, función sensible, proceso crítico o actividad con exposición LA/FT.", "Puede variar entre Proveedores, Clientes/Patronos y Empleados."),
        ("V03", "Ubicación geográfica", "Zona, municipio, país, jurisdicción o exposición territorial.", "Debe permitir catálogos o captura controlada."),
        ("V04", "Antecedentes y coincidencias", "Coincidencias en listas, positivos manuales, historial de seguimientos y hallazgos.", "Puede integrarse con Monitoreo de Listas."),
        ("V05", "Comportamiento transaccional u operativo", "Volumen, frecuencia, variaciones, patrones inusuales o comportamiento operativo.", "Puede provenir de datos internos o captura manual."),
        ("V06", "Canal, producto o relación institucional", "Tipo de interacción, canal de atención, relación contractual o exposición operativa.", "Debe parametrizarse por factor institucional."),
        ("V07", "Control interno y evidencia disponible", "Calidad documental, controles existentes, soportes, seguimiento y capacidad de mitigación.", "Debe vincularse con evidencias y controles."),
    ]
    add_table(doc, ["Código", "Variable / dimensión", "Descripción", "Regla"], variable_rows, [900, 2400, 3900, 2160])

    doc.add_paragraph(
        "Las variables V01-V07 son dimensiones metodológicas iniciales. No sustituyen los tres factores institucionales obligatorios. "
        "Cada factor institucional tendrá su propia matriz de variables, con pesos internos definidos y aprobados por la Sección de Cumplimiento."
    )

    doc.add_heading("8. Ponderación interna de variables por factor", level=1)
    add_table(
        doc,
        ["Factor institucional", "Variables internas", "Regla de ponderación interna"],
        [
            ("Proveedores", "Variables aprobadas por Cumplimiento para evaluación de proveedores.", "Los pesos internos deben sumar 100% dentro del factor Proveedores."),
            ("Clientes/Patronos", "Variables aprobadas por Cumplimiento para evaluación de clientes o patronos.", "Los pesos internos deben sumar 100% dentro del factor Clientes/Patronos."),
            ("Empleados", "Variables aprobadas por Cumplimiento para evaluación de empleados.", "Los pesos internos deben sumar 100% dentro del factor Empleados."),
        ],
        [2300, 4300, 2760],
    )

    doc.add_heading("9. Escala general de calificación", level=1)
    add_table(
        doc,
        ["Valor", "Nivel", "Criterio general"],
        [
            ("1", "Muy bajo", "Exposición mínima, información completa, sin alertas relevantes y controles suficientes."),
            ("2", "Bajo", "Exposición limitada, alertas menores o controles razonables."),
            ("3", "Medio", "Exposición moderada, información parcial, comportamiento a revisar o controles mejorables."),
            ("4", "Alto", "Exposición relevante, alertas importantes, historial sensible o controles débiles."),
            ("5", "Crítico", "Exposición severa, coincidencia confirmada, alerta grave, falta de soporte o control insuficiente."),
        ],
        [900, 1600, 6860],
    )

    doc.add_heading("10. Criterios mínimos por variable", level=1)
    add_table(
        doc,
        ["Variable", "Criterios mínimos", "Fuente sugerida"],
        [
            ("Perfil del sujeto", "Tipo de sujeto, naturaleza jurídica o natural, relación con IHSS, exposición pública o sensibilidad funcional.", "Registro interno, usuario evaluador, catálogos."),
            ("Actividad o función", "Rubro, proceso, cargo, función crítica, actividad con manejo financiero o documental.", "Registro interno, clasificación funcional."),
            ("Ubicación", "Departamento, municipio, país, zona fronteriza o jurisdicción de mayor exposición.", "Catálogo geográfico, captura manual."),
            ("Antecedentes", "Coincidencias, positivos manuales, seguimientos, evidencias, observaciones y reincidencias.", "Monitoreo de Listas, evidencias, auditoría."),
            ("Comportamiento", "Volumen, frecuencia, variación, comportamiento atípico o desviación frente al perfil esperado.", "Datos internos, carga manual o integración futura."),
            ("Canal o relación", "Canal operativo, relación contractual, proveedor, patrono, empleado, proceso o área vinculada.", "Sistema, catálogos y captura manual."),
            ("Control interno", "Existencia de soporte, revisión, evidencia, controles preventivos, responsables y seguimiento.", "Evidencias, controles registrados, planes de acción."),
        ],
        [2100, 5000, 2260],
    )

    doc.add_heading("11. Fórmula base de riesgo inherente", level=1)
    doc.add_paragraph(
        "La fórmula base se calcula en dos niveles. Primero, cada factor institucional obtiene su calificación propia mediante las variables internas definidas por la Sección de Cumplimiento. "
        "Después, el riesgo institucional se calcula aplicando los pesos fijos: Proveedores 50%, Clientes/Patronos 25% y Empleados 25%."
    )
    add_table(
        doc,
        ["Elemento", "Definición"],
        [
            ("Calificación de variable", "Valor entre 1 y 5 asignado a cada variable aprobada por Cumplimiento."),
            ("Calificación del factor institucional", "Resultado ponderado de las variables internas del factor Proveedores, Clientes/Patronos o Empleados."),
            ("Peso institucional fijo", "Proveedores 50%, Clientes/Patronos 25% y Empleados 25%."),
            ("Riesgo inherente institucional", "Suma de cada factor institucional multiplicado por su peso fijo."),
            ("Regla técnica", "El cálculo debe ejecutarse en backend y guardar snapshot de factores, variables, pesos, valores y resultado."),
        ],
        [2300, 7060],
    )
    doc.add_paragraph("Fórmula por factor: Calificación del factor = SUMA(Calificación de variable x Peso interno de variable) / 100.")
    doc.add_paragraph("Fórmula institucional: Riesgo inherente = (Proveedores x 50%) + (Clientes/Patronos x 25%) + (Empleados x 25%).")

    doc.add_heading("12. Rangos de riesgo inherente", level=1)
    add_table(
        doc,
        ["Rango", "Nivel", "Tratamiento mínimo"],
        [
            ("1.00 - 1.80", "Muy bajo", "Registro normal y revisión ordinaria."),
            ("1.81 - 2.60", "Bajo", "Registro normal con observaciones si aplica."),
            ("2.61 - 3.40", "Medio", "Revisión funcional y soporte de criterios relevantes."),
            ("3.41 - 4.20", "Alto", "Revisión obligatoria, controles mitigantes y seguimiento."),
            ("4.21 - 5.00", "Crítico", "Revisión prioritaria, plan de acción obligatorio y aprobación superior."),
        ],
        [1700, 1500, 6160],
    )

    doc.add_heading("13. Controles mitigantes", level=1)
    doc.add_paragraph(
        "Los controles mitigantes reducen el riesgo inherente cuando existe evidencia suficiente de diseño, aplicación y seguimiento. "
        "El control no debe reducir riesgo si no tiene responsable, evidencia o estado verificable."
    )
    add_table(
        doc,
        ["Nivel de control", "Efectividad sugerida", "Criterio"],
        [
            ("Sin control", "0%", "No existe control documentado o no se puede verificar."),
            ("Débil", "10%", "Existe control informal o incompleto, sin evidencia suficiente."),
            ("Moderado", "25%", "Existe control documentado, pero con cobertura parcial o seguimiento limitado."),
            ("Fuerte", "40%", "Existe control documentado, aplicado y con evidencia suficiente."),
            ("Muy fuerte", "55%", "Existe control robusto, aplicado, evidenciado, revisado y con seguimiento formal."),
        ],
        [1900, 1700, 5760],
    )

    doc.add_heading("14. Fórmula base de riesgo residual", level=1)
    doc.add_paragraph(
        "La fórmula base propuesta calcula el riesgo residual aplicando la efectividad mitigante aprobada sobre el riesgo inherente. "
        "Cuando existan varios controles, se debe calcular una efectividad consolidada sin exceder el máximo permitido por la metodología vigente."
    )
    add_table(
        doc,
        ["Elemento", "Definición"],
        [
            ("Efectividad mitigante", "Porcentaje calculado a partir de controles aprobados y evidenciados."),
            ("Tope sugerido", "La reducción máxima sugerida es 55% salvo aprobación metodológica distinta."),
            ("Riesgo residual", "Riesgo inherente multiplicado por el complemento de la mitigación."),
            ("Regla técnica", "El cálculo residual debe ejecutarse en backend y conservar detalle de controles usados."),
        ],
        [2300, 7060],
    )
    doc.add_paragraph("Fórmula propuesta: Riesgo residual = Riesgo inherente x (1 - Efectividad mitigante).")

    doc.add_heading("15. Rangos de riesgo residual", level=1)
    add_table(
        doc,
        ["Rango", "Nivel residual", "Acción requerida"],
        [
            ("1.00 - 1.80", "Muy bajo", "Cierre normal o seguimiento ordinario."),
            ("1.81 - 2.60", "Bajo", "Seguimiento normal con observaciones si aplica."),
            ("2.61 - 3.40", "Medio", "Seguimiento funcional y revisión de controles."),
            ("3.41 - 4.20", "Alto", "Plan de acción obligatorio, responsable y fecha compromiso."),
            ("4.21 - 5.00", "Crítico", "Plan de acción prioritario, revisión superior y control reforzado."),
        ],
        [1700, 1700, 5960],
    )

    doc.add_heading("16. Planes de acción", level=1)
    add_table(
        doc,
        ["Condición", "Regla obligatoria"],
        [
            ("Residual alto", "Debe exigir plan de acción, responsable, fecha compromiso, seguimiento y evidencia de cierre."),
            ("Residual crítico", "Debe exigir plan prioritario, aprobación superior, seguimiento reforzado y evidencia obligatoria."),
            ("Plan vencido", "Debe quedar visible como alerta y registrar auditoría de cambios o reprogramaciones."),
            ("Cierre de plan", "Debe exigir comentario, evidencia cuando aplique y auditoría."),
        ],
        [2300, 7060],
    )

    doc.add_heading("17. Matriz resumen por factor e institucional", level=1)
    add_table(
        doc,
        ["Salida mínima", "Contenido requerido"],
        [
            ("Resumen por Proveedores", "Calificación individual, variables aplicadas, riesgo inherente, controles, residual y nivel."),
            ("Resumen por Clientes/Patronos", "Calificación individual, variables aplicadas, riesgo inherente, controles, residual y nivel."),
            ("Resumen por Empleados", "Calificación individual, variables aplicadas, riesgo inherente, controles, residual y nivel."),
            ("Resumen institucional", "Consolidado calculado con Proveedores 50%, Clientes/Patronos 25% y Empleados 25%."),
            ("Explotación de datos", "Resultados por sujeto, variable, factor, período, nivel de riesgo y estado."),
        ],
        [2600, 6760],
    )

    doc.add_heading("18. Mapa de calor y reportería dinámica", level=1)
    add_table(
        doc,
        ["Elemento", "Regla mínima"],
        [
            ("Mapa actual", "Debe mostrar riesgo inherente y residual vigente por factor e institucional."),
            ("Rango de fechas", "Debe permitir consulta desde-hasta y a fecha específica."),
            ("Nivel de riesgo", "Debe permitir filtrar niveles bajo, medio, alto y crítico."),
            ("Comparativo histórico", "Debe permitir comparar períodos o versiones metodológicas cuando existan datos."),
            ("Perfil institucional", "Debe mostrar distribución por Proveedores, Clientes/Patronos y Empleados."),
            ("Perfil por factor", "Debe mostrar resultados de cada factor institucional y sus variables internas."),
            ("Solidez de controles", "Debe mostrar cantidad de controles definidos, nivel de efectividad, calidad del control, factor asociado y distribución por nivel."),
            ("Exportación", "Toda exportación o impresión debe quedar auditada."),
        ],
        [2300, 7060],
    )

    doc.add_heading("19. Estados funcionales de una matriz", level=1)
    add_table(
        doc,
        ["Estado", "Descripción", "Acciones permitidas"],
        [
            ("Borrador", "Matriz creada sin cálculo final.", "Editar, guardar, cancelar."),
            ("En evaluación", "Datos capturados y criterios en revisión.", "Editar criterios, agregar evidencia, calcular."),
            ("Calculada", "Riesgo inherente y residual generados por backend.", "Revisar, enviar a aprobación, recalcular con motivo."),
            ("En revisión", "Pendiente de revisión funcional.", "Aprobar, observar, devolver."),
            ("Observada", "Devuelta por inconsistencias o falta de soporte.", "Corregir, justificar, reenviar."),
            ("Aprobada", "Resultado aceptado como oficial.", "Cerrar, exportar, generar reporte."),
            ("Cerrada", "Matriz final protegida contra edición metodológica.", "Consultar, exportar, auditar."),
            ("Inactiva", "Matriz anulada lógicamente con motivo.", "Consultar auditoría."),
        ],
        [1500, 4200, 3660],
    )

    doc.add_heading("20. Reglas de auditoría", level=1)
    add_checklist(
        doc,
        [
            "Crear matriz.",
            "Editar criterios, valores o ponderaciones capturadas.",
            "Calcular riesgo inherente.",
            "Registrar, modificar o eliminar lógicamente controles.",
            "Calcular riesgo residual.",
            "Cambiar estado de la matriz.",
            "Aprobar, observar, cerrar o inactivar matriz.",
            "Registrar o cerrar plan de acción.",
            "Visualizar información sensible cuando aplique.",
            "Exportar Excel, PDF, impresión o generación de reporte.",
        ],
    )

    doc.add_heading("21. Datos de origen y captura", level=1)
    add_table(
        doc,
        ["Tipo de dato", "Origen permitido", "Regla"],
        [
            ("Datos internos", "Módulos existentes, catálogos, usuarios y monitoreo.", "Deben consumirse desde backend y con autorización por módulo."),
            ("Datos capturados", "Formulario de matriz.", "Deben validarse en frontend y backend."),
            ("Evidencias", "Módulo de evidencias protegido.", "Deben conservar trazabilidad y descarga auditada."),
            ("Criterios metodológicos", "Versión aprobada de metodología.", "No deben modificarse retroactivamente en matrices cerradas."),
            ("Resultados", "Motor de cálculo backend.", "Deben guardarse con snapshot del cálculo."),
        ],
        [2000, 3600, 3760],
    )

    doc.add_heading("22. Criterios de aceptación de Fase 2", level=1)
    add_checklist(
        doc,
        [
            "La metodología puede explicarse sin depender de código fuente.",
            "Los tres factores institucionales obligatorios están definidos: Proveedores, Clientes/Patronos y Empleados.",
            "La ponderación institucional fija queda establecida: Proveedores 50%, Clientes/Patronos 25% y Empleados 25%.",
            "Las variables internas por factor quedan documentadas como definibles por la Sección de Cumplimiento.",
            "Las ponderaciones institucionales totalizan 100%.",
            "La fórmula oficial respeta los tres factores institucionales y queda aprobada como base metodológica para Fase 3.",
            "Quedan definidas reglas de riesgo inherente, controles, residual y planes de acción.",
            "Quedan definidos los submódulos solicitados: identificación, scoring, resumen, mapa de calor y reportería dinámica.",
            "Las referencias visuales del cliente quedan traducidas a campos mínimos, reglas y salidas por submódulo.",
            "La documentación explica cómo se manejarán identificación del riesgo, mitigadores, plan de acción, mapas, gráficos y resúmenes.",
            "Quedan definidas auditorías mínimas que debe generar el módulo.",
            "Quedan definidas condiciones para pasar a Fase 3 sin mezclar scripts experimentales.",
        ],
    )

    doc.add_heading("23. Decisiones funcionales de cierre", level=1)
    add_table(
        doc,
        ["ID", "Decisión", "Estado"],
        [
            ("M2-D01", "Las variables internas de cada factor institucional se definirán por Cumplimiento durante la parametrización de Fase 3.", "Cerrado para Fase 2"),
            ("M2-D02", "Los pesos internos de variables deberán totalizar 100% dentro de cada factor institucional.", "Cerrado para Fase 2"),
            ("M2-D03", "Los rangos de riesgo inherente y residual se usarán como base metodológica inicial y podrán parametrizarse en Fase 3 con aprobación funcional.", "Cerrado para Fase 2"),
            ("M2-D04", "El tope máximo de mitigación por controles queda como criterio metodológico sujeto a parametrización controlada en Fase 3.", "Cerrado para Fase 2"),
            ("M2-D05", "Los reportes dinámicos mínimos quedan definidos: matriz, resumen institucional, mapa de calor, solidez de controles, planes de acción y exportaciones.", "Cerrado para Fase 2"),
            ("M2-D06", "La calificación de riesgo por patrono deberá contemplarse como integración obligatoria futura hacia DNP, sujeta a contrato técnico, seguridad, autorización institucional y disponibilidad de interfaz.", "Integración futura obligatoria"),
            ("M2-D07", "Las columnas visibles de referencia del cliente se considerarán base funcional mínima; la obligatoriedad/configuración se definirá técnicamente en Fase 3.", "Cerrado para Fase 2"),
        ],
        [1000, 6500, 1860],
        header_fill=LIGHT_GOLD,
    )

    doc.add_heading("24. Cierre formal de Fase 2", level=1)
    doc.add_paragraph(
        "Con la aprobación funcional registrada, la Fase 2 queda cerrada como versión metodológica vigente para el Módulo Matrices de Riesgos. "
        "La siguiente etapa corresponde al diseño técnico de Fase 3, manteniendo sin modificación los pesos institucionales definidos por el cliente."
    )
    add_table(
        doc,
        ["Elemento", "Resultado de cierre"],
        [
            ("Estado final", "Fase 2 aprobada y cerrada."),
            ("Aprobador", "Javier Mejía."),
            ("Fecha de aprobación", date.today().strftime("%d/%m/%Y")),
            ("Regla institucional fija", "Proveedores 50%, Clientes/Patronos 25% y Empleados 25%."),
            ("DNP", "Integración obligatoria futura para calificación de riesgo por patrono, sujeta a contrato técnico, seguridad, autorización institucional y disponibilidad de interfaz."),
            ("Base de Fase 3", "Diseño técnico de tablas, servicios backend, endpoints, auditoría, permisos, frontend y scripts controlados."),
            ("Restricción documental", "No generar versiones paralelas ni modificar esta fase sin aprobación funcional."),
        ],
        [2600, 6760],
        header_fill=LIGHT_GREEN,
    )

    props = doc.core_properties
    props.title = "Fase 2 - Metodología LA/FT - Matrices de Riesgos"
    props.subject = "Sistema de Gestión de Riesgos LA/FT IHSS"
    props.keywords = "IHSS, SGRLA, Matrices de Riesgos, Fase 2, metodología LA/FT"
    props.comments = "Documento elaborado desde el repositorio local RIESGO_LAVADO. Fase 2 aprobada y cerrada."
    props.author = "Javier Mejía"

    doc.save(OUT_FILE)


if __name__ == "__main__":
    build_document()
    print(OUT_FILE)
