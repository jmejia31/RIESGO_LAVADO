from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "7. Politica de Repositorio"
OUT = OUT_DIR / "Politica_Repositorio_SGRLA_IHSS.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
TABLE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
WARN_FILL = "FFF2CC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_document_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Politica de Repositorio")
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = INK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Sistema de Gestion de Riesgos LAFT - IHSS")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_callout(doc, label, text, fill=CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(label + ": ")
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE
    p.add_run(text)
    doc.add_paragraph()


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Detalle"
    for cell in hdr:
        set_cell_shading(cell, TABLE_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    set_table_geometry(table, [2700, 6660])
    doc.add_paragraph()


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, TABLE_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Politica de Repositorio SGRLA-IHSS | Version 1.0")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_document_styles(doc)
    add_title(doc)

    add_kv_table(
        doc,
        [
            ("Documento", "Politica de Repositorio"),
            ("Proyecto", "Sistema de Gestion de Riesgos LAFT - IHSS"),
            ("Version", "1.0"),
            ("Fecha", "01/07/2026"),
            ("Estado", "Version final de trabajo, lista para revision y aprobacion antes de subir cambios."),
            ("Fuente", "Repositorio local C:/RIESGO_LAVADO"),
            ("Responsable de aprobacion", "Usuario responsable del proyecto"),
        ],
    )

    add_callout(
        doc,
        "Regla principal",
        "Ningun cambio documental debe subirse al repositorio como aprobado si antes no fue revisado en Word, depurado de borradores y autorizado como version final.",
        WARN_FILL,
    )

    doc.add_heading("1. Objetivo", 1)
    add_para(
        doc,
        "Establecer una politica clara para mantener el repositorio ordenado, evitar duplicados, controlar documentos finales y asegurar que cada entrega conserve separacion entre informacion de cliente y documentacion tecnica para desarrolladores.",
    )

    doc.add_heading("2. Alcance", 1)
    add_bullets(
        doc,
        [
            "Aplica a documentos Word, diagramas, scripts, capturas, reportes y archivos generados durante el desarrollo.",
            "Aplica antes de cualquier commit, subida a repositorio remoto o entrega formal del proyecto.",
            "Aplica a modulos actuales y futuros, incluyendo Monitoreo de Listas, Seguridad, Auditoria, Evidencias, Base de Datos y Matrices de Riesgo.",
        ],
    )

    doc.add_heading("3. Flujo obligatorio antes de subir cambios", 1)
    add_numbers(
        doc,
        [
            "Preparar el documento o cambio en una carpeta de trabajo local controlada.",
            "Revisar primero el documento en formato Word cuando el entregable sea documental.",
            "Confirmar que el documento tenga version, fecha, estado, fuente, responsable y alcance.",
            "Eliminar borradores, duplicados, imagenes temporales, logs y archivos auxiliares que no formen parte de la entrega.",
            "Validar que los nombres sean claros, consistentes y asociados al modulo correspondiente.",
            "Separar documentos para cliente y documentos para desarrolladores cuando el contenido tenga audiencias distintas.",
            "Solicitar aprobacion de la version final antes de considerar el documento como aprobado.",
            "Subir solamente archivos finales, revisados y necesarios.",
        ],
    )

    doc.add_heading("4. Reglas de control documental", 1)
    add_matrix(
        doc,
        ["Regla", "Aplicacion obligatoria", "Resultado esperado"],
        [
            ("Revisar primero en Word", "Todo documento final debe abrir y revisarse como .docx.", "Se evitan errores de formato y contenido."),
            ("Aprobar version final", "El usuario responsable confirma la version final antes de subirla.", "No se publican documentos sin visto bueno."),
            ("No subir borradores", "Archivos con estado borrador, prueba o temporal no se suben.", "Repositorio limpio y confiable."),
            ("No subir duplicados", "Debe existir una sola version vigente por documento final.", "Se evita confusion de versiones."),
            ("No subir imagenes temporales", "Capturas intermedias se eliminan si ya estan embebidas en Word.", "Se reduce ruido documental."),
            ("No dejar generados sin revisar", "Archivos creados por scripts se validan antes de conservarlos.", "Solo queda material util."),
            ("Usar nombres claros", "El nombre debe indicar modulo, audiencia y tipo de documento.", "Busqueda y mantenimiento mas simples."),
            ("Separar audiencias", "Cliente y desarrollador deben tener documentos separados cuando aplique.", "Cada lector recibe informacion adecuada."),
        ],
        [2100, 3960, 3300],
    )

    doc.add_heading("5. Estructura vigente de docs", 1)
    add_matrix(
        doc,
        ["Carpeta", "Uso correcto"],
        [
            ("0. Documentacion del Cliente", "Material base o documentos orientados al cliente."),
            ("1. Bases de Datos", "Bases SQL de referencia o insumos de base de datos."),
            ("2. Modulo Monitoreo de Listas", "Documentacion especifica del modulo Monitoreo de Listas."),
            ("3. Modulo Matrices de Riesgos", "Analisis y documentos de inicio del modulo Matrices de Riesgo."),
            ("4. Base de Datos", "Guias Word de ejecucion segura y control de base de datos."),
            ("5. Documentacion Modular", "Paquete full vigente por modulo, separado cliente/desarrollador."),
            ("6. Diagramas", "Diagramas funcionales, tecnicos y versiones finales de arquitectura visual."),
            ("7. Politica de Repositorio", "Politicas operativas para orden, revision y subida de cambios."),
            ("Herramientas", "Instaladores o utilidades necesarias para el proyecto."),
        ],
        [3300, 6060],
    )

    doc.add_heading("6. Convencion de nombres", 1)
    add_bullets(
        doc,
        [
            "Usar prefijo numerico cuando el documento pertenezca a un orden modular: 02_Usuarios, 04_Monitoreo_Listas, 11_Base_Datos.",
            "Indicar audiencia: Version_Cliente, Version_Desarrollador, Version_Tecnica o Version_Full.",
            "Evitar nombres genericos como final, nuevo, copia, prueba, temp o version2.",
            "Usar guiones bajos para documentos tecnicos generados y nombres legibles para documentos manuales existentes.",
            "Mantener una sola version final vigente en la carpeta correspondiente.",
        ],
    )

    doc.add_heading("7. Politica para imagenes y capturas", 1)
    add_bullets(
        doc,
        [
            "Las capturas deben ser actuales del sistema y no deben reutilizarse como evidencia si pertenecen a un periodo anterior.",
            "Las imagenes temporales usadas para construir un Word deben eliminarse cuando ya esten embebidas y no sean necesarias como fuente.",
            "No se deben subir capturas con datos sensibles reales sin aprobacion.",
            "Las capturas de documentacion deben provenir de ambiente local o de pruebas con datos controlados.",
        ],
    )

    doc.add_heading("8. Separacion cliente/desarrollador", 1)
    add_matrix(
        doc,
        ["Tipo", "Debe contener", "No debe contener"],
        [
            ("Cliente", "Flujo funcional, pantallas, reglas operativas, estados y acciones permitidas.", "Detalles internos innecesarios, codigo fuente o implementacion sensible."),
            ("Desarrollador", "Rutas, endpoints, tablas, servicios, validaciones, auditoria, permisos y pendientes tecnicos.", "Explicaciones comerciales extensas o contenido duplicado del cliente sin valor tecnico."),
            ("Tecnico transversal", "Politicas, guias de base de datos, auditoria, seguridad y arquitectura.", "Borradores sin control de version o instrucciones no aprobadas."),
        ],
        [1800, 3930, 3630],
    )

    doc.add_heading("9. Checklist previo a commit o subida", 1)
    add_bullets(
        doc,
        [
            "El documento final fue revisado en Word.",
            "La version final fue aprobada o esta marcada claramente como pendiente de aprobacion.",
            "No quedan borradores ni duplicados en la carpeta.",
            "No quedan imagenes temporales ni logs de generacion.",
            "Los archivos generados fueron revisados antes de conservarse.",
            "El nombre del archivo identifica modulo, audiencia y tipo.",
            "La carpeta docs conserva la estructura vigente.",
            "Cliente y desarrollador estan separados cuando corresponde.",
            "No se mezclan scripts experimentales con scripts aprobados.",
            "No se sube informacion sensible sin validacion.",
        ],
    )

    doc.add_heading("10. Criterios de rechazo", 1)
    add_callout(
        doc,
        "No se aprueba para subir",
        "Cualquier entrega que contenga borradores, duplicados, capturas temporales, documentos sin revisar, archivos generados sin validacion o mezcla de audiencia cliente/desarrollador.",
        WARN_FILL,
    )

    doc.add_heading("11. Pendientes controlados", 1)
    add_bullets(
        doc,
        [
            "Cuando se cree un nuevo modulo, registrar su documento cliente y desarrollador en la carpeta correspondiente.",
            "Cuando se generen nuevas capturas, confirmar si deben quedar embebidas, archivadas como fuente o eliminadas.",
            "Cuando se agreguen scripts de base de datos, clasificarlos como aprobados, utilitarios o experimentales.",
        ],
    )

    add_footer(doc)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
