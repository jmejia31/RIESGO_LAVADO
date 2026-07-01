import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"

TARGETS = [
    "Politica",
    "Repositorio",
    "Gestion",
    "Riesgos",
    "Modulo",
    "Modulos",
    "Documentacion",
    "Version",
    "Tecnica",
    "Tecnico",
    "Tecnicos",
    "Configuracion",
    "Auditoria",
    "Evidencia",
    "Evidencias",
    "Seguimiento",
    "Seguimientos",
    "Ejecucion",
    "Aprobacion",
    "Revision",
    "Validacion",
    "Informacion",
    "Operacion",
    "Eliminacion",
    "Accion",
    "Acciones",
    "Critica",
    "Critico",
    "Fisico",
    "Fisica",
    "Logica",
    "Logico",
    "Autorizacion",
    "Exportacion",
    "Generacion",
    "Instalacion",
    "Actualizacion",
    "Ubicacion",
    "Clasificacion",
    "Calificacion",
    "Parametrizacion",
    "Contraseña",
    "contrasena",
    "sesion",
    "Sesion",
    "pagina",
    "Pagina",
    "Indice",
    "analisis",
    "Analisis",
    "codigo",
    "Codigo",
    "Practica",
    "practica",
    "Ambito",
    "ambito",
    "proposito",
    "Proposito",
    "unica",
    "Unica",
    "unico",
    "Unico",
]


def iter_text(doc):
    for p in doc.paragraphs:
        if p.text:
            yield p.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        yield p.text
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                if p.text:
                    yield p.text


def main():
    counts = Counter()
    docs = defaultdict(set)
    patterns = {w: re.compile(rf"\b{re.escape(w)}\b") for w in TARGETS}
    for path in sorted(DOCS.rglob("*.docx")):
        try:
            doc = Document(path)
        except Exception as exc:
            print(f"ERROR\t{path}\t{exc}")
            continue
        text = "\n".join(iter_text(doc))
        for word, pattern in patterns.items():
            matches = len(pattern.findall(text))
            if matches:
                counts[word] += matches
                docs[word].add(str(path.relative_to(ROOT)))

    for word, count in counts.most_common():
        print(f"{word}\t{count}\t{len(docs[word])} docs")
        for doc_path in sorted(docs[word])[:8]:
            print(f"  - {doc_path}")


if __name__ == "__main__":
    main()
