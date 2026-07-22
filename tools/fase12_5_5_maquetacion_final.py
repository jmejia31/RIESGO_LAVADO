from __future__ import annotations

import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "docs/3. Módulo Matrices de Riesgos/Fase 12 - Mejora ejecutiva UXUI y mapa de calor"
DOCX = BASE / "Cierre_Tecnico_Fase_12_Matrices_Riesgos_SGRLA_IHSS.docx"
EVIDENCE = BASE / "Evidencia_Fase_12_5_5/fase12_5_5_cierre_definitivo.json"
XLSX_SOURCE = ROOT / "backend/RL.API/Infrastructure/Reporting/InstitutionalXlsxWorkbook.cs"


def replace_once(path: Path, old: str, new: str, label: str) -> None:
    text = path.read_text(encoding="utf-8")
    count = text.count(old)
    if count != 1:
        raise RuntimeError(f"{label}: se esperaba una coincidencia y se encontraron {count}")
    path.write_text(text.replace(old, new, 1), encoding="utf-8")


def patch_xlsx_widths() -> None:
    old = '''    private static decimal[] CalculateWidths(InstitutionalXlsxSheet sheet)
    {
        var widths = new decimal[sheet.Headers.Count];
        for (var column = 0; column < widths.Length; column++)
        {
            var maximum = sheet.Headers[column].Length;
            foreach (var row in sheet.Rows)
            {
                if (column < row.Count)
                    maximum = Math.Max(maximum, Convert.ToString(row[column], CultureInfo.InvariantCulture)?.Length ?? 0);
            }
            widths[column] = Math.Clamp(maximum + 2, 10, 48);
        }
        return widths;
    }
'''
    new = '''    private static decimal[] CalculateWidths(InstitutionalXlsxSheet sheet)
    {
        // Anchos institucionales para columnas recurrentes. El objetivo es evitar
        // cortes de palabras en tipos, estados y niveles cuando la hoja se ajusta
        // a una página de ancho, sin sobredimensionar las columnas descriptivas.
        var preferredWidths = new Dictionary<string, decimal>(StringComparer.OrdinalIgnoreCase)
        {
            ["ID"] = 8m,
            ["Sujeto"] = 30m,
            ["Documento"] = 18m,
            ["Tipo"] = 15m,
            ["Estado"] = 15m,
            ["Puntaje inherente"] = 14m,
            ["Nivel inherente"] = 14m,
            ["Puntaje residual"] = 14m,
            ["Nivel residual"] = 14m,
            ["Plan requerido"] = 14m,
            ["Fecha"] = 12m,
            ["Código"] = 10m,
            ["Factor"] = 16m,
            ["Matrices"] = 10m,
            ["Promedio inherente"] = 17m,
            ["Promedio residual"] = 17m,
            ["Alto / Crítico"] = 14m,
            ["Total"] = 11m,
            ["Vencidos"] = 12m
        };

        var widths = new decimal[sheet.Headers.Count];
        for (var column = 0; column < widths.Length; column++)
        {
            var header = sheet.Headers[column];
            if (preferredWidths.TryGetValue(header, out var preferred))
            {
                widths[column] = preferred;
                continue;
            }

            var maximum = header.Length;
            foreach (var row in sheet.Rows)
            {
                if (column < row.Count)
                    maximum = Math.Max(maximum, Convert.ToString(row[column], CultureInfo.InvariantCulture)?.Length ?? 0);
            }
            widths[column] = Math.Clamp(maximum + 2, 10, 34);
        }

        // Las hojas pequeñas deben conservar una anchura visual suficiente para
        // que títulos y valores no se compriman en una columna excesivamente estrecha.
        if (widths.Length > 0 && widths.Length <= 4)
        {
            const decimal minimumTotalWidth = 60m;
            var currentTotal = widths.Sum();
            if (currentTotal < minimumTotalWidth)
            {
                var extraPerColumn = (minimumTotalWidth - currentTotal) / widths.Length;
                for (var column = 0; column < widths.Length; column++)
                    widths[column] += extraPerColumn;
            }
        }

        return widths;
    }
'''
    replace_once(XLSX_SOURCE, old, new, "Ajuste final de anchos XLSX")


def remove_extra_paragraphs(cell) -> None:
    paragraphs = list(cell.paragraphs)
    for paragraph in paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)


def set_cell_preserving_first_run(cell, text: str) -> None:
    remove_extra_paragraphs(cell)
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_paragraph(document: Document, old: str, new: str) -> int:
    changed = 0
    for paragraph in document.paragraphs:
        if old not in paragraph.text:
            continue
        replacement = paragraph.text.replace(old, new)
        if paragraph.runs:
            paragraph.runs[0].text = replacement
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(replacement)
        changed += 1
    return changed


def patch_document() -> None:
    document = Document(DOCX)
    if len(document.tables) < 29:
        raise RuntimeError("Estructura inesperada del Documento Maestro.")

    # Eliminar párrafos heredados que quedaron visibles debajo de los textos finales.
    for table_index in (0, 2, 4, 24):
        remove_extra_paragraphs(document.tables[table_index].cell(0, 0))

    set_cell_preserving_first_run(
        document.tables[1].cell(0, 1),
        "Cierre Técnico Fase 12 y Actualización 12.5 - Versión Única para Aprobación",
    )
    set_cell_preserving_first_run(document.tables[21].cell(10, 0), "PR técnicos temporales")
    set_cell_preserving_first_run(document.tables[21].cell(10, 1), "#12, #13 y #14")
    set_cell_preserving_first_run(
        document.tables[21].cell(10, 2),
        "Ejecución, control visual y maquetación final; cerrados sin merge a main.",
    )

    old_approval = (
        "Complete esta sección para autorizar o rechazar el cierre formal de la Fase 12. "
        "La aprobación constituye el requisito previo para integrar el PR #2 a main."
    )
    new_approval = (
        "Complete esta sección para autorizar o rechazar el cierre formal de la Fase 12 y su actualización 12.5. "
        "La aprobación constituye el requisito previo para integrar el PR #2 a main."
    )
    if replace_paragraph(document, old_approval, new_approval) != 1:
        raise RuntimeError("No se pudo actualizar la instrucción de aprobación formal.")

    document.core_properties.modified = datetime.now()
    document.save(DOCX)


def update_evidence(run_id: int) -> None:
    data = json.loads(EVIDENCE.read_text(encoding="utf-8"))
    data["estado"] = "maquetacion_final_generada_pendiente_revision_visual_manual"
    data["maquetacion_final"] = {
        "run_id": run_id,
        "anchos_xlsx_institucionales": True,
        "palabras_categoricas_sin_cortes": "pendiente_verificacion_render",
        "bloques_documentales_sin_redundancia": True,
        "revision_visual_manual": "pendiente_revision_asistente",
    }
    data["siguiente_paso"] = "Revisión visual manual definitiva y cierre de evidencia"
    EVIDENCE.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Uso: fase12_5_5_maquetacion_final.py <run_id>")
    patch_xlsx_widths()
    patch_document()
    update_evidence(int(sys.argv[1]))
    print(json.dumps({
        "xlsx_widths": "institutional",
        "docx_redundancies_removed": True,
        "run_id": int(sys.argv[1]),
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
