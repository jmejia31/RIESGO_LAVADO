from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_ROOT = ROOT / "docs"
MATRICES_DIR = next(DOCS_ROOT.rglob("3. Módulo Matrices de Riesgos"))
OUT_DIR = MATRICES_DIR / "Fase 4 - Diseño funcional y experiencia de usuario"
OUT_FILE = OUT_DIR / "Fase_4_Diseno_Funcional_Experiencia_Usuario_Matrices_Riesgos_SGRLA_IHSS.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "E2F0D9"
LIGHT_GOLD = "FFF2CC"
LIGHT_RED = "FCE4D6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_run_font(run, size=10.5, color=INK, bold=False, italic=False, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, before, after, color in [
        (1, 16, 16, 8, BLUE),
        (2, 13, 12, 6, BLUE),
        (3, 12, 8, 4, DARK_BLUE),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(doc: Document) -> None:
    header_p = doc.sections[0].header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("IHSS - SGRLA/FT | Módulo Matrices de Riesgos | Fase 4")
    set_run_font(run, size=8, color=MUTED, bold=True)

    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Documento funcional aprobado y cerrado como base para Fase 5")
    set_run_font(run, size=8, color=MUTED)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[int], header_fill: str = LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, size=9.5, bold=True)
    set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=9.5)
            if len(value) <= 12:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def add_checklist(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("DOCUMENTO FUNCIONAL DE FASE 4")
    set_run_font(run, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Fase 4. Diseño funcional y experiencia de usuario")
    set_run_font(run, size=22, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Módulo Matrices de Riesgos - Sistema de Gestión de Riesgos LA/FT IHSS")
    set_run_font(run, size=13, color=MUTED)

    meta = [
        ("Proyecto", "RIESGO_LAVADO - IHSS"),
        ("Módulo", "Matrices de Riesgos"),
        ("Fase", "Fase 4. Diseño funcional y experiencia de usuario"),
        ("Versión", "1.3"),
        ("Fecha", date.today().strftime("%d/%m/%Y")),
        ("Estado", "Fase 4 aprobada y cerrada con permisos por módulo"),
        ("Responsable", "Javier Mejía"),
        ("Fuente", "Plan de fases, análisis final maestro, Fase 1 aprobada, Fase 2 aprobada y Fase 3 ejecutada/validada"),
        ("Ubicación", "docs/3. Módulo Matrices de Riesgos/Fase 4 - Diseño funcional y experiencia de usuario"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [2450, 6910])
    for row, (label, value) in zip(table.rows, meta):
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_title_block(doc)

    doc.add_heading("0. Control de cambios", level=1)
    add_table(
        doc,
        ["Versión", "Fecha", "Descripción del cambio", "Responsable", "Estado"],
        [
            ("1.0", date.today().strftime("%d/%m/%Y"), "Creación de documento funcional de Fase 4: pantallas, flujos, roles, permisos, estados, navegación, validaciones, mensajes, wireframes textuales y criterios de aceptación.", "Javier Mejía", "Revisión funcional"),
            ("1.1", date.today().strftime("%d/%m/%Y"), "Alineación con Fase 3: el estado EN_EVALUACION queda respaldado por la restricción física CK_RL_MR_MAT_ESTADO mediante script incremental DBA 05.", "Javier Mejía", "Alineada para aprobación"),
            ("1.2", date.today().strftime("%d/%m/%Y"), "Aprobación formal y cierre documental de Fase 4 como base funcional válida para iniciar Fase 5.", "Javier Mejía", "Aprobada y cerrada"),
            ("1.3", date.today().strftime("%d/%m/%Y"), "Aclaración de regla institucional: los permisos se mantienen por módulo, conforme al esquema inicial del sistema; no se implementarán permisos por acción.", "Javier Mejía", "Aclaración cerrada"),
        ],
        [900, 1350, 5000, 1450, 660],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("1. Propósito de la fase", level=1)
    doc.add_paragraph(
        "La Fase 4 define cómo deberá operar el módulo Matrices de Riesgos desde la experiencia de usuario y el flujo funcional antes de construir pantallas, endpoints definitivos o lógica de cálculo. "
        "Este documento convierte la metodología aprobada y el modelo Oracle ejecutado en una especificación funcional aprobada y cerrada para continuar con contratos backend, seguridad y auditoría en Fase 5."
    )

    doc.add_heading("2. Dependencias y reglas heredadas", level=1)
    add_table(
        doc,
        ["Fuente", "Decisión heredada para Fase 4"],
        [
            ("Fase 1", "Gobierno documental, responsable Javier Mejía, documentación en Word y control de aprobación antes de subir cambios finales."),
            ("Fase 2", "Metodología LA/FT aprobada: Proveedores 50%, Clientes/Patronos 25% y Empleados 25%; variables internas totalizan 100% por factor."),
            ("Fase 3", "Modelo Oracle RL_MR_* ejecutado y validado; módulo registrado como /matrices-riesgos con MOD_ID 10; CK_RL_MR_MAT_ESTADO alineado con EN_EVALUACION."),
            ("Análisis maestro", "Backend calcula, frontend presenta; DNP queda como integración futura obligatoria sujeta a contrato técnico e institucional."),
            ("Base de datos", "El esquema RIESGO_LAVADO quedó con comentarios completos y sin codificación dañada en metadatos."),
            ("Sistema actual", "Angular usa rutas protegidas con moduloGuard; backend usa ModuloAuthorize; menú lateral se alimenta de RL_MODULOS."),
        ],
        [2500, 6860],
        header_fill=LIGHT_GOLD,
    )

    doc.add_heading("3. Alcance y fuera de alcance", level=1)
    add_table(
        doc,
        ["Tipo", "Definición"],
        [
            ("Incluye", "Mapa de pantallas, flujo funcional por rol, estados de matriz, acceso por módulo, rutas frontend, validaciones visibles, mensajes, wireframes textuales y criterios de aceptación."),
            ("Incluye", "Diseño funcional de submódulos: dashboard, listado, creación, evaluación, controles, planes de acción, evidencias, historial, reportes, parametrización metodológica y bandeja futura DNP."),
            ("No incluye", "Programación de componentes Angular, endpoints definitivos, DTOs finales, repositorios Oracle ni motor de cálculo productivo."),
            ("No incluye", "Cambio de ponderaciones institucionales 50/25/25 ni modificación retroactiva de matrices cerradas."),
            ("No incluye", "Escritura directa hacia DNP sin contrato técnico, seguridad, autorización institucional y disponibilidad de interfaz."),
        ],
        [1800, 7560],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("4. Perfiles funcionales", level=1)
    add_table(
        doc,
        ["Perfil", "Responsabilidad", "Acceso esperado"],
        [
            ("Operativo", "Captura matrices, completa variables, adjunta evidencias y registra controles preliminares.", "Crear, editar borrador, guardar evaluación y consultar sus matrices permitidas."),
            ("Revisor", "Valida integridad, observa inconsistencias y solicita correcciones antes de aprobación.", "Consultar, observar, devolver a borrador controlado y enviar a aprobación."),
            ("Aprobador", "Aprueba resultados, autoriza cierre y valida planes de acción cuando el residual lo exige.", "Revisar, aprobar, cerrar y rechazar cierre incompleto."),
            ("Administrador metodológico", "Administra modelo vigente, variables, escalas y criterios sin afectar matrices cerradas.", "Consultar y proponer cambios metodológicos versionados."),
            ("Consulta / Auditoría", "Consulta matrices, historial, evidencias y reportes sin alterar información.", "Solo lectura, exportación autorizada y trazabilidad de accesos sensibles."),
        ],
        [1900, 4300, 3160],
    )

    doc.add_heading("5. Submódulos funcionales", level=1)
    add_table(
        doc,
        ["Submódulo", "Objetivo", "Resultado esperado"],
        [
            ("Dashboard", "Presentar estado general de matrices, distribución por riesgo, pendientes y vencimientos.", "Usuario entiende carga operativa, riesgos críticos y acciones pendientes."),
            ("Listado de matrices", "Consultar, filtrar y abrir matrices por estado, factor, sujeto, fechas y nivel de riesgo.", "Entrada principal para operación y seguimiento."),
            ("Creación de matriz", "Registrar sujeto o alcance, tipo de matriz, origen de datos y modelo metodológico vigente.", "Matriz creada en BORRADOR con snapshot inicial controlado."),
            ("Evaluación", "Capturar variables, probabilidad, impacto, justificación y soporte por factor.", "Datos suficientes para cálculo backend de riesgo inherente."),
            ("Controles", "Registrar mitigadores, solidez, responsable, evidencia y efectividad.", "Mitigación trazable para cálculo residual."),
            ("Planes de acción", "Gestionar actividades, responsables, fechas, medios de prueba y estados.", "Plan obligatorio cuando residual sea alto o crítico."),
            ("Evidencias", "Adjuntar y consultar documentos asociados a matriz, control o plan.", "Soporte documental protegido y auditable."),
            ("Historial", "Mostrar eventos funcionales y cambios de estado.", "Trazabilidad clara para revisión y auditoría."),
            ("Reportes", "Consultar mapa de calor, solidez de controles, matrices por estado y exportaciones.", "Información ejecutiva y operativa sin depender de hojas externas."),
            ("Parametrización", "Administrar metodología vigente y versiones futuras con aprobación.", "Cambios metodológicos controlados sin afectar matrices cerradas."),
            ("Integración DNP", "Preparar calificación de patrono para envío futuro cuando exista contrato técnico.", "Bandeja de pendientes, errores y estado de integración sin escritura directa no autorizada."),
        ],
        [2050, 3650, 3660],
    )

    doc.add_heading("6. Mapa de navegación propuesto", level=1)
    add_table(
        doc,
        ["Ruta frontend", "Pantalla", "Guard", "Notas"],
        [
            ("/matrices-riesgos", "Dashboard del módulo", "moduloGuard(10)", "Ruta principal desde menú Riesgos LA/FT."),
            ("/matrices-riesgos/listado", "Listado de matrices", "moduloGuard(10)", "Filtros por estado, factor, nivel, sujeto, fechas y responsable."),
            ("/matrices-riesgos/nueva", "Crear matriz", "moduloGuard(10)", "Selecciona tipo de matriz y modelo vigente."),
            ("/matrices-riesgos/:id/evaluacion", "Evaluación de matriz", "moduloGuard(10)", "Captura variables por factor y justificaciones."),
            ("/matrices-riesgos/:id/controles", "Controles mitigantes", "moduloGuard(10)", "Solidez, evidencia, responsable y efectividad."),
            ("/matrices-riesgos/:id/planes", "Planes de acción", "moduloGuard(10)", "Requerido para residual alto o crítico."),
            ("/matrices-riesgos/:id/historial", "Historial funcional", "moduloGuard(10)", "Eventos y cambios de estado."),
            ("/matrices-riesgos/reportes", "Reportes y mapa de calor", "moduloGuard(10)", "Incluye solidez de controles y exportaciones."),
            ("/matrices-riesgos/metodologia", "Parametrización metodológica", "moduloGuard(10)", "Restringido a administrador metodológico."),
            ("/matrices-riesgos/integracion-dnp", "Bandeja de integración DNP", "moduloGuard(10)", "Futura, sujeta a contrato técnico."),
        ],
        [2450, 2400, 1800, 2710],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("7. Flujo funcional principal", level=1)
    add_table(
        doc,
        ["Paso", "Acción", "Regla funcional"],
        [
            ("1", "Ingresar al módulo /matrices-riesgos.", "El usuario debe estar autenticado y tener MOD_ID 10."),
            ("2", "Crear o seleccionar matriz.", "Debe usar modelo metodológico vigente y no inactivo."),
            ("3", "Definir sujeto o alcance.", "Permitidos: proveedor, cliente/patrono, empleado, proceso, caso positivo o matriz institucional."),
            ("4", "Capturar variables por factor.", "Cada factor conserva ponderación institucional y variables internas aprobadas."),
            ("5", "Guardar evaluación.", "El sistema permite borrador incompleto, pero no cálculo final incompleto."),
            ("6", "Calcular riesgo inherente.", "El cálculo lo ejecuta backend y registra snapshot."),
            ("7", "Registrar controles.", "Cada control debe tener descripción, responsable, solidez y evidencia si aplica."),
            ("8", "Calcular riesgo residual.", "Backend aplica mitigación aprobada y conserva resultado vigente."),
            ("9", "Gestionar plan de acción.", "Obligatorio si residual queda alto o crítico."),
            ("10", "Enviar a revisión.", "Debe existir evaluación completa, resultado calculado y evidencias requeridas."),
            ("11", "Revisar, observar o aprobar.", "Toda observación exige motivo visible y queda en historial."),
            ("12", "Cerrar matriz.", "Cierre bloquea edición retroactiva de metodología, pesos y resultados."),
        ],
        [700, 3300, 5360],
    )

    doc.add_heading("8. Estados de matriz y acciones permitidas", level=1)
    doc.add_paragraph(
        "El estado EN_EVALUACION se conserva por necesidad funcional y queda alineado con el modelo físico Oracle de Fase 3. "
        "La restricción CK_RL_MR_MAT_ESTADO permite este estado para diferenciar matrices creadas de matrices con captura activa de variables, evidencias y justificaciones."
    )
    add_table(
        doc,
        ["Estado", "Descripción", "Acciones permitidas"],
        [
            ("BORRADOR", "Matriz creada, aún editable y sin cálculo definitivo.", "Editar, guardar, adjuntar evidencia, cancelar con motivo."),
            ("EN_EVALUACION", "Variables y evidencias en proceso de captura.", "Editar evaluación, guardar avance, calcular preliminar."),
            ("CALCULADA", "Backend generó inherente, mitigación y residual.", "Revisar resultados, agregar controles, generar plan si aplica."),
            ("EN_REVISION", "Matriz enviada para revisión funcional.", "Observar, aprobar o devolver con motivo."),
            ("OBSERVADA", "Matriz devuelta por inconsistencias o falta de soporte.", "Corregir observaciones y reenviar."),
            ("APROBADA", "Resultado aprobado por perfil autorizado.", "Cerrar, exportar, consultar historial."),
            ("CERRADA", "Matriz finalizada y protegida contra cambios retroactivos.", "Consultar, exportar, ver evidencia e historial."),
            ("INACTIVA", "Matriz anulada lógicamente con motivo.", "Consultar trazabilidad; no editar ni recalcular."),
        ],
        [1650, 3600, 4110],
        header_fill=LIGHT_GREEN,
    )

    doc.add_heading("9. Acceso por módulo y roles funcionales", level=1)
    doc.add_paragraph(
        "La regla institucional aprobada mantiene el esquema inicial del sistema: los permisos se administran por módulo mediante RL_MODULOS y RL_USUARIO_MODULOS. "
        "No se implementarán permisos finos por acción. Las acciones listadas en esta sección describen el flujo funcional esperado por perfil operativo, pero no crean un nuevo modelo de autorización por botón, operación o acción individual."
    )
    add_table(
        doc,
        ["Actividad funcional", "Operativo", "Revisor", "Aprobador", "Admin. metodológico", "Consulta"],
        [
            ("Consultar dashboard", "Sí", "Sí", "Sí", "Sí", "Sí"),
            ("Crear matriz", "Sí", "Opcional", "No", "No", "No"),
            ("Editar borrador/evaluación", "Sí", "No", "No", "No", "No"),
            ("Calcular", "Sí", "Sí", "No", "No", "No"),
            ("Observar matriz", "No", "Sí", "Sí", "No", "No"),
            ("Aprobar matriz", "No", "No", "Sí", "No", "No"),
            ("Cerrar matriz", "No", "No", "Sí", "No", "No"),
            ("Administrar metodología", "No", "No", "No", "Sí", "No"),
            ("Exportar reportes", "Opcional", "Sí", "Sí", "Sí", "Opcional"),
            ("Inactivar matriz", "No", "No", "Sí", "No", "No"),
        ],
        [2300, 1100, 1100, 1100, 1600, 2160],
    )

    doc.add_heading("10. Wireframes textuales aprobados", level=1)
    add_table(
        doc,
        ["Pantalla", "Estructura funcional esperada"],
        [
            ("Dashboard", "Fila superior con tarjetas: matrices abiertas, críticas, vencidas y cerradas. Zona central con mapa de calor. Panel lateral con pendientes por rol. Accesos rápidos a nueva matriz, listado y reportes."),
            ("Listado", "Barra de filtros por fecha, estado, factor, sujeto, nivel, responsable y origen. Tabla con estado visual, riesgo inherente, residual, última acción y botones contextuales."),
            ("Nueva matriz", "Formulario en pasos: tipo de matriz, sujeto/alcance, modelo vigente, origen de datos, responsable y confirmación. Validación visible antes de crear."),
            ("Evaluación", "Tabs por factor: Proveedores, Clientes/Patronos y Empleados. Cada tab muestra variables, peso interno, valor, justificación, soporte y estado de completitud."),
            ("Controles", "Lista de controles por factor con solidez, periodicidad, responsable, evidencia, mitigación sugerida y observaciones."),
            ("Planes de acción", "Tabla editable por actividad, responsable, fecha inicio, fecha fin, medio de prueba, estado y observaciones."),
            ("Detalle/cierre", "Resumen de sujeto, metodología, resultados, controles, planes, evidencias e historial. Botones según estado y permiso."),
            ("Reportes", "Mapa de calor, solidez de controles, matrices por estado, promedio residual, distribución por factor, exportaciones Excel/PDF auditadas."),
            ("Metodología", "Vista de solo administración: modelos, factores, variables, escalas, vigencia, estado y control de versiones."),
        ],
        [2100, 7260],
        header_fill=LIGHT_GOLD,
    )

    doc.add_heading("11. Validaciones visibles y mensajes", level=1)
    add_table(
        doc,
        ["Situación", "Validación", "Mensaje funcional sugerido"],
        [
            ("Sin acceso al módulo", "Usuario no tiene MOD_ID 10 asignado.", "No tiene acceso al módulo Matrices de Riesgos."),
            ("Modelo no vigente", "No existe modelo APROBADO activo.", "No existe metodología vigente para crear matrices. Solicite revisión metodológica."),
            ("Variables incompletas", "Faltan valores obligatorios por factor.", "Complete todas las variables obligatorias antes de calcular."),
            ("Peso inválido", "La suma interna por factor no totaliza 100%.", "La ponderación interna del factor no es válida. Revise la metodología."),
            ("Cálculo incompleto", "No hay datos suficientes para cálculo backend.", "No se puede calcular la matriz porque existen datos pendientes."),
            ("Residual alto/crítico", "Resultado exige plan de acción.", "Debe registrar un plan de acción antes de enviar a aprobación."),
            ("Cierre inválido", "Matriz sin aprobación o con plan pendiente.", "La matriz no puede cerrarse mientras existan acciones obligatorias pendientes."),
            ("Observación", "Revisor devuelve matriz.", "Ingrese el motivo de la observación para continuar."),
            ("Exportación", "Usuario genera reporte.", "La exportación será registrada en auditoría."),
        ],
        [2100, 3300, 3960],
    )

    doc.add_heading("12. Auditoría y trazabilidad funcional", level=1)
    add_checklist(
        doc,
        [
            "Crear, editar, calcular, recalcular, observar, aprobar, cerrar, inactivar y exportar debe registrar auditoría transversal.",
            "Visualización o descarga de evidencias debe quedar auditada.",
            "Cambios de estado deben conservar motivo obligatorio cuando aplique.",
            "RL_MR_HISTORIAL debe complementar RL_AUDITORIA con estado anterior, estado nuevo, datos anteriores, datos nuevos, usuario, IP y fecha.",
            "Toda exportación Excel/PDF debe registrar usuario, filtros, fecha, formato y módulo.",
            "El cálculo debe guardar snapshot de modelo, factores, variables, pesos, escalas, controles y resultado.",
        ],
    )

    doc.add_heading("13. Reportería mínima", level=1)
    add_table(
        doc,
        ["Reporte", "Contenido mínimo", "Filtros"],
        [
            ("Mapa de calor", "Distribución por probabilidad, impacto, inherente y residual.", "Fecha, factor, sujeto, estado, modelo."),
            ("Matrices por estado", "Cantidad y detalle por BORRADOR, EN_EVALUACION, CALCULADA, EN_REVISION, OBSERVADA, APROBADA y CERRADA.", "Fecha, responsable, estado."),
            ("Riesgo residual", "Promedio residual, nivel y variación por factor y sujeto.", "Factor, nivel, periodo."),
            ("Solidez de controles", "Cantidad de controles, nivel de efectividad, calidad, factor asociado y distribución por nivel.", "Factor, matriz, periodo, responsable."),
            ("Planes de acción", "Pendientes, vencidos, cerrados, responsable y medio de prueba.", "Estado, fecha, responsable, riesgo."),
            ("DNP futuro", "Calificaciones por patrono preparadas, pendientes, enviadas o con error.", "Estado integración, fecha, patrono."),
            ("Auditoría del módulo", "Eventos sensibles por usuario, acción, IP, fecha y matriz.", "Usuario, acción, fecha, estado."),
        ],
        [2000, 4700, 2660],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("14. Integración futura DNP", level=1)
    doc.add_paragraph(
        "La calificación de riesgo por patrono deberá contemplarse como integración obligatoria futura hacia DNP, sujeta a contrato técnico, seguridad, autorización institucional y disponibilidad de interfaz. "
        "Fase 4 no autoriza escritura directa hacia DNP; solo define la experiencia funcional de bandeja, consulta de estado, errores y reintentos controlados."
    )
    add_table(
        doc,
        ["Elemento", "Diseño funcional"],
        [
            ("Bandeja", "Mostrar matriz, número de patrono, calificación, residual, estado de envío, fecha y respuesta técnica."),
            ("Estados", "PENDIENTE, ENVIADO, ERROR y ANULADO."),
            ("Acciones", "Consultar, reintentar cuando sea autorizado, anular con motivo y exportar evidencia."),
            ("Seguridad", "Solo usuarios autorizados; toda acción auditada."),
            ("Condición", "No se implementa envío real hasta Fase técnica aprobada con DNP."),
        ],
        [2500, 6860],
        header_fill=LIGHT_GOLD,
    )

    doc.add_heading("15. Componentes futuros sugeridos", level=1)
    add_table(
        doc,
        ["Capa", "Elemento sugerido", "Uso"],
        [
            ("Angular", "matrices-riesgos/dashboard", "Vista principal del módulo."),
            ("Angular", "matrices-riesgos/listado", "Consulta y filtros de matrices."),
            ("Angular", "matrices-riesgos/formulario", "Creación y edición controlada."),
            ("Angular", "matrices-riesgos/evaluacion", "Captura de variables por factor."),
            ("Angular", "matrices-riesgos/reportes", "Mapa de calor y reportería."),
            ("Backend", "MatricesRiesgosController", "Entrada HTTP en Fase 5."),
            ("Backend", "MatricesRiesgoService", "Reglas funcionales y cálculo en backend."),
            ("Backend", "MatricesRiesgosRepository", "Persistencia Oracle RL_MR_*."),
            ("Oracle", "RL_MR_*", "Persistencia ejecutada y validada en Fase 3."),
        ],
        [1600, 3200, 4560],
    )

    doc.add_heading("16. Criterios de aceptación de Fase 4", level=1)
    add_table(
        doc,
        ["Criterio", "Estado esperado para cierre"],
        [
            ("Mapa de pantallas definido.", "Cumplido en este documento."),
            ("Flujo funcional por rol definido.", "Cumplido en perfiles funcionales y flujo principal, sin crear permisos por acción."),
            ("Estados y acciones permitidas definidos.", "Cumplido con matriz de estados y alineación física de EN_EVALUACION en Fase 3."),
            ("Rutas frontend propuestas.", "Cumplido con mapa de navegación."),
            ("Validaciones visibles y mensajes definidos.", "Cumplido con catálogo funcional de mensajes."),
            ("Reportería mínima definida.", "Cumplido, incluyendo solidez de controles."),
            ("Integración DNP tratada como futura obligatoria.", "Cumplido sin autorizar escritura directa."),
            ("Backend sigue siendo única fuente de cálculo.", "Cumplido como regla funcional heredada."),
            ("Listo para Fase 5.", "Cumplido: Fase 4 aprobada formalmente por Javier Mejía el 06/07/2026."),
        ],
        [4300, 5060],
        header_fill=LIGHT_GREEN,
    )

    doc.add_heading("17. Pendientes controlados y reglas cerradas", level=1)
    add_table(
        doc,
        ["Pendiente", "Motivo", "Momento de definición"],
        [
            ("Permisos por módulo", "Regla cerrada: se mantiene el esquema inicial por módulo; no se implementan permisos finos por acción.", "Vigente para Fase 5 y fases posteriores."),
            ("DTOs definitivos", "Deben derivarse del contrato REST y no del diseño visual.", "Fase 5."),
            ("Fórmulas finales de cálculo", "La metodología está aprobada, pero el motor debe implementarse con casos controlados.", "Fase 6."),
            ("Diseño visual final", "Fase 4 define estructura funcional; implementación visual se construye en Angular.", "Fase 8."),
            ("Integración real DNP", "Depende de contrato técnico institucional.", "Fase específica de integración."),
        ],
        [2600, 3900, 2860],
        header_fill=LIGHT_RED,
    )

    doc.add_heading("18. Cierre formal de Fase 4", level=1)
    add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ("Estado de Fase 4", "Aprobada y cerrada."),
            ("Responsable de aprobación", "Javier Mejía."),
            ("Fecha de aprobación", "06/07/2026."),
            ("Alcance aprobado", "Mapa de pantallas, flujo funcional, estados, acceso por módulo, rutas frontend, validaciones visibles, mensajes, reportería, DNP futuro y criterios de aceptación."),
            ("Base técnica", "Fase 3 versión 1.9, modelo Oracle validado y CK_RL_MR_MAT_ESTADO alineado con EN_EVALUACION."),
            ("Siguiente fase autorizada", "Fase 5. Contratos backend, seguridad y auditoría."),
            ("Condición de control", "No modificar el diseño aprobado sin registrar nueva versión documental y aprobación correspondiente."),
        ],
        [2900, 6460],
        header_fill=LIGHT_GREEN,
    )

    doc.add_heading("19. Decisión de cierre", level=1)
    doc.add_paragraph(
        "La Fase 4 queda aprobada y cerrada como documento funcional alineado con el modelo Oracle de Fase 3 y con la regla institucional de permisos por módulo. "
        "La versión 1.3 sirve como base formal para iniciar Fase 5: contratos backend, seguridad y auditoría. "
        "No se debe modificar el diseño aprobado ni iniciar cambios fuera de alcance sin registrar una nueva versión del documento y obtener aprobación correspondiente."
    )

    props = doc.core_properties
    props.title = "Fase 4 - Diseño funcional y experiencia de usuario - Matrices de Riesgos"
    props.subject = "Sistema de Gestión de Riesgos LA/FT IHSS"
    props.keywords = "IHSS, SGRLA, Matrices de Riesgos, Fase 4, diseño funcional"
    props.comments = "Documento funcional de Fase 4 aprobado y cerrado como base para Fase 5 del módulo Matrices de Riesgos, con permisos por módulo."
    props.author = "Javier Mejía"
    doc.save(OUT_FILE)


if __name__ == "__main__":
    build_document()
    print(OUT_FILE)
