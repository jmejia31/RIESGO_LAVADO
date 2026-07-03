from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_ROOT = ROOT / "docs"
MATRICES_DIR = next(DOCS_ROOT.rglob("3. Módulo Matrices de Riesgos"))
FASE_DIR = MATRICES_DIR / "Fase 3 - Modelo de datos y arquitectura Oracle"
EVIDENCE_DIR = FASE_DIR / "Evidencia_DBA"
LOG_FILE = EVIDENCE_DIR / "05_F3_align_estado_en_evaluacion_matrices_20260703.log"
OUT_FILE = EVIDENCE_DIR / "Evidencia_DBA_Alineacion_Estado_EN_EVALUACION_Fase_3_Matrices_Riesgos_SGRLA_IHSS.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "E2F0D9"


def set_run_font(run, size=10.5, color=INK, bold=False, italic=False, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, before, after, color in [
        (1, 16, 16, 8, BLUE),
        (2, 13, 12, 6, BLUE),
        (3, 12, 8, 4, DARK_BLUE),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(doc: Document) -> None:
    header_p = doc.sections[0].header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("IHSS - SGRLA/FT | Evidencia DBA Fase 3")
    set_run_font(run, size=8, color=MUTED, bold=True)

    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Documento de evidencia técnica - Javier Mejía")
    set_run_font(run, size=8, color=MUTED)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[int], header_fill: str = LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=9.5)
            if len(value) <= 12:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("EVIDENCIA DBA DE ALINEACIÓN CONTROLADA")
    set_run_font(run, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Estado EN_EVALUACION en RL_MR_MATRICES")
    set_run_font(run, size=22, color=INK, bold=True)

    p = doc.add_paragraph()
    run = p.add_run(
        "Fase 3. Modelo de datos y arquitectura Oracle | Módulo 3. Matrices de Riesgos | "
        f"Fecha: {date.today().strftime('%d/%m/%Y')}"
    )
    set_run_font(run, size=10, color=MUTED)


def add_log_excerpt(doc: Document) -> None:
    doc.add_heading("5. Salida SQLPlus registrada", level=1)
    if not LOG_FILE.exists():
        text = "No se encontró el archivo de log esperado."
    else:
        text = LOG_FILE.read_text(encoding="cp1252", errors="replace").strip()

    allowed_controls = {"\n", "\r", "\t"}
    text = "".join(ch for ch in text if ch in allowed_controls or ord(ch) >= 32)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for line in text.splitlines():
        run = p.add_run(line)
        set_run_font(run, size=8.5, color=INK, name="Consolas")
        p.add_run("\n")


def build_doc() -> None:
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_title(doc)

    doc.add_heading("1. Propósito", level=1)
    doc.add_paragraph(
        "Documentar la corrección controlada realizada para alinear el estado funcional EN_EVALUACION de Fase 4 "
        "con la restricción física CK_RL_MR_MAT_ESTADO de la tabla RL_MR_MATRICES en el modelo Oracle de Fase 3."
    )

    doc.add_heading("2. Resumen ejecutivo", level=1)
    add_table(
        doc,
        ["Elemento", "Resultado"],
        [
            ("Base validada", "RIESGO_LAVADO."),
            ("Tabla", "RL_MR_MATRICES."),
            ("Restricción", "CK_RL_MR_MAT_ESTADO."),
            ("Estado incorporado", "EN_EVALUACION."),
            ("Tipo de acción", "Cambio estructural controlado sobre restricción CHECK; sin eliminación de datos."),
            ("Script ejecutado", "05_F3_align_estado_en_evaluacion_matrices.sql."),
            ("Responsable", "Javier Mejía."),
            ("Resultado", "Restricción alineada correctamente; EN_EVALUACION presente; sin objetos inválidos reportados."),
        ],
        [2500, 6860],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("3. Validaciones realizadas", level=1)
    add_table(
        doc,
        ["Validación", "Resultado"],
        [
            ("Estados existentes", "Se validó que no existieran estados incompatibles antes de recrear la restricción."),
            ("Restricción física", "Se recreó CK_RL_MR_MAT_ESTADO incorporando EN_EVALUACION."),
            ("Presencia del estado", "La consulta posterior devolvió SI para EN_EVALUACION_PRESENTE."),
            ("Datos productivos", "No se ejecutaron DROP, TRUNCATE ni DELETE de tablas o datos."),
            ("Objetos inválidos", "La validación posterior no reportó objetos inválidos."),
        ],
        [3000, 6360],
    )

    doc.add_heading("4. Conclusión DBA", level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        "La observación queda corregida. El flujo funcional de Fase 4 puede conservar el estado EN_EVALUACION, "
        "porque el modelo físico de Fase 3 ya lo permite en la restricción CK_RL_MR_MAT_ESTADO. "
        "Esta evidencia no autoriza cambios adicionales en producción sin aprobación institucional y protocolo DBA."
    )
    set_run_font(run, size=10.5)

    add_log_excerpt(doc)
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build_doc()
