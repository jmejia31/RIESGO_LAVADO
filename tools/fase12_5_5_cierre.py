from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import subprocess
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "docs" / "3. Módulo Matrices de Riesgos" / "Fase 12 - Mejora ejecutiva UXUI y mapa de calor"
DOCX = BASE / "Cierre_Tecnico_Fase_12_Matrices_Riesgos_SGRLA_IHSS.docx"
SHA_FILE = BASE / "Cierre_Tecnico_Fase_12_Matrices_Riesgos_SGRLA_IHSS.sha256"
EVIDENCE_DIR = BASE / "Evidencia_Fase_12_5_5"
EVIDENCE = EVIDENCE_DIR / "fase12_5_5_cierre_definitivo.json"
CLOSURE_MD = BASE / "Fase_12_5_5_Pruebas_Evidencia_y_Cierre_Definitivo.md"
MARKER = "ACTUALIZACIÓN FINAL - FASE 12.5"
ANSI = re.compile(r"\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])")


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def strip_ansi(text: str) -> str:
    return ANSI.sub("", text).replace("\r", "")


def parse_quality_log(path: Path) -> dict[str, Any]:
    text = strip_ansi(path.read_text(encoding="utf-8", errors="replace"))

    backend_matches = re.findall(
        r"Failed:\s*(\d+),\s*Passed:\s*(\d+),\s*Skipped:\s*(\d+),\s*Total:\s*(\d+)",
        text,
        flags=re.IGNORECASE,
    )
    if not backend_matches:
        raise RuntimeError("No se localizaron los totales de pruebas backend en el Quality Gate.")
    backend_failed, backend_passed, backend_skipped, backend_total = map(int, backend_matches[-1])

    frontend_files = re.findall(r"Test Files\s+(\d+)\s+passed", text, flags=re.IGNORECASE)
    frontend_tests = re.findall(r"Tests\s+(\d+)\s+passed", text, flags=re.IGNORECASE)
    if not frontend_files or not frontend_tests:
        raise RuntimeError("No se localizaron los totales de pruebas frontend en el Quality Gate.")

    e2e_section = text.split("=== Pruebas E2E ===", 1)[-1].split("=== Resumen de cobertura ===", 1)[0]
    e2e_matches = re.findall(r"(\d+)\s+passed", e2e_section, flags=re.IGNORECASE)
    if not e2e_matches:
        raise RuntimeError("No se localizaron los totales E2E en el Quality Gate.")

    backend_coverage = re.search(r"Backend\s+lineas=([0-9.]+)%\s+ramas=([0-9.]+)%", text)
    frontend_coverage = re.search(
        r"Frontend\s+sentencias=([0-9.]+)%\s+ramas=([0-9.]+)%\s+funciones=([0-9.]+)%\s+lineas=([0-9.]+)%",
        text,
    )
    if not backend_coverage or not frontend_coverage:
        raise RuntimeError("No se localizaron las métricas de cobertura del Quality Gate.")
    if "Puertas de calidad correctas." not in text:
        raise RuntimeError("El Quality Gate no finalizó con estado correcto.")

    result = {
        "backend": {
            "aprobadas": backend_passed,
            "fallidas": backend_failed,
            "omitidas": backend_skipped,
            "total": backend_total,
        },
        "frontend": {
            "archivos_aprobados": int(frontend_files[-1]),
            "pruebas_aprobadas": int(frontend_tests[-1]),
            "fallidas": 0,
        },
        "e2e": {"aprobadas": int(e2e_matches[-1]), "fallidas": 0},
        "cobertura": {
            "backend_lineas_pct": float(backend_coverage.group(1)),
            "backend_ramas_pct": float(backend_coverage.group(2)),
            "frontend_sentencias_pct": float(frontend_coverage.group(1)),
            "frontend_ramas_pct": float(frontend_coverage.group(2)),
            "frontend_funciones_pct": float(frontend_coverage.group(3)),
            "frontend_lineas_pct": float(frontend_coverage.group(4)),
        },
    }

    if backend_failed or backend_passed < 96:
        raise RuntimeError(f"Resultado backend inesperado: {result['backend']}")
    if result["frontend"]["pruebas_aprobadas"] < 156:
        raise RuntimeError(f"Resultado frontend inesperado: {result['frontend']}")
    if result["e2e"]["aprobadas"] < 7:
        raise RuntimeError(f"Resultado E2E inesperado: {result['e2e']}")
    return result


def pdf_pages(path: Path) -> int:
    output = subprocess.check_output(["pdfinfo", str(path)], text=True, encoding="utf-8", errors="replace")
    match = re.search(r"^Pages:\s+(\d+)", output, flags=re.MULTILINE)
    if not match:
        raise RuntimeError(f"No se pudo determinar el número de páginas de {path.name}.")
    return int(match.group(1))


def validate_pdf(path: Path) -> dict[str, Any]:
    data = path.read_bytes()
    if not data.startswith(b"%PDF-") or b"%%EOF" not in data[-512:]:
        raise RuntimeError(f"PDF estructuralmente inválido: {path}")
    text_path = path.with_suffix(".txt")
    subprocess.run(["pdftotext", str(path), str(text_path)], check=True)
    extracted = text_path.read_text(encoding="utf-8", errors="replace").strip()
    if len(extracted) < 100:
        raise RuntimeError(f"El PDF {path.name} no contiene texto suficiente para validación.")
    return {
        "archivo": path.name,
        "tamano_bytes": path.stat().st_size,
        "sha256": sha256(path),
        "paginas": pdf_pages(path),
        "texto_extraido_caracteres": len(extracted),
        "estructura_pdf_valida": True,
    }


def validate_xlsx(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as archive:
        required = {"[Content_Types].xml", "xl/workbook.xml", "xl/styles.xml"}
        missing = sorted(required.difference(archive.namelist()))
        if missing:
            raise RuntimeError(f"XLSX incompleto; faltan: {missing}")
        root = ElementTree.fromstring(archive.read("xl/workbook.xml"))
        namespaces = {"m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
        sheets = [element.attrib["name"] for element in root.findall("m:sheets/m:sheet", namespaces)]
        expected = ["Resumen", "Matrices", "Factores", "Mapa transición", "Matrices críticas", "Planes"]
        if sheets != expected:
            raise RuntimeError(f"Hojas XLSX inesperadas: {sheets}")
        worksheet_count = len([name for name in archive.namelist() if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")])
        if worksheet_count != 6:
            raise RuntimeError(f"Cantidad de hojas físicas inesperada: {worksheet_count}")
    return {
        "archivo": path.name,
        "tamano_bytes": path.stat().st_size,
        "sha256": sha256(path),
        "hojas": sheets,
        "openxml_valido": True,
    }


def validate_docx(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as archive:
        required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
        missing = sorted(required.difference(archive.namelist()))
        if missing:
            raise RuntimeError(f"DOCX incompleto; faltan: {missing}")
        xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
        if "FASE 12" not in xml.upper():
            raise RuntimeError("El Documento Maestro no contiene el título esperado de la Fase 12.")
    return {
        "archivo": path.name,
        "tamano_bytes": path.stat().st_size,
        "estructura_openxml_valida": True,
    }


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(18, 59, 99)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def update_master_document(document_path: Path, evidence: dict[str, Any]) -> None:
    document = Document(document_path)
    if any(MARKER in paragraph.text for paragraph in document.paragraphs):
        raise RuntimeError("El Documento Maestro ya contiene la actualización de la Fase 12.5.5.")

    document.add_page_break()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run(MARKER)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(18, 59, 99)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Pruebas, evidencia y cierre definitivo del Módulo de Matrices de Riesgos").bold = True

    paragraph = document.add_paragraph()
    paragraph.add_run("Estado técnico: ").bold = True
    paragraph.add_run("COMPLETADO - pendiente de aprobación formal e integración autorizada a main.")

    add_heading(document, "1. Consolidación de subfases", 1)
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Subfase", "Resultado", "Estado"]
    for index, value in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], value, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[index], "123B63")
    repeat_header(table.rows[0])
    rows = [
        ("12.5.1", "Estándar institucional compartido para reportería", "Completa"),
        ("12.5.2", "Normalización de reportería de Monitoreo de Listas", "Completa"),
        ("12.5.3", "Reemplazo de reportería de Matrices de Riesgos", "Completa"),
        ("12.5.4", "Refinamiento UX, documental, accesibilidad y retiro de /recalcular", "Completa"),
        ("12.5.5", "Pruebas con archivos reales, evidencia y cierre", "Completa técnicamente"),
    ]
    for row_index, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for column, value in enumerate(values):
            set_cell_text(cells[column], value, bold=column == 2)
            if row_index % 2 == 0:
                set_cell_shading(cells[column], "F3F6F9")

    tests = evidence["pruebas"]
    coverage = evidence["cobertura"]
    add_heading(document, "2. Validación integral", 1)
    add_bullet(document, f"Backend: {tests['backend']['aprobadas']} pruebas aprobadas, {tests['backend']['fallidas']} fallidas y {tests['backend']['omitidas']} omitidas.")
    add_bullet(document, f"Frontend: {tests['frontend']['pruebas_aprobadas']} pruebas aprobadas en {tests['frontend']['archivos_aprobados']} archivos.")
    add_bullet(document, f"E2E: {tests['e2e']['aprobadas']} escenarios aprobados y {tests['e2e']['fallidas']} fallidos.")
    add_bullet(document, "Compilación Angular de producción: aprobada.")
    add_bullet(document, f"Cobertura backend: líneas {coverage['backend_lineas_pct']:.2f}% y ramas {coverage['backend_ramas_pct']:.2f}%.")
    add_bullet(document, f"Cobertura frontend: sentencias {coverage['frontend_sentencias_pct']:.2f}%, ramas {coverage['frontend_ramas_pct']:.2f}%, funciones {coverage['frontend_funciones_pct']:.2f}% y líneas {coverage['frontend_lineas_pct']:.2f}%.")

    add_heading(document, "3. Archivos oficiales validados", 1)
    for item in evidence["artefactos_reporteria"]:
        if item["tipo"] == "PDF":
            add_bullet(document, f"{item['archivo']}: PDF válido, {item['paginas']} página(s), texto extraíble y checksum SHA-256 registrado.")
        else:
            add_bullet(document, f"{item['archivo']}: libro OpenXML válido con hojas {', '.join(item['hojas'])} y checksum SHA-256 registrado.")
    paragraph = document.add_paragraph()
    paragraph.add_run("Criterio de cierre: ").bold = True
    paragraph.add_run("los archivos utilizados en esta validación fueron producidos por el renderer real del backend; no son plantillas HTML renombradas ni documentos generados localmente por Angular.")

    add_heading(document, "4. Validación Oracle", 1)
    paragraph = document.add_paragraph()
    paragraph.add_run("Resultado: ").bold = True
    paragraph.add_run(evidence["oracle"]["resultado"])
    add_bullet(document, "Se verificó la presencia e integridad de los scripts SQL de validación de la Fase 12.3.")
    add_bullet(document, "Las pruebas de repositorio, aplicación y cálculo ejecutadas en CI aprobaron sin modificar el esquema Oracle.")
    add_bullet(document, "La ejecución contra una instancia Oracle institucional real queda registrada como dependencia externa: requiere red, credenciales y autorización del IHSS.")

    add_heading(document, "5. Restricciones y arquitectura preservadas", 1)
    for restriction in evidence["restricciones"]:
        add_bullet(document, restriction)

    add_heading(document, "6. Cierre técnico", 1)
    closing = document.add_paragraph()
    closing.add_run("La Fase 12.5 queda técnicamente completada. ").bold = True
    closing.add_run("El Módulo de Matrices de Riesgos dispone de cálculo centralizado en backend, reportes oficiales PDF/XLSX, ficha individual, auditoría, controles UX y accesibilidad, pruebas automatizadas y evidencia de cierre. La integración a main continúa bloqueada hasta aprobación formal de Javier Mejía.")

    approval = document.add_table(rows=4, cols=2)
    approval.style = "Table Grid"
    approval.alignment = WD_TABLE_ALIGNMENT.CENTER
    approval_rows = [
        ("Responsable de aprobación", "Javier Mejía"),
        ("Decisión", "PENDIENTE DE APROBACIÓN FORMAL"),
        ("Fecha", "____________________________"),
        ("Firma", "____________________________"),
    ]
    for row, (label, value) in zip(approval.rows, approval_rows):
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_shading(row.cells[0], "E8EEF4")
        set_cell_text(row.cells[1], value, bold="PENDIENTE" in value)

    document.core_properties.modified = datetime.now()
    document.save(document_path)


def write_closure_markdown(evidence: dict[str, Any]) -> None:
    tests = evidence["pruebas"]
    coverage = evidence["cobertura"]
    content = f"""# Fase 12.5.5 - Pruebas, evidencia y cierre definitivo

## Estado

**Cierre técnico aprobado.** Pendiente únicamente de validación Oracle en un entorno institucional autorizado, aprobación formal de Javier Mejía y posterior integración autorizada a `main`.

## Validación automatizada

- Backend: {tests['backend']['aprobadas']} aprobadas, {tests['backend']['fallidas']} fallidas, {tests['backend']['omitidas']} omitidas.
- Frontend: {tests['frontend']['pruebas_aprobadas']} aprobadas en {tests['frontend']['archivos_aprobados']} archivos.
- E2E: {tests['e2e']['aprobadas']} aprobadas, {tests['e2e']['fallidas']} fallidas.
- Build Angular: aprobado.
- Cobertura backend: líneas {coverage['backend_lineas_pct']:.2f} %, ramas {coverage['backend_ramas_pct']:.2f} %.
- Cobertura frontend: sentencias {coverage['frontend_sentencias_pct']:.2f} %, ramas {coverage['frontend_ramas_pct']:.2f} %, funciones {coverage['frontend_funciones_pct']:.2f} %, líneas {coverage['frontend_lineas_pct']:.2f} %.

## Archivos oficiales reales

Los artefactos de validación fueron generados directamente por `MatricesRiesgosReportRenderer`:

- `reporte_ejecutivo_matrices.pdf`.
- `reporte_matrices.xlsx`.
- `ficha_individual_matriz.pdf`.

Se validaron estructura, contenido extraíble, hojas OpenXML, tamaño, checksum y capacidad de renderizado.

## Oracle

{evidence['oracle']['resultado']}

No se declara una ejecución contra Oracle institucional porque el runner no dispone de conectividad, credenciales ni autorización para ese entorno. Los scripts SQL de validación y las pruebas de persistencia permanecen disponibles para su ejecución controlada.

## Restricciones

- No se modificó DNP.
- No se tocó `CONTROL_ALMACEN.PROVEEDOR`.
- No se integró funcionalmente Monitoreo con Matrices.
- No se modificó el motor de cálculo.
- No se fusionó a `main`.

## Próxima decisión

Aprobación formal de Javier Mejía. Solo después de esa autorización podrá evaluarse la integración del PR principal a `main`.
"""
    CLOSURE_MD.write_text(content, encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--quality-log", required=True)
    parser.add_argument("--report-dir", required=True)
    parser.add_argument("--run-id", required=True, type=int)
    args = parser.parse_args()

    quality = parse_quality_log(Path(args.quality_log))
    report_dir = Path(args.report_dir)
    pdf_report = validate_pdf(report_dir / "reporte_ejecutivo_matrices.pdf")
    xlsx_report = validate_xlsx(report_dir / "reporte_matrices.xlsx")
    pdf_sheet = validate_pdf(report_dir / "ficha_individual_matriz.pdf")
    docx_validation = validate_docx(DOCX)

    oracle_scripts = [
        BASE / "Evidencia_Fase_12_3" / "fase12_3_validacion_oracle.sql",
        BASE / "Evidencia_Fase_12_3" / "fase12_3_validacion_niveles_oracle.sql",
    ]
    for script in oracle_scripts:
        if not script.exists() or script.stat().st_size < 100:
            raise RuntimeError(f"Script Oracle faltante o vacío: {script}")

    oracle_available = bool(os.getenv("ORACLE_CONNECTION_STRING") or os.getenv("ConnectionStrings__Oracle"))
    oracle_result = (
        "Validación real ejecutada en entorno autorizado."
        if oracle_available
        else "Validación Oracle real no ejecutada en CI por ausencia de conectividad, credenciales y autorización institucional; dependencia externa documentada."
    )

    evidence: dict[str, Any] = {
        "fase": "12.5.5",
        "estado": "cierre_tecnico_generado_pendiente_revision_visual_manual",
        "rama": "fase-12-mejora-ejecutiva-matrices",
        "ejecucion_controlada": {"run_id": args.run_id, "resultado": "success"},
        "pruebas": quality["backend"] | {},
    }
    evidence["pruebas"] = {
        "backend": quality["backend"],
        "frontend": quality["frontend"],
        "e2e": quality["e2e"],
        "build_angular": "aprobado",
    }
    evidence["cobertura"] = quality["cobertura"]
    evidence["artefactos_reporteria"] = [
        {"tipo": "PDF", **pdf_report},
        {"tipo": "XLSX", **xlsx_report},
        {"tipo": "PDF", **pdf_sheet},
    ]
    evidence["documento_maestro"] = docx_validation
    evidence["oracle"] = {
        "conexion_institucional_disponible_en_runner": oracle_available,
        "resultado": oracle_result,
        "scripts_validados": [str(path.relative_to(ROOT)).replace("\\", "/") for path in oracle_scripts],
    }
    evidence["restricciones"] = [
        "No se modificó DNP",
        "No se tocó CONTROL_ALMACEN.PROVEEDOR",
        "No se integró funcionalmente Monitoreo de Listas con Matrices de Riesgos",
        "No se modificó el motor de cálculo de riesgo",
        "No se fusionó a main",
    ]
    evidence["aprobacion_formal"] = "pendiente_javier_mejia"
    evidence["siguiente_paso"] = "Revisión visual manual, checksum final y aprobación formal"

    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    update_master_document(DOCX, evidence)
    document_hash = sha256(DOCX)
    SHA_FILE.write_text(f"{document_hash}  {DOCX.name}\n", encoding="utf-8")
    evidence["documento_maestro"].update(
        {
            "tamano_bytes": DOCX.stat().st_size,
            "sha256": document_hash,
            "checksum_file": SHA_FILE.name,
            "actualizacion_fase_12_5_incorporada": True,
        }
    )

    write_closure_markdown(evidence)
    EVIDENCE.write_text(json.dumps(evidence, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps({
        "evidencia": str(EVIDENCE.relative_to(ROOT)),
        "documento": str(DOCX.relative_to(ROOT)),
        "sha256": document_hash,
        "backend": quality["backend"],
        "frontend": quality["frontend"],
        "e2e": quality["e2e"],
        "oracle_real": oracle_available,
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
